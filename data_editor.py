"""
資料編輯器與報告生成模組

包含：
- 物件資料編輯視窗
- Word 報告生成（買賣/租賃）
- 地圖影像生成
- Excel 報告生成
"""

# === 必要的 imports ===
import psutil
import subprocess
import threading
import tkinter as tk
from tkinter import font, scrolledtext, filedialog, messagebox, ttk
import tkinter.simpledialog as simpledialog
import re
import ctypes
from ctypes import wintypes
import os
import time
import json
from datetime import datetime
import sys
from config_manager import open_config_manager, load_config, save_config

# 🔥 基準目錄設定（data.json 和工作資料夾的位置）
from base_dir_helper import BASE_DIR, get_data_json_path, get_work_folder


def _get_with_ssl_fallback(url, **kwargs):
    """先嘗試正常 SSL 驗證，失敗才退到不驗證（政府網站偶有憑證鏈問題；避免下載失效）。"""
    import requests
    kwargs.pop('verify', None)
    try:
        return requests.get(url, verify=True, **kwargs)
    except requests.exceptions.SSLError:
        try:
            import urllib3
            urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        except Exception:
            pass
        return requests.get(url, verify=False, **kwargs)


# 🔥 處理 PyInstaller 打包後的路徑問題
def get_resource_path(relative_path):
    """
    獲取資源的絕對路徑，適用於開發環境和 PyInstaller 打包後的環境
    """
    try:
        # PyInstaller 會創建臨時資料夾，並將路徑存儲在 _MEIPASS 中
        base_path = sys._MEIPASS
    except Exception:
        # 如果不是打包環境，使用當前工作目錄
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# === 檢查 send2trash 模組 ===
try:
    import send2trash
    SEND2TRASH_AVAILABLE = True
except ImportError:
    SEND2TRASH_AVAILABLE = False
    print("警告：未安裝 send2trash，刪除檔案將直接刪除而非送入資源回收筒")
    print("建議執行：pip install send2trash")

# === 檢查 requests 模組 ===
try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False
    print("警告：未安裝 requests，無法使用學區查詢功能")
    print("請執行：pip install requests")

# === 檢查 python-docx 是否可用（lazy：只探測有沒有，不真的載入 module）===
# docx 載入很慢，實際 import 移到要用 Document() 的函式內部
try:
    import importlib.util
    DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None
except Exception:
    DOCX_AVAILABLE = False
if not DOCX_AVAILABLE:
    print("警告：未安裝 python-docx，無法生成 Word 報告")

# === 從 main 模組需要的全域變數和函數 ===
# 這些會在 main.py 中 import 此模組時設定
load_final_data = None
save_final_data = None
load_contacts_from_excel = None
load_urban_planning_data = None
get_city_from_district = None
extract_district_from_address = None
calculate_total_land_area = None
convert_transcript_to_final_data = None  # 轉換謄本資料為 data_final.json
update_message = None  # 訊息輸出函數
show_large_message = None  # 顯示大型訊息對話框
show_large_yesno = None  # 顯示 Yes/No 對話框
show_large_yesnocancel = None  # 顯示 Yes/No/Cancel 對話框


# === 備用函數：當無法從 main.py 載入時使用 ===
def _fallback_load_contacts_from_excel(file_path='通訊錄.xlsx'):
    """
    備用的通訊錄載入函數
    當 data_editor.py 被獨立執行時使用
    """
    result = {
        'stores': [],
        'contacts_by_store': {}
    }

    # 將錯誤訊息記錄到文件（已停用）
    def log_error(message):
        pass  # 不再寫入 log 檔案

    try:
        import pandas as pd
        from base_dir_helper import get_base_dir
        import os
        import sys

        # 使用 base_dir 來定位通訊錄檔案
        if not os.path.isabs(file_path):
            base_dir = get_base_dir()

            # 🔥 在打包環境下，通訊錄放在 _internal 目錄中
            if getattr(sys, 'frozen', False):
                # 打包環境：通訊錄在 _internal 子目錄
                file_path = os.path.join(base_dir, '_internal', file_path)
            else:
                # 開發環境：通訊錄在 base_dir 根目錄
                file_path = os.path.join(base_dir, file_path)

            log_error(f"通訊錄完整路徑：{file_path}")

        # 檢查檔案是否存在
        if not os.path.exists(file_path):
            error_msg = f"找不到通訊錄檔案：{file_path}"
            log_error(f"錯誤：{error_msg}")
            # 嘗試顯示錯誤訊息框（如果可用）
            try:
                import tkinter.messagebox as messagebox
                messagebox.showwarning("通訊錄載入失敗", error_msg)
            except:
                pass
            return result

        log_error(f"開始讀取通訊錄：{file_path}")
        # 讀取 Excel（第11行是標題，pandas索引為10）
        df = pd.read_excel(file_path, sheet_name=0, header=10)
        log_error(f"成功讀取 Excel，共 {len(df)} 行")
        log_error(f"Excel 欄位：{df.columns.tolist()}")

        current_store = None
        total_rows_processed = 0
        total_contacts_added = 0

        # 遍歷所有行
        for idx, row in df.iterrows():
            total_rows_processed += 1
            store_col = row.get('店別', '')
            title_col = row.get('職稱', '')
            name = row.get('姓名', '')
            phone = row.get('電話', '')
            code_col = row.get('編號', '')  # 🔥 讀取編號欄位

            # 檢查是否為店別標題行
            if pd.notna(store_col) and isinstance(store_col, str) and store_col.strip():
                store_name = store_col.strip()
                log_error(f"第 {idx+1} 行：發現店別標題 [{store_name}]")
                # 🔥 保留真正的店別 + 董事長（排除總公司管理部、店別標題）
                if store_name not in ['店別', '總公司管理部']:
                    current_store = store_name
                    if current_store not in result['stores']:
                        result['stores'].append(current_store)
                        result['contacts_by_store'][current_store] = []
                        log_error(f"  → 新增店別：{current_store}")
                else:
                    current_store = None
                    log_error(f"  → 跳過：{store_name}")

            # 🔥 新邏輯：只列出有編號（且編號不是「離職」）的聯絡人
            # 檢查是否有有效的編號
            clean_code = str(code_col).strip() if pd.notna(code_col) and code_col else ''
            has_valid_code = clean_code and clean_code not in ['離職', '成仁', '離退', '']

            # 檢查是否為有效聯絡人（phone 可能是 float 或 str）
            if pd.notna(name) and pd.notna(phone) and isinstance(name, str):
                if ('＜' not in name and '＞' not in name and
                    '姓名' not in name and '店別' not in name and
                    '：' not in name):

                    # 🔥 核心過濾：必須有有效編號才加入
                    if has_valid_code and current_store:
                        clean_name = name.strip()
                        # 🔥 處理電話號碼：如果是數字，補回開頭的 0
                        clean_phone = str(phone).strip()
                        if clean_phone and clean_phone[0] != '0' and clean_phone.replace('.', '').isdigit():
                            # 移除小數點後的 .0，並補上開頭的 0
                            clean_phone = '0' + clean_phone.split('.')[0]

                        result['contacts_by_store'][current_store].append({
                            'name': clean_name,
                            'phone': clean_phone,
                            'code': clean_code  # 儲存員工編號
                        })
                        total_contacts_added += 1
                        log_error(f"  → 加入聯絡人：{clean_name} ({clean_code})")
                    else:
                        # 記錄為什麼被過濾（用於除錯）
                        if not has_valid_code:
                            pass  # 沒有有效編號，不記錄
                        elif not current_store:
                            pass  # 無當前店別，不記錄

        log_error(f"=" * 40)
        log_error(f"處理統計：")
        log_error(f"  總處理行數：{total_rows_processed}")
        log_error(f"  新增聯絡人數：{total_contacts_added}")
        log_error(f"  找到店別數：{len(result['stores'])}")
        log_error(f"=" * 40)
        log_error(f"通訊錄載入完成，共找到 {len(result['stores'])} 個店別")
        for store in result['stores']:
            log_error(f"  {store}: {len(result['contacts_by_store'][store])} 位聯絡人")

        return result

    except ImportError as e:
        error_msg = f"無法導入 pandas：{e}"
        log_error(f"錯誤：{error_msg}")
        try:
            import tkinter.messagebox as messagebox
            messagebox.showwarning("通訊錄載入失敗", f"無法導入 pandas 模組\n\n{e}\n\n通訊錄功能將無法使用")
        except:
            pass
        return result
    except Exception as e:
        error_msg = f"載入通訊錄時發生錯誤：{e}"
        log_error(f"錯誤：{error_msg}")
        import traceback
        log_error(f"詳細錯誤：\n{traceback.format_exc()}")
        try:
            import tkinter.messagebox as messagebox
            messagebox.showwarning("通訊錄載入失敗", f"載入通訊錄時發生錯誤\n\n{e}")
        except:
            pass
        return result
root = None  # 主視窗
get_data_final_path = None  # 取得 data_final.json 路徑的函數
get_selected_json_path = None  # 取得已選擇的 JSON 檔案路徑的函數
get_screenshot_path = None  # 取得截圖檔案路徑的函數
report_type = None
entries = {}  # 🔥 編輯器中所有輸入欄位的全域字典

# === 報告生成和編輯器函數 ===

def calculate_rent_with_tax(net_income_amount):
    """
    計算租金含稅金額（台灣稅法 - 情境B：反推稅前總額）

    Args:
        net_income_amount: 業主實收金額（萬元）

    Returns:
        tax_included_amount: 房客應付總額（元，精確數字）

    情境說明：
    - 業主想實收 X 元（稅後淨收入）
    - 房客需支付 Y 元（稅前總額）
    - 稅金從 Y 中扣除後，業主實收 X 元
    - 計算公式：Y = X ÷ (1 - 稅率)

    台灣租金所得稅法：
    1. 所得稅扣繳：月租金 >= 20,000 元，扣繳率 10%
    2. 二代健保補充保費：單次給付 >= 20,001 元，費率 2.11%
    """
    try:
        # 轉換為數字（萬元 -> 元）
        net_income = float(net_income_amount) * 10000

        # 判斷適用的稅率
        if net_income >= 20001:
            # 兩種稅都要扣：10% + 2.11% = 12.11%
            # 含稅金額 = 淨收入 ÷ (1 - 0.1211) = 淨收入 ÷ 0.8789
            tax_included = net_income / 0.8789
        elif net_income >= 20000:
            # 只扣所得稅：10%
            # 含稅金額 = 淨收入 ÷ (1 - 0.10) = 淨收入 ÷ 0.90
            tax_included = net_income / 0.90
        else:
            # 低於門檻，不需扣稅
            tax_included = net_income

        # 返回精確的元數（四捨五入到整數）
        return round(tax_included)

    except (ValueError, TypeError):
        return None

def generate_report_from_button():
    """從按鈕手動生成報告"""
    global report_type, entries

    # 檢查是否安裝 python-docx
    if not DOCX_AVAILABLE:
        messagebox.showerror(
            "缺少套件",
            "未安裝 python-docx 套件\n\n請在命令提示字元執行：\npip install python-docx"
        )
        return

    # 檢查是否有資料
    data_final_path = get_data_final_path()
    if not os.path.exists(data_final_path):
        result = show_large_yesno(
            "找不到資料",
            "尚未建立物調資料 (data_final.json)\n\n是否要讀取 JSON 並轉換？"
        )
        if result:
            choose_json_file()
        return

    # 🔥 直接從編輯器中的交易類型判斷，不再詢問
    try:
        # 取得UI當前選擇的交易類型
        transaction_type = "售件"  # 預設值

        if "交易類型" in entries:
            transaction_type = entries["交易類型"].get()
            if not transaction_type:  # 如果為空，使用預設值
                transaction_type = "售件"
        else:
            transaction_type = "售件"

        update_message(f"準備生成【{transaction_type}】報告...")

        # 設定報告類型
        if transaction_type == "租件":
            report_type = "rent"
        else:  # 預設為售件
            report_type = "buy"

    except Exception as e:
        update_message(f"讀取交易類型失敗：{e}，使用預設值「售件」")
        print(f"[DEBUG] 讀取交易類型發生異常: {e}")
        # 發生錯誤時，直接使用售件作為預設值
        report_type = "buy"
        update_message("準備生成【售件】報告...")

    # 生成報告
    generate_property_report()

def generate_property_report():
    """生成物件個案調查表 Word 文件"""
    global report_type
    
    try:
        update_message("=" * 50)
        update_message("開始生成報告...")

        # 讀取 data_final.json
        data_final_path = get_data_final_path()
        if not os.path.exists(data_final_path):
            update_message("[錯誤] 找不到 data_final.json")
            return

        with open(data_final_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 取得基本資訊來決定檔案名稱
        土地標示 = data.get('土地標示', '')

        # 🔥 從 data.json 讀取第一筆地號
        地號 = ''
        try:
            if os.path.exists(get_data_json_path()):
                with open(get_data_json_path(), 'r', encoding='utf-8') as f:
                    data_list = json.load(f)
                if data_list and len(data_list) > 0:
                    first_data = data_list[0]
                    地號 = first_data.get('lot_number', '')
                    print(f"[檔名] 從 data.json 讀取第一筆地號: {地號}")
        except Exception as e:
            print(f"[提示] 無法讀取 data.json 地號：{e}")

        # 簡化檔名
        if '市' in 土地標示:
            parts = 土地標示.split('市')[1].strip()
            if '段' in parts:
                簡化名稱 = parts.split('段')[0] + '段'
            else:
                簡化名稱 = parts[:20]
        else:
            簡化名稱 = '物件'

        # 生成今天的日期
        today = datetime.now().strftime("%Y年%m月%d日")

        # 根據報告類型生成不同格式
        report_suffix = "售" if report_type == "buy" else "租"

        # 確保資料夾存在
        # 🔥 使用 get_work_folder 確保在主程式目錄下建立資料夾
        output_folder = get_work_folder("報告輸出")
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        # 🔥 刪除舊版本報告（同一物件的舊報告）
        try:
            import glob
            # 根據地號和簡化名稱搜尋舊報告
            if 地號:
                pattern = f"物件個案調查表({report_suffix})_{簡化名稱}-{地號}_*.docx"
            else:
                pattern = f"物件個案調查表({report_suffix})_{簡化名稱}_*.docx"

            old_reports = glob.glob(os.path.join(output_folder, pattern))
            for old_report in old_reports:
                try:
                    os.remove(old_report)
                    update_message(f"[刪除] 已刪除舊版本報告：{os.path.basename(old_report)}")
                except Exception as e:
                    update_message(f"[警告] 無法刪除舊報告 {os.path.basename(old_report)}：{e}")
        except Exception as e:
            update_message(f"[提示] 搜尋舊報告時發生錯誤：{e}")

        # 🔥 生成 Word 文件路徑（地段-地號_日期）
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if 地號:
            word_filename = f"物件個案調查表({report_suffix})_{簡化名稱}-{地號}_{timestamp}.docx"
        else:
            word_filename = f"物件個案調查表({report_suffix})_{簡化名稱}_{timestamp}.docx"
        word_path = os.path.join(output_folder, word_filename)
        
        update_message(f"正在生成 {report_suffix} 件報告...")
        
        if report_type == "buy":
            generate_word_report_buy(word_path, data, today)
        else:
            generate_word_report_rent(word_path, data, today)

        update_message(f"[完成] 已生成報告：{word_path}")

        # 🔥 複製一份報告到建立的資料夾中
        try:
            # 取得建立的資料夾路徑
            if os.path.exists(get_data_json_path()):
                with open(get_data_json_path(), 'r', encoding='utf-8') as f:
                    data_list = json.load(f)
                    if data_list:
                        first_data = data_list[0]
                        area = first_data.get('area', '')
                        section = first_data.get('section', '')
                        lot_number = first_data.get('lot_number', '')
                        # 🔥 使用 get_work_folder 確保在主程式目錄下建立資料夾
                        folder_name = get_work_folder(f"{area}{section}-{lot_number}")

                        # 確保資料夾存在
                        if os.path.exists(folder_name):
                            import shutil
                            import glob

                            # 🔥 先刪除該資料夾中的舊版本報告
                            if 地號:
                                pattern = f"物件個案調查表({report_suffix})_{簡化名稱}-{地號}_*.docx"
                            else:
                                pattern = f"物件個案調查表({report_suffix})_{簡化名稱}_*.docx"

                            old_reports_in_folder = glob.glob(os.path.join(folder_name, pattern))
                            for old_report in old_reports_in_folder:
                                try:
                                    os.remove(old_report)
                                    update_message(f"[刪除] 已刪除資料夾中的舊報告：{os.path.basename(old_report)}")
                                except Exception as e:
                                    update_message(f"[警告] 無法刪除資料夾中的舊報告：{e}")

                            # 複製到建立的資料夾根目錄
                            dest_path = os.path.join(folder_name, word_filename)
                            shutil.copy2(word_path, dest_path)
                            update_message(f"[完成] 已複製報告到：{dest_path}")
                        else:
                            update_message("[提示] 建立的資料夾不存在，僅保存在報告輸出資料夾")
        except Exception as e:
            update_message(f"[警告] 複製報告到資料夾失敗：{e}")

        update_message("=" * 50)

        # 詢問是否開啟
        result = show_large_yesno(
            "報告生成完成",
            f"報告已生成：\n{word_filename}\n\n是否要立即開啟？"
        )

        if result:
            try:
                os.startfile(word_path)
                update_message("[完成] 已開啟報告")
            except Exception as e:
                update_message(f"[錯誤] 開啟失敗：{e}")
        
    except Exception as e:
        update_message(f"[錯誤] 生成報告時發生錯誤：{e}")
        import traceback
        update_message(traceback.format_exc())
        messagebox.showwarning("警告", f"生成報告失敗：{e}")

def format_address(address):
    """格式化建物門牌地址（全形轉半形、中文數字轉阿拉伯數字）"""
    if not address:
        return ""

    # 移除「市」之前的內容
    if '市' in address:
        address = address.split('市', 1)[1]

    # 全形數字對照表
    full_to_half = {
        '０': '0', '１': '1', '２': '2', '３': '3', '４': '4',
        '５': '5', '６': '6', '７': '7', '８': '8', '９': '9',
    }

    # 中文數字對照表（樓層）
    chinese_to_num = {
        '一': '1', '二': '2', '三': '3', '四': '4', '五': '5',
        '六': '6', '七': '7', '八': '8', '九': '9', '十': '10',
        '十一': '11', '十二': '12', '十三': '13', '十四': '14', '十五': '15',
        '十六': '16', '十七': '17', '十八': '18', '十九': '19', '二十': '20',
        '二十一': '21', '二十二': '22', '二十三': '23', '二十四': '24', '二十五': '25',
        '二十六': '26', '二十七': '27', '二十八': '28', '二十九': '29', '三十': '30',
    }

    # 全形轉半形
    for full, half in full_to_half.items():
        address = address.replace(full, half)

    # 處理樓層：「號七樓」→「號7樓」
    for chinese, num in chinese_to_num.items():
        address = address.replace(f'號{chinese}樓', f'號{num}樓')

    return address.strip()

def export_archive(force_reset=False):
    """匯出封存資料夾（智慧模式）

    Args:
        force_reset: 是否強制重新設定（Shift+點擊時為 True）
    """
    try:
        update_message("=" * 50)
        update_message("開始智慧匯出封存...")

        # 檢查是否有 data.json
        if not os.path.exists(get_data_json_path()):
            messagebox.showerror("錯誤", "找不到 data.json，無法確定資料夾名稱")
            return

        # 讀取 data.json 取得地段地號
        with open(get_data_json_path(), 'r', encoding='utf-8') as f:
            data_list = json.load(f)

        if not data_list:
            messagebox.showerror("錯誤", "data.json 中沒有資料")
            return

        first_data = data_list[0]
        area = first_data.get('area', '')
        section = first_data.get('section', '')
        lot_number = first_data.get('lot_number', '')

        # 工作區資料夾名稱
        # 🔥 使用 get_work_folder 確保在主程式目錄下建立資料夾
        source_folder = get_work_folder(f"{area}{section}-{lot_number}")

        if not os.path.exists(source_folder):
            messagebox.showerror("錯誤", f"找不到工作區資料夾：{source_folder}")
            return

        update_message(f"工作區資料夾：{source_folder}")

        # 🔥 讀取上次的匯出設定（智慧模式核心）
        config = load_config()
        export_settings = config.get('export_settings', {})

        # 檢查是否為同一資料夾
        last_folder = export_settings.get('current_folder', '')
        is_same_folder = (last_folder == source_folder)
        has_previous_export = bool(export_settings.get('last_export_path'))

        # 決定使用模式
        use_smart_mode = (is_same_folder and has_previous_export and not force_reset)

        if use_smart_mode:
            update_message("[模式] 智慧更新模式")
        else:
            update_message("[模式] 首次匯出模式")
            export_settings = {}  # 清空設定

        # 檢查並補足子資料夾
        required_subfolders = [
            "0.謄本",
            "1.基本資料",
            "2.照片",
            "3.行情",
            "4.其他相關",
            "5.委託",
            "6.成交"
        ]

        for subfolder in required_subfolders:
            subfolder_path = os.path.join(source_folder, subfolder)
            if not os.path.exists(subfolder_path):
                os.makedirs(subfolder_path)
                update_message(f"[建立] {subfolder}")
            else:
                update_message(f"[存在] {subfolder}")

        # 讀取 data_final.json 取得建物門牌
        data_final_path = get_data_final_path()
        門牌 = ""
        if os.path.exists(data_final_path):
            with open(data_final_path, 'r', encoding='utf-8') as f:
                data_final = json.load(f)
                門牌_原始 = data_final.get('建物門牌', '')
                if 門牌_原始:
                    門牌 = format_address(門牌_原始)
                    update_message(f"建物門牌（原始）：{門牌_原始}")
                    update_message(f"建物門牌（格式化）：{門牌}")

        # 🔥 智慧模式邏輯 - 檢查現有封存並決定處理方式
        update_mode = None  # 更新模式
        selected_format = None  # 匯出格式
        dest_parent = None  # 匯出位置
        archive_name = None  # 封存名稱

        if use_smart_mode:
            # 從設定中取得上次的資訊
            last_export_path = export_settings.get('last_export_path', '')
            last_export_format = export_settings.get('last_export_format', 'zip')
            last_archive_name = export_settings.get('last_archive_name', '')
            auto_export = export_settings.get('auto_export', False)

            # 🔥 保存搜尋封存的函數，但延後執行（在使用者選擇「更新日期」之後才清理舊檔）
            # 這裡先不執行清理動作

            # 構建完整路徑檢查檔案是否存在
            if last_export_format in ['folder', 'both']:
                last_full_path = os.path.join(last_export_path, last_archive_name)
            else:
                last_full_path = os.path.join(last_export_path, last_archive_name + '.zip')

            archive_exists = os.path.exists(last_full_path)

            if archive_exists:
                # 取得檔案資訊
                if os.path.isdir(last_full_path):
                    total_size = sum(
                        os.path.getsize(os.path.join(dirpath, filename))
                        for dirpath, dirnames, filenames in os.walk(last_full_path)
                        for filename in filenames
                    )
                    mod_time = os.path.getmtime(last_full_path)
                else:
                    total_size = os.path.getsize(last_full_path)
                    mod_time = os.path.getmtime(last_full_path)

                size_mb = total_size / (1024 * 1024)
                mod_date = datetime.fromtimestamp(mod_time).strftime('%Y-%m-%d %H:%M')

                update_message(f"[發現] 現有封存：{last_archive_name}")
                update_message(f"  位置：{last_export_path}")
                update_message(f"  大小：{size_mb:.2f} MB")

                if auto_export:
                    # 自動模式：直接更新，保持日期
                    update_message("[自動] 更新內容，保持日期")
                    update_mode = "keep_date"
                    selected_format = last_export_format
                    dest_parent = last_export_path
                    archive_name = last_archive_name
                else:
                    # 顯示智慧更新對話框
                    from datetime import datetime as dt
                    today_obj = dt.now()
                    民國年 = today_obj.year - 1911
                    今天日期 = f"{民國年}.{today_obj.month:02d}.{today_obj.day:02d}"

                    update_result = show_smart_update_dialog(
                        last_archive_name, last_export_path, size_mb, mod_date,
                        area, section, lot_number, 門牌, 今天日期
                    )

                    if update_result is None:
                        update_message("[取消] 使用者取消匯出")
                        return

                    update_mode, auto_export_choice = update_result

                    if auto_export_choice:
                        export_settings['auto_export'] = True
                        config['export_settings'] = export_settings
                        save_config(config)
                        update_message("[設定] 已記住選擇")

                    selected_format = last_export_format
                    dest_parent = last_export_path

                    if update_mode == "keep_date":
                        archive_name = last_archive_name
                        update_message("[選擇] 保持原日期")
                    elif update_mode == "update_date":
                        # 🔥 產生新的封存名稱
                        today = datetime.now()
                        民國年 = today.year - 1911
                        新日期 = f"{民國年}.{today.month:02d}.{today.day:02d}"
                        if 門牌:
                            new_archive_name = f"{新日期}{area}{section}-{lot_number},{門牌}"
                        else:
                            new_archive_name = f"{新日期}{area}{section}-{lot_number}"
                        update_message(f"[選擇] 更新為今天：{新日期}")

                        # 🔥 搜尋所有相關的舊封存（相同地段地號但不同日期）
                        old_archives_to_delete = []
                        pattern_prefix = f"{area}{section}-{lot_number}"

                        try:
                            # 列出目標目錄中所有項目
                            for item in os.listdir(dest_parent):
                                item_path = os.path.join(dest_parent, item)
                                # 檢查是否為資料夾或ZIP檔
                                if os.path.isdir(item_path) or item.endswith('.zip'):
                                    # 移除 .zip 副檔名以比較
                                    item_name = item[:-4] if item.endswith('.zip') else item
                                    # 檢查是否包含地段地號且不是新名稱
                                    if pattern_prefix in item_name and item_name != new_archive_name:
                                        # 檢查是否符合日期格式（民國年.月.日）
                                        if re.match(r'^\d+\.\d+\.\d+', item_name):
                                            old_archives_to_delete.append(item_path)
                        except Exception as e:
                            update_message(f"[警告] 搜尋舊封存時發生錯誤：{e}")

                        # 如果找到舊封存，詢問使用者是否刪除
                        if old_archives_to_delete:
                            old_list = "\n".join([os.path.basename(p) for p in old_archives_to_delete])
                            result = show_large_yesnocancel(
                                "發現舊封存",
                                f"舊封存（{len(old_archives_to_delete)}個）：\n{old_list}\n\n"
                                f"最新：{new_archive_name}\n\n"
                                f"刪除舊封存？"
                            )

                            if result is None:
                                # 使用者選擇取消
                                update_message("[取消] 使用者取消匯出")
                                return
                            elif result:
                                # 使用者選擇刪除
                                import shutil
                                import stat
                                import time

                                def safe_send_to_trash(path, max_retries=3):
                                    """安全地將檔案送到回收筒，支援重試機制"""
                                    for attempt in range(max_retries):
                                        try:
                                            if SEND2TRASH_AVAILABLE:
                                                send2trash.send2trash(path)
                                                return True, "回收筒"
                                            else:
                                                # 傳統刪除方式
                                                if os.path.isdir(path):
                                                    def remove_readonly(func, filepath, excinfo):
                                                        os.chmod(filepath, stat.S_IWRITE)
                                                        func(filepath)
                                                    shutil.rmtree(path, onerror=remove_readonly)
                                                else:
                                                    os.remove(path)
                                                return True, "刪除"
                                        except Exception as e:
                                            if attempt < max_retries - 1:
                                                # 還有重試機會，等待後重試
                                                time.sleep(1.0)
                                                continue
                                            else:
                                                # 最後一次嘗試失敗
                                                return False, str(e)
                                    return False, "未知錯誤"

                                deleted_successfully = []
                                for old_path in old_archives_to_delete:
                                    basename = os.path.basename(old_path)
                                    update_message(f"[處理] 正在刪除：{basename}")
                                    success, msg = safe_send_to_trash(old_path, max_retries=3)

                                    if success:
                                        update_message(f"[{msg}] 已移至{msg}：{basename}")
                                        deleted_successfully.append(old_path)
                                        # 每刪除一個檔案後等待，讓 Dropbox 同步
                                        time.sleep(0.5)
                                    else:
                                        update_message(f"[警告] 無法刪除 {basename}：{msg}")

                                # 等待檔案系統完成操作（特別是 Dropbox 同步）
                                if deleted_successfully:
                                    update_message("[等待] 等待檔案系統同步...")
                                    time.sleep(2.0)

                                # 記錄已刪除的路徑，避免後續重複處理
                                export_settings['deleted_archives'] = deleted_successfully
                            else:
                                update_message("[保留] 使用者選擇保留舊封存")

                        archive_name = new_archive_name
                    elif update_mode == "create_new":
                        today = datetime.now()
                        民國年 = today.year - 1911
                        日期 = f"{民國年}.{today.month:02d}.{today.day:02d}"
                        時間戳記 = today.strftime("%H%M")
                        if 門牌:
                            archive_name = f"{日期}{area}{section}-{lot_number},{門牌}_{時間戳記}"
                        else:
                            archive_name = f"{日期}{area}{section}-{lot_number}_{時間戳記}"
                        update_message("[選擇] 建立新封存")
                    else:  # reset
                        use_smart_mode = False
                        export_settings = {}
                        update_message("[選擇] 重新設定")
            else:
                # 檔案不存在，使用上次設定重建
                update_message("[警告] 上次封存不存在，將重新建立")
                selected_format = last_export_format
                dest_parent = last_export_path
                today = datetime.now()
                民國年 = today.year - 1911
                日期 = f"{民國年}.{today.month:02d}.{today.day:02d}"
                if 門牌:
                    archive_name = f"{日期}{area}{section}-{lot_number},{門牌}"
                else:
                    archive_name = f"{日期}{area}{section}-{lot_number}"

                # 🔥 搜尋是否有其他封存（包括較新的），並識別最新的
                all_related_archives = []
                pattern_prefix = f"{area}{section}-{lot_number}"

                try:
                    for item in os.listdir(dest_parent):
                        item_path = os.path.join(dest_parent, item)
                        if os.path.isdir(item_path) or item.endswith('.zip'):
                            item_name = item[:-4] if item.endswith('.zip') else item
                            if pattern_prefix in item_name:
                                if re.match(r'^\d+\.\d+\.\d+', item_name):
                                    all_related_archives.append((item_path, item_name))
                except Exception as e:
                    update_message(f"[警告] 搜尋舊封存時發生錯誤：{e}")

                # 如果找到其他封存，識別最新的並詢問是否刪除舊的
                if all_related_archives:
                    # 🔥 使用日期比較找出最新的封存
                    def extract_date(base_name):
                        """從封存名稱中提取日期並轉換為可比較的數字"""
                        match = re.match(r'^(\d+)\.(\d+)\.(\d+)', base_name)
                        if match:
                            year, month, day = match.groups()
                            return int(year) * 10000 + int(month) * 100 + int(day)
                        return 0

                    sorted_archives = sorted(all_related_archives, key=lambda x: extract_date(x[1]), reverse=True)
                    newest_base_name = sorted_archives[0][1]

                    # 將非最新的封存標記為舊封存
                    old_archives_to_delete = [path for path, name in all_related_archives if name != newest_base_name]

                    if old_archives_to_delete:
                        old_list = "\n".join([os.path.basename(p) for p in old_archives_to_delete])
                        result = show_large_yesno(
                            "發現舊封存",
                            f"舊封存（{len(old_archives_to_delete)}個）：\n{old_list}\n\n"
                            f"最新：{newest_base_name}\n\n"
                            f"刪除舊封存？"
                        )

                        if result:
                            # 使用者選擇刪除
                            import shutil
                            import stat
                            import time

                            def safe_send_to_trash(path, max_retries=3):
                                """安全地將檔案送到回收筒，支援重試機制"""
                                for attempt in range(max_retries):
                                    try:
                                        if SEND2TRASH_AVAILABLE:
                                            send2trash.send2trash(path)
                                            return True, "回收筒"
                                        else:
                                            if os.path.isdir(path):
                                                def remove_readonly(func, filepath, excinfo):
                                                    os.chmod(filepath, stat.S_IWRITE)
                                                    func(filepath)
                                                shutil.rmtree(path, onerror=remove_readonly)
                                            else:
                                                os.remove(path)
                                            return True, "刪除"
                                    except Exception as e:
                                        if attempt < max_retries - 1:
                                            time.sleep(1.0)
                                            continue
                                        else:
                                            return False, str(e)
                                return False, "未知錯誤"

                            deleted_successfully = []
                            for old_path in old_archives_to_delete:
                                basename = os.path.basename(old_path)
                                update_message(f"[處理] 正在刪除：{basename}")
                                success, msg = safe_send_to_trash(old_path, max_retries=3)

                                if success:
                                    update_message(f"[{msg}] 已移至{msg}：{basename}")
                                    deleted_successfully.append(old_path)
                                    time.sleep(0.5)
                                else:
                                    update_message(f"[警告] 無法刪除 {basename}：{msg}")

                            if deleted_successfully:
                                update_message("[等待] 等待檔案系統同步...")
                                time.sleep(2.0)

                            export_settings['deleted_archives'] = deleted_successfully
                        else:
                            update_message("[保留] 使用者選擇保留舊封存")

        # 首次匯出模式：需要選擇格式和位置
        if not use_smart_mode or update_mode == "reset":
            # 生成新的封存名稱
            today = datetime.now()
            民國年 = today.year - 1911
            日期 = f"{民國年}.{today.month:02d}.{today.day:02d}"

            if 門牌:
                archive_name = f"{日期}{area}{section}-{lot_number},{門牌}"
            else:
                archive_name = f"{日期}{area}{section}-{lot_number}"

            update_message(f"封存資料夾名稱：{archive_name}")

            # 🔥 選擇匯出格式
            format_dialog = tk.Toplevel()
            format_dialog.title("選擇匯出格式")
            # 🔥 調整視窗大小以適應高 DPI（500x300 → 600x420）
            dialog_width = 600
            dialog_height = 420
            format_dialog.geometry(f"{dialog_width}x{dialog_height}")
            format_dialog.transient(root if 'root' in globals() else None)
            format_dialog.grab_set()

            # 置中顯示
            screen_width = format_dialog.winfo_screenwidth()
            screen_height = format_dialog.winfo_screenheight()
            x = (screen_width - dialog_width) // 2
            y = (screen_height - dialog_height) // 2
            format_dialog.geometry(f"{dialog_width}x{dialog_height}+{x}+{y}")

            tk.Label(format_dialog, text="請選擇匯出格式：",
                    font=("Microsoft JhengHei", 14, "bold")).pack(pady=20)

            export_format = tk.StringVar(value="zip")

            tk.Radiobutton(format_dialog, text="資料夾（方便查閱）",
                          variable=export_format, value="folder",
                          font=("Microsoft JhengHei", 12)).pack(pady=10, anchor="w", padx=50)

            tk.Radiobutton(format_dialog, text="資料夾 + ZIP 壓縮檔（兼顧兩者）",
                          variable=export_format, value="both",
                          font=("Microsoft JhengHei", 12)).pack(pady=10, anchor="w", padx=50)

            tk.Radiobutton(format_dialog, text="只建立 ZIP 壓縮檔（節省空間）",
                          variable=export_format, value="zip",
                          font=("Microsoft JhengHei", 12)).pack(pady=10, anchor="w", padx=50)

            btn_frame = tk.Frame(format_dialog)
            btn_frame.pack(pady=20)

            format_selected = [False]

            def on_confirm():
                format_selected[0] = True
                format_dialog.destroy()

            def on_cancel():
                format_dialog.destroy()

            tk.Button(btn_frame, text="確定", command=on_confirm,
                     font=("Microsoft JhengHei", 12, "bold"),
                     bg='#4CAF50', fg='white', width=10, height=2).pack(side=tk.LEFT, padx=10)

            tk.Button(btn_frame, text="取消", command=on_cancel,
                     font=("Microsoft JhengHei", 12),
                     bg='#F44336', fg='white', width=10, height=2).pack(side=tk.LEFT, padx=10)

            format_dialog.wait_window()

            if not format_selected[0]:
                update_message("[取消] 使用者取消選擇格式")
                return

            selected_format = export_format.get()
            update_message(f"選擇的格式：{selected_format}")

            # 選擇目的地
            from tkinter import filedialog
            dest_parent = filedialog.askdirectory(title="選擇封存位置")

            if not dest_parent:
                update_message("[取消] 使用者取消選擇位置")
                return

        dest_path = os.path.join(dest_parent, archive_name)
        zip_path = dest_path + ".zip"

        # 檢查是否剛剛已經將此檔案送到回收筒
        deleted_archives = export_settings.get('deleted_archives', [])
        already_deleted = dest_path in deleted_archives

        # 檢查目的地是否已存在（但跳過剛剛已送到回收筒的）
        if os.path.exists(dest_path) and not already_deleted:
            result = show_large_yesno(
                "資料夾已存在",
                f"目的地已存在相同名稱的資料夾：\n{archive_name}\n\n是否要覆蓋？"
            )
            if not result:
                update_message("[取消] 使用者取消覆蓋")
                return
            else:
                # 🔥 使用更可靠的方式刪除舊資料夾（處理 Windows 權限問題）
                import shutil
                import stat
                import time

                def remove_readonly(func, path, excinfo):
                    """移除唯讀屬性後重試刪除"""
                    os.chmod(path, stat.S_IWRITE)
                    func(path)

                try:
                    update_message("[刪除] 正在刪除舊資料夾...")
                    shutil.rmtree(dest_path, onerror=remove_readonly)
                    # 等待一下確保檔案系統完成刪除
                    time.sleep(0.5)
                    update_message("[刪除] 已刪除舊資料夾")
                except Exception as e:
                    update_message(f"[錯誤] 刪除舊資料夾失敗：{e}")
                    messagebox.showerror("錯誤", f"無法刪除舊資料夾，可能被其他程式占用：\n\n{e}")
                    return
        elif already_deleted:
            update_message(f"[略過] 檔案已送到回收筒，略過檢查：{archive_name}")

        # 根據選擇的格式進行匯出
        import shutil
        import zipfile
        from tkinter import ttk

        files_created = []

        # 建立進度條視窗
        progress_window = tk.Toplevel()
        progress_window.title("匯出進度")
        progress_window.geometry("500x150")
        progress_window.transient(root if 'root' in globals() else None)
        progress_window.grab_set()

        # 置中顯示
        screen_width = progress_window.winfo_screenwidth()
        screen_height = progress_window.winfo_screenheight()
        x = (screen_width - 500) // 2
        y = (screen_height - 150) // 2
        progress_window.geometry(f"500x150+{x}+{y}")

        progress_label = tk.Label(progress_window, text="準備中...", font=("Microsoft JhengHei", 12))
        progress_label.pack(pady=20)

        progress_bar = ttk.Progressbar(progress_window, length=400, mode='determinate')
        progress_bar.pack(pady=10)

        progress_detail = tk.Label(progress_window, text="", font=("Microsoft JhengHei", 9), fg="gray")
        progress_detail.pack(pady=5)

        progress_window.update()

        # 1. 複製資料夾（folder 或 both）
        if selected_format in ["folder", "both"]:
            update_message("開始複製資料夾...")
            progress_label.config(text="正在複製資料夾...")
            progress_window.update()

            # 計算總檔案數
            total_files = sum(len(files) for _, _, files in os.walk(source_folder))
            progress_bar['maximum'] = total_files
            copied_files = 0

            def copy_with_progress(src, dst, *, follow_symlinks=True):
                nonlocal copied_files
                # 🔥 跳過 Excel 暫存檔（~$ 開頭的檔案）
                filename = os.path.basename(src)
                if filename.startswith('~$'):
                    return  # 跳過此檔案
                shutil.copy2(src, dst, follow_symlinks=follow_symlinks)
                copied_files += 1
                progress_bar['value'] = copied_files
                progress_detail.config(text=f"{copied_files}/{total_files} 個檔案")
                if copied_files % 10 == 0:  # 每 10 個檔案更新一次畫面
                    progress_window.update()

            def ignore_temp_files(dir, files):
                """忽略 Excel 暫存檔"""
                return [f for f in files if f.startswith('~$')]

            shutil.copytree(source_folder, dest_path, copy_function=copy_with_progress, ignore=ignore_temp_files)
            update_message(f"[完成] 資料夾已匯出")
            files_created.append(("資料夾", dest_path))

        # 2. 建立 ZIP 壓縮檔（zip 或 both）
        if selected_format in ["zip", "both"]:
            update_message("開始建立 ZIP 壓縮檔...")
            progress_label.config(text="正在建立 ZIP 壓縮檔...")
            progress_window.update()

            # 檢查 ZIP 是否已存在
            if os.path.exists(zip_path):
                try:
                    os.remove(zip_path)
                    update_message("[刪除] 已刪除舊的 ZIP 檔案")
                except Exception as e:
                    update_message(f"[警告] 無法刪除舊 ZIP：{e}")

            try:
                # 計算總檔案數（排除 Excel 暫存檔），並收集空目錄
                all_files = []
                empty_dirs = []
                for root_dir, dirs, files in os.walk(source_folder):
                    filtered_files = [f for f in files if not f.startswith('~$')]
                    for file in filtered_files:
                        all_files.append(os.path.join(root_dir, file))
                    # 如果該目錄沒有子目錄也沒有檔案（或檔案全被過濾），記為空目錄
                    if not dirs and not filtered_files:
                        empty_dirs.append(root_dir)

                total_files = len(all_files)
                progress_bar['maximum'] = total_files
                progress_bar['value'] = 0

                with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    # 寫入空目錄（以斜線結尾的路徑）
                    for empty_dir in empty_dirs:
                        arcname = os.path.join(
                            archive_name,
                            os.path.relpath(empty_dir, source_folder)
                        ) + '/'
                        zipf.writestr(arcname, '')

                    for idx, file_path in enumerate(all_files, 1):
                        # 計算相對路徑
                        arcname = os.path.join(
                            archive_name,
                            os.path.relpath(file_path, source_folder)
                        )
                        zipf.write(file_path, arcname)

                        progress_bar['value'] = idx
                        progress_detail.config(text=f"{idx}/{total_files} 個檔案")
                        if idx % 10 == 0:  # 每 10 個檔案更新一次畫面
                            progress_window.update()

                update_message(f"[完成] ZIP 壓縮檔已建立")
                files_created.append(("ZIP檔", zip_path))

            except Exception as e:
                progress_window.destroy()
                update_message(f"[錯誤] 建立 ZIP 失敗：{e}")
                messagebox.showerror("錯誤", f"建立 ZIP 壓縮檔失敗：\n\n{e}")
                return

        # 關閉進度條視窗
        progress_window.destroy()

        update_message("=" * 50)
        update_message("匯出結果：")
        for file_type, file_path in files_created:
            update_message(f"  ✓ {file_type}：{os.path.basename(file_path)}")

        # 🔥 儲存匯出設定（智慧模式）
        export_settings = {
            'last_export_path': dest_parent,
            'last_export_format': selected_format,
            'last_archive_name': archive_name,
            'last_export_date': datetime.now().strftime('%Y-%m-%d'),
            'current_folder': source_folder,
            'auto_export': export_settings.get('auto_export', False) if 'export_settings' in locals() else False
        }

        config['export_settings'] = export_settings
        save_config(config)
        update_message("[設定] 已儲存匯出設定")

        # 詢問是否開啟
        if files_created:
            result = show_large_yesno(
                "匯出完成",
                f"已成功匯出封存：\n\n{archive_name}\n\n位置：{dest_parent}\n\n是否要開啟資料夾？"
            )

            if result:
                try:
                    os.startfile(dest_parent)  # 開啟父資料夾
                except Exception as e:
                    update_message(f"[錯誤] 開啟資料夾失敗：{e}")

    except Exception as e:
        update_message(f"[錯誤] 匯出封存失敗：{e}")
        import traceback
        update_message(traceback.format_exc())
        messagebox.showerror("錯誤", f"匯出封存失敗：\n\n{e}")

def show_smart_update_dialog(archive_name, export_path, size_mb, mod_date,
                             area, section, lot_number, 門牌, 今天日期):
    """顯示智慧更新對話框

    Returns:
        (update_mode, auto_export): 更新模式和是否記住選擇
        update_mode: "keep_date", "update_date", "create_new", "reset", None
    """
    # 取得主視窗
    try:
        parent_window = root
    except:
        parent_window = None

    dialog = tk.Toplevel(parent_window)
    dialog.title("智慧匯出 - 發現現有封存")

    # 🔥 增加高度確保按鈕可見
    dialog_width = 620
    dialog_height = 700  # 從 600 增加到 700

    # 置中顯示
    screen_width = dialog.winfo_screenwidth()
    screen_height = dialog.winfo_screenheight()
    x = (screen_width - dialog_width) // 2
    y = (screen_height - dialog_height) // 2
    dialog.geometry(f"{dialog_width}x{dialog_height}+{x}+{y}")

    # 🔥 設定最小尺寸，防止被壓縮
    dialog.minsize(dialog_width, dialog_height)
    # 設定最大尺寸，防止被拉太大
    dialog.maxsize(dialog_width, dialog_height)

    if parent_window:
        dialog.transient(parent_window)
    dialog.grab_set()

    # 🔥 確保對話框在最上層且無法調整大小
    dialog.resizable(False, False)
    dialog.lift()
    dialog.focus_force()

    # 標題（減少間距）
    tk.Label(dialog, text="📦 發現現有封存",
            font=("Microsoft JhengHei", 16, "bold"),
            fg='#FF9800').pack(pady=10)

    # 現有封存資訊框（減少間距）
    info_frame = tk.Frame(dialog, relief=tk.RIDGE, borderwidth=2, bg='#F5F5F5')
    info_frame.pack(padx=20, pady=8, fill=tk.BOTH)

    tk.Label(info_frame, text=f"位置：{export_path}",
            font=("Microsoft JhengHei", 12), bg='#F5F5F5',
            anchor='w').pack(padx=15, pady=3, fill=tk.X)

    tk.Label(info_frame, text=f"檔案：{archive_name}",
            font=("Microsoft JhengHei", 12, "bold"), bg='#F5F5F5',
            anchor='w').pack(padx=15, pady=3, fill=tk.X)

    tk.Label(info_frame, text=f"大小：{size_mb:.2f} MB",
            font=("Microsoft JhengHei", 12), bg='#F5F5F5',
            anchor='w').pack(padx=15, pady=3, fill=tk.X)

    tk.Label(info_frame, text=f"修改：{mod_date}",
            font=("Microsoft JhengHei", 12), bg='#F5F5F5',
            anchor='w').pack(padx=15, pady=3, fill=tk.X)

    # 分隔線（減少間距）
    tk.Frame(dialog, height=2, bg='#E0E0E0').pack(fill=tk.X, padx=20, pady=6)

    # 選擇更新方式（減少間距）
    tk.Label(dialog, text="請選擇更新方式：",
            font=("Microsoft JhengHei", 13, "bold")).pack(pady=8)

    update_mode = tk.StringVar(value="keep_date")

    # 提取原日期
    原日期 = archive_name.split(area)[0] if area in archive_name else "原日期"

    # 選項（減少間距）
    tk.Radiobutton(dialog,
                  text=f"● 更新檔案內容，保持原日期（{原日期}）",
                  variable=update_mode, value="keep_date",
                  font=("Microsoft JhengHei", 12),
                  fg='#2196F3', selectcolor='white').pack(pady=5, anchor="w", padx=50)

    tk.Radiobutton(dialog,
                  text=f"○ 更新檔案內容，並更新日期為今天（{今天日期}）",
                  variable=update_mode, value="update_date",
                  font=("Microsoft JhengHei", 12),
                  selectcolor='white').pack(pady=5, anchor="w", padx=50)

    tk.Radiobutton(dialog,
                  text="○ 建立新的封存（保留舊的，加上時間戳記）",
                  variable=update_mode, value="create_new",
                  font=("Microsoft JhengHei", 12),
                  selectcolor='white').pack(pady=5, anchor="w", padx=50)

    tk.Radiobutton(dialog,
                  text="○ 重新選擇位置和格式",
                  variable=update_mode, value="reset",
                  font=("Microsoft JhengHei", 12),
                  selectcolor='white').pack(pady=5, anchor="w", padx=50)

    # 分隔線（減少間距）
    tk.Frame(dialog, height=1, bg='#E0E0E0').pack(fill=tk.X, padx=50, pady=8)

    # 記住我的選擇
    auto_export_var = tk.BooleanVar(value=False)
    checkbox = tk.Checkbutton(dialog,
                              text="☑ 記住我的選擇（下次直接更新，不再詢問）",
                              variable=auto_export_var,
                              font=("Microsoft JhengHei", 11),
                              fg='#666666')
    checkbox.pack(pady=10)

    # 🔥 彈性空間，把按鈕推到底部
    spacer = tk.Frame(dialog, height=20)
    spacer.pack(fill=tk.BOTH, expand=True)

    # 按鈕區域（固定在底部）
    btn_frame = tk.Frame(dialog, bg='#F0F0F0', relief=tk.RAISED, borderwidth=1)
    btn_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=0)

    # 內部按鈕容器（置中）
    btn_container = tk.Frame(btn_frame, bg='#F0F0F0')
    btn_container.pack(pady=15)

    result = [None]

    def on_confirm():
        result[0] = (update_mode.get(), auto_export_var.get())
        dialog.destroy()

    def on_cancel():
        dialog.destroy()

    # 🔥 更大、更明顯的按鈕
    tk.Button(btn_container, text="✓ 開始匯出", command=on_confirm,
             font=("Microsoft JhengHei", 14, "bold"),
             bg='#4CAF50', fg='white',
             width=14, height=2,
             cursor='hand2').pack(side=tk.LEFT, padx=15)

    tk.Button(btn_container, text="✕ 取消", command=on_cancel,
             font=("Microsoft JhengHei", 14, "bold"),
             bg='#F44336', fg='white',
             width=14, height=2,
             cursor='hand2').pack(side=tk.LEFT, padx=15)

    dialog.wait_window()
    return result[0]

def set_cell_border(cell):
    """設定表格儲存格邊框"""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '4')
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), '000000')
            tcBorders.append(edge_element)
        
        tcPr.append(tcBorders)
    except Exception as e:
        print(f"設定邊框失敗：{e}")

def set_cell_text_color(cell, text, color_rgb=(0, 112, 192)):
    """設定儲存格文字並套用顏色
    
    Args:
        cell: 儲存格物件
        text: 要填入的文字
        color_rgb: RGB 顏色 tuple，預設為亮藍色 (0, 112, 192)
    """
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(*color_rgb)


def generate_map_image(lat, lng, zoom=16, layer='EMAP'):
    """使用地圖服務生成地圖（支援多種圖層）"""
    try:
        start_time = time.time()
        print(f"開始生成地圖：{layer}, zoom={zoom}")
        import requests
        from io import BytesIO
        from PIL import Image, ImageDraw
        import math
        
        # 確保參數是正確型態
        lat = float(lat)
        lng = float(lng)
        zoom = int(zoom)
        
        # 從下拉選單值提取圖層代碼
        if ' - ' in layer:
            layer = layer.split(' - ')[0].strip()
        
        print(f"生成地圖...")
        print(f"圖層：{layer}")
        print(f"座標：{lat}, {lng} (縮放={zoom})")
        
        # 經緯度轉換為圖磚座標
        def lat_lng_to_tile(lat, lng, zoom):
            """將經緯度轉換為圖磚座標"""
            lat_rad = math.radians(lat)
            n = 2.0 ** zoom
            x_tile = int((lng + 180.0) / 360.0 * n)
            y_tile = int((1.0 - math.asinh(math.tan(lat_rad)) / math.pi) / 2.0 * n)
            return x_tile, y_tile
        
        # 計算中心點的圖磚座標
        center_x, center_y = lat_lng_to_tile(lat, lng, zoom)
        
        # 下載 3x3 的圖磚
        tile_size = 256
        tiles_per_side = 3
        
        # 創建空白畫布
        map_width = tile_size * tiles_per_side
        map_height = tile_size * tiles_per_side
        final_image = Image.new('RGB', (map_width, map_height), color='white')
        
        # 🔥 根據圖層選擇不同的圖磚服務

        # Google 圖層
        if layer == 'GOOGLE_SATELLITE':
            base_url = "https://mt0.google.com/vt/lyrs=s"
            url_format = "google"
        elif layer == 'GOOGLE_HYBRID':
            base_url = "https://mt0.google.com/vt/lyrs=y"
            url_format = "google"
        elif layer == 'GOOGLE_STREET':
            base_url = "https://mt0.google.com/vt/lyrs=m"
            url_format = "google"

        # OpenStreetMap 圖層
        elif layer == 'OSM':
            base_url = "https://tile.openstreetmap.org"
            url_format = "tile"
        elif layer == 'OSM_HOT':
            # Humanitarian OpenStreetMap（正確的網址）
            base_url = "https://a.tile.openstreetmap.fr/hot"
            url_format = "tile"  # 標記為 OSM 格式
        elif layer == 'OSM_DE':
            # OpenStreetMap 德國伺服器（備用，速度較快）
            base_url = "https://a.tile.openstreetmap.de"
            url_format = "tile"
        elif layer == 'PHOTO_MIX':
            # 正射影像地籍套疊
            base_url = "https://wmts.nlsc.gov.tw/wmts/PHOTO_MIX/default/GoogleMapsCompatible"
            url_format = "wmts"  # 標記為 WMTS 格式
        elif layer == 'PHOTO2':
            # 純正射影像
            base_url = "https://wmts.nlsc.gov.tw/wmts/PHOTO2/default/GoogleMapsCompatible"
            url_format = "wmts"
        else:
            # 內政部其他圖層
            base_url = f"https://wmts.nlsc.gov.tw/wmts/{layer}/default/GoogleMapsCompatible"
            url_format = "wmts"

        print(f"開始下載地圖圖磚...")
        print(f"服務：{base_url}")
        downloaded = 0

        # 下載並拼接圖磚
        for dy in range(-1, 2):
            for dx in range(-1, 2):
                tile_x = center_x + dx
                tile_y = center_y + dy
                
                # 🔥 根據不同的格式構建 URL
                if url_format == "tile":
                    # OSM 格式：/{z}/{x}/{y}.png
                    tile_url = f"{base_url}/{zoom}/{tile_x}/{tile_y}.png"
                elif url_format == "google":
                    # 🔥 Google 格式：&x={x}&y={y}&z={z}
                    tile_url = f"{base_url}&x={tile_x}&y={tile_y}&z={zoom}"
                else:  # wmts
                    # 內政部 WMTS 格式：/{z}/{y}/{x}
                    tile_url = f"{base_url}/{zoom}/{tile_y}/{tile_x}"
                
                try:
                    # 🔥 設定 User-Agent（OSM 要求）
                    headers = {
                        'User-Agent': 'PropertySurveyReport/1.0'
                    }
                    
                    response = requests.get(tile_url, headers=headers, timeout=10)
                    
                    if response.status_code == 200:
                        tile_img = Image.open(BytesIO(response.content))
                        
                        paste_x = (dx + 1) * tile_size
                        paste_y = (dy + 1) * tile_size
                        
                        final_image.paste(tile_img, (paste_x, paste_y))
                        downloaded += 1
                        print(f"  [OK] 圖磚 {downloaded}/9")
                    else:
                        print(f"  ✗ 圖磚 HTTP {response.status_code}")
                        if downloaded == 0:  # 第一張就失敗，印出 URL 除錯
                            print(f"     URL: {tile_url}")
                except Exception as e:
                    print(f"  ✗ 圖磚錯誤：{e}")
                    continue
        
        print(f"下載完成：{downloaded}/9 張圖磚")
        
        if downloaded >= 6:
            # 🔥 修正：計算錨點在拼接大圖中的精確像素位置
            
            import math
            
            def lat_to_tile_y(lat, zoom):
                """經緯度轉圖磚座標"""
                lat_rad = math.radians(lat)
                n = 2.0 ** zoom
                return (1.0 - math.asinh(math.tan(lat_rad)) / math.pi) / 2.0 * n
            
            def lon_to_tile_x(lon, zoom):
                """經度轉圖磚座標"""
                n = 2.0 ** zoom
                return (lon + 180.0) / 360.0 * n
            
            # 1. 計算錨點在圖磚座標系統中的精確位置
            tile_x_float = lon_to_tile_x(lng, zoom)
            tile_y_float = lat_to_tile_y(lat, zoom)
            
            # 2. 計算錨點在拼接大圖中的像素位置
            tile_x_offset = (tile_x_float - center_x) * tile_size
            tile_y_offset = (tile_y_float - center_y) * tile_size
            
            # 錨點在大圖中的位置（從左上角算起）
            marker_x_in_big_image = tile_size + tile_x_offset
            marker_y_in_big_image = tile_size + tile_y_offset
            
            print(f"錨點在大圖中的位置：x={marker_x_in_big_image:.2f}, y={marker_y_in_big_image:.2f}")
            
            # 3. 繪製錨點
            from PIL import ImageDraw
            draw = ImageDraw.Draw(final_image)
            
            # 🔥 優先使用 Google Maps 標記圖片
            marker_added = False
            try:
                print("下載 Google Maps 標記圖片...")
                marker_url = "https://maps.gstatic.com/mapfiles/api-3/images/spotlight-poi2.png"
                
                marker_response = requests.get(marker_url, timeout=5)
                if marker_response.status_code == 200:
                    from io import BytesIO
                    marker_img = Image.open(BytesIO(marker_response.content)).convert('RGBA')
                    
                    # 調整標記大小
                    marker_size = (40, 58)
                    marker_img = marker_img.resize(marker_size, Image.Resampling.LANCZOS)
                    
                    # 計算貼上位置（針尖在錨點位置）
                    paste_x = int(marker_x_in_big_image - marker_size[0] // 2)
                    paste_y = int(marker_y_in_big_image - marker_size[1])
                    
                    # 轉換為 RGBA
                    final_image = final_image.convert('RGBA')
                    
                    # 貼上標記
                    final_image.paste(marker_img, (paste_x, paste_y), marker_img)
                    
                    # 轉回 RGB
                    final_image = final_image.convert('RGB')
                    
                    marker_added = True
                    print("[OK] 標記圖片已載入")
            except Exception as e:
                print(f"[警告] 無法載入標記圖片：{e}")
            
            # 🔥 備用：繪製標記
            if not marker_added:
                print("使用繪製方式建立標記...")
                
                def draw_google_marker(draw, x, y, size=40):
                    """繪製 Google Maps 風格標記"""
                    width = size * 0.7
                    radius = width / 2
                    center_x = x
                    center_y = y - size * 0.65
                    
                    points = []
                    
                    # 上半部：圓弧
                    for angle in range(180, -1, -10):
                        rad = math.radians(angle)
                        px = center_x + radius * math.cos(rad)
                        py = center_y + radius * math.sin(rad)
                        points.append((px, py))
                    
                    # 右側曲線到尖端
                    for t in range(0, 11):
                        t = t / 10.0
                        px = (1-t)**2 * (center_x + radius) + 2*(1-t)*t * (center_x + radius*0.3) + t**2 * x
                        py = (1-t)**2 * center_y + 2*(1-t)*t * (y - size*0.2) + t**2 * y
                        points.append((px, py))
                    
                    # 左側曲線回到圓弧
                    for t in range(0, 11):
                        t = t / 10.0
                        px = (1-t)**2 * x + 2*(1-t)*t * (center_x - radius*0.3) + t**2 * (center_x - radius)
                        py = (1-t)**2 * y + 2*(1-t)*t * (y - size*0.2) + t**2 * center_y
                        points.append((px, py))
                    
                    # 繪製陰影
                    shadow_offset = 3
                    shadow_points = [(px, py + shadow_offset) for px, py in points]
                    draw.polygon(shadow_points, fill=(0, 0, 0, 60))
                    
                    # 繪製外框
                    draw.polygon(points, fill='#C5221F')
                    
                    # 繪製內部
                    highlight_points = []
                    scale = 0.92
                    for px, py in points:
                        new_x = center_x + (px - center_x) * scale
                        new_y = center_y + (py - center_y) * scale
                        highlight_points.append((new_x, new_y))
                    draw.polygon(highlight_points, fill='#EA4335')
                    
                    # 繪製中心白圓
                    inner_radius = radius * 0.35
                    inner_box = [
                        center_x - inner_radius,
                        center_y - inner_radius,
                        center_x + inner_radius,
                        center_y + inner_radius
                    ]
                    draw.ellipse(inner_box, fill='white')
                
                draw_google_marker(draw, int(marker_x_in_big_image), int(marker_y_in_big_image), size=40)
                print("[OK] 標記已繪製")
            
            # 4. 裁切成 600x400（以錨點為中心）
            crop_width = 600
            crop_height = 400
            
            # 計算裁切範圍（以錨點為中心）
            left = int(marker_x_in_big_image - crop_width // 2)
            top = int(marker_y_in_big_image - crop_height // 2)
            
            # 確保不超出邊界
            left = max(0, min(left, map_width - crop_width))
            top = max(0, min(top, map_height - crop_height))
            
            final_image = final_image.crop((left, top, left + crop_width, top + crop_height))
            
            print(f"裁切範圍：left={left}, top={top}, right={left+crop_width}, bottom={top+crop_height}")
            
            # 🔥 如果是 OSM 圖層，在圖片右下角加上文字
            if layer.startswith('OSM'):
                from PIL import ImageFont
                
                draw = ImageDraw.Draw(final_image)
                
                credit_text = "© OpenStreetMap contributors"
                
                try:
                    font = ImageFont.truetype("arial.ttf", 10)
                except:
                    try:
                        font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 10)
                    except:
                        font = ImageFont.load_default()
                
                bbox = draw.textbbox((0, 0), credit_text, font=font)
                text_width = bbox[2] - bbox[0]
                text_height = bbox[3] - bbox[1]
                
                text_x = crop_width - text_width - 5
                text_y = crop_height - text_height - 5
                
                padding = 2
                background_box = [
                    text_x - padding,
                    text_y - padding,
                    text_x + text_width + padding,
                    text_y + text_height + padding
                ]
                
                overlay = Image.new('RGBA', final_image.size, (255, 255, 255, 0))
                overlay_draw = ImageDraw.Draw(overlay)
                overlay_draw.rectangle(background_box, fill=(255, 255, 255, 180))
                
                final_image = final_image.convert('RGBA')
                final_image = Image.alpha_composite(final_image, overlay)
                final_image = final_image.convert('RGB')
                
                draw = ImageDraw.Draw(final_image)
                draw.text((text_x, text_y), credit_text, fill=(80, 80, 80), font=font)
            
            # 儲存圖片
            map_path = "temp_map.png"
            final_image.save(map_path, 'PNG')
            
            print(f"[完成] 地圖已生成：{map_path}", flush=True)
            print(f"地圖生成耗時：{time.time() - start_time:.2f} 秒")
            return map_path
        else:
            print(f"[錯誤] 圖磚下載不足：{downloaded}/9")
            return None
            
    except Exception as e:
        print(f"[錯誤] 生成地圖失敗：{e}")
        import traceback
        traceback.print_exc()
        return None


def add_watermark_to_word(doc, watermark_text):
    """
    添加浮水印到 Word 文檔（使用圖片方式，放在最底層）

    參數：
        doc: Document 對象
        watermark_text: 浮水印文字（例如："C 1 4 5"）
    """
    # 🔥🔥🔥 浮水印功能暫時停用 🔥🔥🔥
    # 原因：
    # 1. 使用 VML/XML 方式會破壞文件結構
    # 2. 使用圖片插入到頁首會把表格擠下去
    #
    # 正確做法需要：
    # - 使用 WordArt 背景層
    # - 或使用第三方工具（如 python-docx-template）
    # - 或手動編輯 document.xml 添加背景水印
    #
    # 暫時停用，等找到穩定的解決方案
    print(f"[提示] 浮水印功能暫時停用")
    return False


def get_employee_code_from_contacts(employee_name, contacts_data):
    """
    從通訊錄數據中獲取員工編號

    參數：
        employee_name: 員工姓名
        contacts_data: 通訊錄數據（字典格式）

    返回：
        員工編號字符串，如果找不到則返回空字符串
    """
    try:
        contacts_by_store = contacts_data.get('contacts_by_store', {})

        # 遍歷所有店別查找員工
        for store_contacts in contacts_by_store.values():
            for contact in store_contacts:
                if contact.get('name') == employee_name:
                    code = contact.get('code', '')
                    if code:
                        return code

        return ''
    except Exception as e:
        print(f"警告：獲取員工編號失敗：{e}")
        return ''


def format_watermark_text(code):
    """
    格式化浮水印文字（在字母和數字之間加空格）

    參數：
        code: 員工編號（例如："C145"）

    返回：
        格式化後的文字（例如："C 1 4 5"）
    """
    if not code:
        return ''

    # 將編號拆分為字母和數字，並在每個字符之間加空格
    result = []
    for char in code:
        result.append(char)

    return ' '.join(result)


def generate_word_report_buy(output_path, data, report_date):
    """生成 Word 格式報告（售件）- A4 直向，固定表格結構"""


    # ========================================
    # 🔥 開始生成報告
    # ========================================
    try:
        from docx import Document
        from docx.shared import Cm, Pt, Inches, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ... (原有的報告生成程式碼)

        # === 設定頁面為 A4 直向 ===
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.9)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        
        # 預設字體
        style = doc.styles['Normal']
        style.font.name = 'PMingLiU'  # 新細明體
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
        style.font.size = Pt(11)  # 🔥 加大字體
        style.font.bold = True  # 🔥 加粗
        
        # === 標題區域（使用表格佈局）===
        import os
        from docx.shared import Cm, Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL

        # 建立一個 1 列 3 欄的表格（這會是整體的一部分）
        header_table = doc.add_table(rows=1, cols=3)
        header_table.autofit = False
        header_table.allow_autofit = False

        # 設定欄寬
        header_table.columns[0].width = Cm(5)   # LOGO 欄（加寬）
        header_table.columns[1].width = Cm(10)  # 標題欄
        header_table.columns[2].width = Cm(3)   # 日期欄

        # 🔥 隱藏表格邊框
        def remove_table_borders(table):
            """移除表格所有邊框"""
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tblBorders.append(border)
            tblPr.append(tblBorders)

        remove_table_borders(header_table)

        # 設定列高
        header_table.rows[0].height = Cm(1.2)  # 加高以容納 LOGO

        # 🔥 左側：LOGO（放大）
        logo_cell = header_table.rows[0].cells[0]
        logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM  # 靠下對齊
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        logo_path = "logo-1.png"
        if os.path.exists(logo_path):
            run_logo = logo_para.add_run()
            run_logo.add_picture(logo_path, height=Cm(1.0))  # 🔥 放大 LOGO

        # 🔥 中間：標題 - 大字體、置中、靠下
        title_cell = header_table.rows[0].cells[1]
        title_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM  # 靠下對齊
        title_para = title_cell.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_title = title_para.add_run('物件個案調查表(售)')
        run_title.font.name = 'PMingLiU'
        run_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
        run_title.font.size = Pt(18)  # 🔥 字體加大
        run_title.font.bold = True

        # 🔥 右側：日期 - 小字體、靠右、靠下
        date_cell = header_table.rows[0].cells[2]
        date_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM  # 靠下對齊
        date_para = date_cell.paragraphs[0]
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run_date = date_para.add_run(f'製表日期：{report_date}')
        run_date.font.name = 'PMingLiU'
        run_date._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
        run_date.font.size = Pt(11)  # 🔥 加大字體
        run_date.font.bold = True  # 🔥 加粗

        # 🔥 不要加空白行，直接接主表格

        
        # ========================================
        # 建立主表格：20 列 × 28 欄（加倍以便更精細控制）
        # ========================================
        table = doc.add_table(rows=20, cols=28)
        table.style = 'Table Grid'

        # 🔥 設定表格填滿整頁
        # 計算可用寬度
        available_width = section.page_width - section.left_margin - section.right_margin

        # 設定表格寬度
        table.width = available_width

        # 🔥 精確設定每列高度
        列高設定 = {
            0: Cm(1.0),   # 物件編號
            1: Cm(1.0),   # 土地標示
            2: Cm(1.0),   # 建物門牌
            3: Cm(0.9),   # 型式/用途
            4: Cm(1.0),   # 總建坪/總地坪
            5: Cm(0.95),   # 主建物第1列
            6: Cm(0.95),   # 主建物第2列
            7: Cm(0.95),   # 主建物第3列
            8: Cm(0.95),   # 主建物第4列
            9: Cm(0.95),   # 主建物第5列
            10: Cm(0.95),  # 主建物第6列
            11: Cm(0.95),  # 主建物第7列
            12: Cm(0.95),  # 主建物第8列
            13: Cm(0.95),  # 主建物第9列
            14: Cm(0.95),  # 附屬建物
            15: Cm(0.95),  # 公共設施
            16: Cm(0.95),  # 增建面積等
            17: Cm(0.95),  # 訴求重點（標題+內容，合併）
            18: Cm(2.0),  # 訴求重點（與17合併）
            19: Cm(1.0),  # 承辦人
        }

        for row_idx, height in 列高設定.items():
            table.rows[row_idx].height = height

        # 設定欄寬（平均分配給28欄）
        for row in table.rows:
            for cell in row.cells:
                cell.width = available_width / 28

        
        # ========================================
        # 第 0 列：物件編號、契約編號、案名（28欄版本）
        # ========================================
        # 原本14欄：物件編號標題(1欄) + 內容(3欄) + 契約編號標題(1欄) + 內容(3欄) + 案名標題(1欄) + 內容(5欄)
        # 現在28欄：物件編號標題(2欄) + 內容(6欄) + 契約編號標題(2欄) + 內容(6欄) + 案名標題(2欄) + 內容(10欄)

        # 物件編號標題（2欄，折行、置中）
        table.rows[0].cells[0].merge(table.rows[0].cells[1]).text = '物件\n編號'
        table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 物件編號內容（6欄，cells[2] 到 cells[7]）
        table.rows[0].cells[2].merge(table.rows[0].cells[7]).text = data.get('物件編號', '')

        # 契約編號標題（2欄，折行、置中）
        table.rows[0].cells[8].merge(table.rows[0].cells[9]).text = '契約\n編號'
        table.rows[0].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 契約編號內容（6欄，cells[10] 到 cells[15]）
        table.rows[0].cells[10].merge(table.rows[0].cells[15]).text = data.get('契約編號', '')

        # 案名標題（2欄，置中）
        table.rows[0].cells[16].merge(table.rows[0].cells[17]).text = '案名'
        table.rows[0].cells[16].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 案名內容（10欄，置中，cells[18] 到 cells[27]）
        table.rows[0].cells[18].merge(table.rows[0].cells[27]).text = data.get('案名', '')
        table.rows[0].cells[18].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ========================================
        # 第 1 列：土地標示（28欄版本）
        # ========================================
        # 原本14欄：土地標示標題(1欄) + 內容(13欄)
        # 現在28欄：土地標示標題(2欄) + 內容(26欄)

        # 土地標示標題（2欄，折行、置中）
        cell = table.rows[1].cells[0].merge(table.rows[1].cells[1])
        cell.text = '土地\n標示'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 土地標示內容（26欄，cells[2] 到 cells[27]）
        cell = table.rows[1].cells[2].merge(table.rows[1].cells[27])
        cell.text = data.get('土地標示', '')
        
        # ========================================
        # 第 2 列：建物門牌、總價（28欄版本）
        # ========================================
        # 原本14欄：建物門牌標題(1欄) + 內容(6欄) + 總價(7欄)
        # 現在28欄：建物門牌標題(2欄) + 內容(12欄) + 總價(14欄)

        # 建物門牌標題（2欄，折行、置中）
        cell = table.rows[2].cells[0].merge(table.rows[2].cells[1])
        cell.text = '建物\n門牌'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 建物門牌內容（12欄，cells[2] 到 cells[13]）
        cell = table.rows[2].cells[2].merge(table.rows[2].cells[13])
        cell.text = data.get('建物門牌', '')
        
        總價 = data.get('總價', '')
        含車位 = data.get('含車位價格', '')

        # 🔥 格式化總價為中文單位
        def format_price_for_report(price_wan):
            """將萬元轉換為中文單位顯示（保留小數點後一位）"""
            try:
                price_float = float(price_wan)
                if price_float == 0:
                    return ""

                # 判斷是否有小數
                has_decimal = (price_float != int(price_float))

                units = [
                    ("載", 10**44), ("正", 10**40), ("澗", 10**36), ("溝", 10**32),
                    ("穰", 10**28), ("秭", 10**24), ("垓", 10**20), ("京", 10**16),
                    ("兆", 10**8), ("億", 10**4), ("萬", 1)
                ]

                result = []
                remainder = price_float

                for i, (unit_name, unit_value) in enumerate(units):
                    if remainder >= unit_value:
                        count = remainder / unit_value

                        # 最後一個單位保留小數（如果有的話）
                        if i == len(units) - 1 or remainder < units[i+1][1] if i+1 < len(units) else True:
                            if has_decimal and count < 10000:  # 小於下一個單位才顯示小數
                                result.append(f"{count:.1f}{unit_name}")
                                break
                            else:
                                count_int = int(count)
                                result.append(f"{count_int}{unit_name}")
                                break
                        else:
                            count_int = int(count)
                            remainder = remainder - (count_int * unit_value)
                            if count_int > 0:
                                result.append(f"{count_int}{unit_name}")

                return "".join(result) if result else ""
            except (ValueError, TypeError):
                return price_wan  # 如果轉換失敗，返回原值

        if 總價:
            總價_中文 = format_price_for_report(總價)
            含車位_中文 = format_price_for_report(含車位) if 含車位 else ""

            if 含車位_中文:
                價格文字 = f"總價：{總價_中文}\n(含車位 {含車位_中文})"
            else:
                價格文字 = f"總價：{總價_中文}\n(含車位      )"
        else:
            價格文字 = "總價：\n(含車位      )"

        # 總價內容（14欄，cells[14] 到 cells[27]）
        table.rows[2].cells[14].merge(table.rows[2].cells[27]).text = 價格文字
        
        # ========================================
        # 第 3 列：型式、用途（28欄版本）
        # ========================================
        # 原本14欄：型式(7欄) + 用途(7欄)
        # 現在28欄：型式(14欄) + 用途(14欄)

        型式 = data.get('型式', '')
        型式dict = {'土地': '□', '透天': '□', '大樓': '□', '華廈': '□', '公寓': '□'}
        if 型式 in 型式dict:
            型式dict[型式] = '■'
        型式選項 = ' '.join([f"{v}{k}" for k, v in 型式dict.items()])
        # 型式（14欄，cells[0] 到 cells[13]）
        table.rows[3].cells[0].merge(table.rows[3].cells[13]).text = f'型式：{型式選項}'

        用途 = data.get('用途', '')
        用途dict = {'建地': '□', '工業地': '□', '廠房': '□', '店面': '□',
                   '辦公': '□', '住家': '□', '農地': '□', '其他': '□'}
        if 用途 in 用途dict:
            用途dict[用途] = '■'
        用途選項 = ' '.join([f"{v}{k}" for k, v in 用途dict.items()])
        # 用途（14欄，cells[14] 到 cells[27]）
        table.rows[3].cells[14].merge(table.rows[3].cells[27]).text = f'用途：{用途選項}'
        
        # ========================================
        # 第 4 列：總建坪、總地坪、竣工日期、建築結構（28欄版本）
        # ========================================
        # 原本14欄：總建坪(4欄) + 總地坪(3欄) + 竣工日期(4欄) + 建築結構(3欄)
        # 現在28欄：總建坪(7欄) + 總地坪(7欄) + 竣工日期(7欄) + 建築結構(7欄) - 平均分配

        總建坪 = data.get('總建坪', '')
        if 總建坪:
            try:
                坪數 = float(總建坪)
                平方公尺 = 坪數 / 0.3025
                總建坪文字 = f"總建坪：{總建坪}坪\n({平方公尺:.2f}平方公尺)"
            except:
                總建坪文字 = f"總建坪：{總建坪}坪"
        else:
            總建坪文字 = "總建坪：   坪\n(    平方公尺)"
        # 總建坪（7欄，cells[0] 到 cells[6]）
        table.rows[4].cells[0].merge(table.rows[4].cells[6]).text = 總建坪文字

        總地坪 = data.get('總地坪', '')
        if 總地坪:
            try:
                坪數 = float(總地坪)
                平方公尺 = 坪數 / 0.3025
                總地坪文字 = f"總地坪：{總地坪}坪\n({平方公尺:.2f}平方公尺)"
            except:
                總地坪文字 = f"總地坪：{總地坪}坪"
        else:
            總地坪文字 = "總地坪：   坪\n(    平方公尺)"
        # 總地坪（7欄，cells[7] 到 cells[13]）
        table.rows[4].cells[7].merge(table.rows[4].cells[13]).text = 總地坪文字

        # 竣工日期（7欄，cells[14] 到 cells[20]）
        table.rows[4].cells[14].merge(table.rows[4].cells[20]).text = f"竣工日期：{data.get('竣工日期', '')}"
        # 建築結構（7欄，cells[21] 到 cells[27]）
        table.rows[4].cells[21].merge(table.rows[4].cells[27]).text = f"建築結構：{data.get('建築結構', '')}"

        
        # ========================================
        # 第 5-13 列：主建物區域（固定 9 列）（28欄版本）
        # ========================================
        # 結構：主建物標題(1欄) + 層名坪數(6欄) + 右側三區(7+7+7欄)

        table.rows[5].cells[0].merge(table.rows[13].cells[0]).text = '主\n建\n物'

        主建物 = data.get('主建物', {})
        主建物列表 = list(主建物.items())

        for i in range(9):
            row_idx = 5 + i

            # 🔥 左側：層名和坪數（6欄，cells[1-6]）
            if i < len(主建物列表):
                層名, 坪數 = 主建物列表[i]
                merged_cell = table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[6])
                merged_cell.text = f"{層名}：{坪數} 坪"
                merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # 沒有資料時，完全空白
                merged_cell = table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[6])
                merged_cell.text = ""

            # 右側：根據列數填入不同資訊
            if i == 0:
                # 第 6 列（一層）
                土地單價 = data.get('土地單價', '')
                土地單價_顯示 = data.get('土地單價_顯示', True)
                # 土地單價（7欄，cells[7-13]）- 根據顯示開關決定是否顯示數值
                if 土地單價_顯示:
                    table.rows[row_idx].cells[7].merge(table.rows[row_idx].cells[13]).text = f"土地單價：{土地單價}萬元/坪" if 土地單價 else "土地單價："
                else:
                    table.rows[row_idx].cells[7].merge(table.rows[row_idx].cells[13]).text = "土地單價："
                # 座向（7欄，cells[14-20]）
                table.rows[row_idx].cells[14].merge(table.rows[row_idx].cells[20]).text = f"座向：{data.get('座向', '')}"

                格局 = data.get('格局', {})
                房 = 格局.get('房', '')
                廳 = 格局.get('廳', '')
                衛 = 格局.get('衛', '')
                # 格局選項處理
                格局選項_data = data.get('格局選項', [])
                格局選項 = ["開放式", "空地"]
                格局選項文字 = "  ".join([f"{'■' if opt in 格局選項_data else '□'}{opt}" for opt in 格局選項])
                格局文字 = f"格局：{格局選項文字}\n{房}房 {廳}廳 {衛}衛"
                # 格局（7欄，合併第6-7列的 cells[21-27]）
                table.rows[5].cells[21].merge(table.rows[6].cells[27])
                table.rows[5].cells[21].text = 格局文字

            elif i == 1:
                # 第 7 列（二層）
                建物單價 = data.get('建物單價', '')
                建物單價_顯示 = data.get('建物單價_顯示', True)
                # 建物單價（7欄，cells[7-13]）- 根據顯示開關決定是否顯示數值
                if 建物單價_顯示:
                    table.rows[row_idx].cells[7].merge(table.rows[row_idx].cells[13]).text = f"建物單價：{建物單價}萬元/坪" if 建物單價 else "建物單價："
                else:
                    table.rows[row_idx].cells[7].merge(table.rows[row_idx].cells[13]).text = "建物單價："

                # 🔥 鑰匙：如果是「未填」就顯示空白
                鑰匙 = data.get('鑰匙', '')
                鑰匙文字 = f"鑰匙：{鑰匙}" if 鑰匙 and 鑰匙 != "未填" else "鑰匙："
                # 鑰匙（7欄，cells[14-20]）
                table.rows[row_idx].cells[14].merge(table.rows[row_idx].cells[20]).text = 鑰匙文字

            elif i == 2:
                # 第 8 列（三層）
                # 所有權型態（7欄，cells[7-13]）
                table.rows[row_idx].cells[7].merge(table.rows[row_idx].cells[13]).text = f"所有權型態：{data.get('所有權型態', '')}"

                # 🔥 現況：根據類型顯示說明
                現況 = data.get('現況', '')
                現況說明 = data.get('現況說明', '')
                if 現況 == "出租中" and 現況說明:
                    現況文字 = f"現況：{現況}（出租至{現況說明}）"
                elif 現況 == "其他" and 現況說明:
                    現況文字 = f"現況：{現況}（{現況說明}）"
                else:
                    現況文字 = f"現況：{現況}"

                # 現況（7欄，cells[14-20]）
                table.rows[row_idx].cells[14].merge(table.rows[row_idx].cells[20]).text = 現況文字

                車位 = data.get('車位', {})
                車位有無 = 車位.get('有無', '')
                車位型式 = 車位.get('型式', '')
                車位編號 = 車位.get('編號', '')

                # 處理型式的複選框
                型式選項 = {'平面': '□', '機械': '□'}
                if 車位型式:
                    if '平面' in 車位型式:
                        型式選項['平面'] = '■'
                    if '機械' in 車位型式:
                        型式選項['機械'] = '■'
                型式文字 = ''.join([f"{v}{k}" for k, v in 型式選項.items()])

                車位文字 = f"車位：{車位有無}\n型式：{型式文字}\n編號：{車位編號}"
                # 車位（7欄，cells[21-27]）
                table.rows[row_idx].cells[21].merge(table.rows[row_idx].cells[27]).text = 車位文字

            elif i == 3:
                # 第 9 列（四層）
                地上層數 = data.get('地上層數', '')
                地下層數 = data.get('地下層數', '')
                建物文字 = f"建物：\n地上 {地上層數} 地下 {地下層數} 樓"
                # 建物（7欄，cells[7-13]）
                table.rows[row_idx].cells[7].merge(table.rows[row_idx].cells[13]).text = 建物文字

                # 🔥 電梯含數量和一層幾戶
                電梯 = data.get('電梯', '')
                電梯數量 = data.get('電梯數量', '')
                一層幾戶 = data.get('一層幾戶', '')

                if 電梯 == '有':
                    電梯parts = []
                    if 電梯數量:
                        電梯parts.append(f"{電梯數量}部")
                    if 一層幾戶:
                        電梯parts.append(f"一層{一層幾戶}戶")

                    if 電梯parts:
                        電梯文字 = f"電梯：{電梯}（{' '.join(電梯parts)}）"
                    else:
                        電梯文字 = f"電梯：{電梯}"
                else:
                    電梯文字 = f"電梯：{電梯}"

                # 電梯（7欄，cells[14-20]）
                table.rows[row_idx].cells[14].merge(table.rows[row_idx].cells[20]).text = 電梯文字

                # 履約保證
                履約保證 = data.get('履約保證', '')
                if 履約保證 == '同意':
                    履約文字 = "履約保證：\n■同意 □不同意"
                elif 履約保證 == '不同意':
                    履約文字 = "履約保證：\n□同意 ■不同意"
                else:
                    履約文字 = "履約保證：\n□同意 □不同意"
                # 履約保證（7欄，cells[21-27]）
                table.rows[row_idx].cells[21].merge(table.rows[row_idx].cells[27]).text = 履約文字

            elif i == 4:
                # 第 10 列（五層）
                # 公告現值（7欄，cells[7-13]）
                table.rows[row_idx].cells[7].merge(table.rows[row_idx].cells[13]).text = f"公告現值：{data.get('公告現值', '')}元/M²"
                # 管理費（7欄，cells[14-20]）
                table.rows[row_idx].cells[14].merge(table.rows[row_idx].cells[20]).text = f"管理費：{data.get('管理費', '')}元/月"

                # 廠房保存登記（7欄，cells[21-27]）
                廠房保存登記 = data.get('廠房保存登記', '')
                table.rows[row_idx].cells[21].merge(table.rows[row_idx].cells[27]).text = f"廠房保存登記：{廠房保存登記}"

            elif i == 5:
                # 第 11 列（六層）
                # 自用增值稅（7欄，cells[7-13]）
                table.rows[row_idx].cells[7].merge(table.rows[row_idx].cells[13]).text = f"自用增值稅：{data.get('自用增值稅', '')}元"
                面臨道路 = data.get('面臨道路', '')
                # 面臨道路（7欄，cells[14-20]）
                table.rows[row_idx].cells[14].merge(table.rows[row_idx].cells[20]).text = f"面臨道路：約{面臨道路}米" if 面臨道路 else "面臨道路：約  米"

                # 工廠登記（7欄，cells[21-27]）
                工廠登記 = data.get('工廠登記', '')
                table.rows[row_idx].cells[21].merge(table.rows[row_idx].cells[27]).text = f"工廠登記：{工廠登記}"

            elif i == 6:
                # 第 12 列（七層）
                # 一般增值稅（7欄，cells[7-13]）
                table.rows[row_idx].cells[7].merge(table.rows[row_idx].cells[13]).text = f"一般增值稅：{data.get('一般增值稅', '')}元"
                面寬 = data.get('面寬', '')
                # 面寬（7欄，cells[14-20]）
                table.rows[row_idx].cells[14].merge(table.rows[row_idx].cells[20]).text = f"面寬：約{面寬}米" if 面寬 else "面寬：約  米"

                # 分割銷售（7欄，cells[21-27]）
                分割銷售 = data.get('分割銷售', '')
                table.rows[row_idx].cells[21].merge(table.rows[row_idx].cells[27]).text = f"分割銷售：{分割銷售}"

            elif i == 7:
                # 第 13 列（八層）
                # 貸款銀行（7欄，cells[7-13]）
                table.rows[row_idx].cells[7].merge(table.rows[row_idx].cells[13]).text = f"貸款銀行：{data.get('貸款銀行', '')}"
                深度 = data.get('深度', '')
                # 深度（7欄，cells[14-20]）
                table.rows[row_idx].cells[14].merge(table.rows[row_idx].cells[20]).text = f"深度：約{深度}米" if 深度 else "深度：約  米"

                # 天車（7欄，cells[21-27]）
                天車 = data.get('天車', '')
                table.rows[row_idx].cells[21].merge(table.rows[row_idx].cells[27]).text = f"天車：{天車}"

            elif i == 8:
                # 第 14 列（九層）
                設定金額 = data.get('設定金額', '')
                # 🔥 設定金額已包含「萬元正」，只需加上「新台幣」前綴
                # 設定金額（7欄，cells[7-13]）
                table.rows[row_idx].cells[7].merge(table.rows[row_idx].cells[13]).text = f"設定金額：新台幣{設定金額}" if 設定金額 else "設定金額："
                # 建蔽率（7欄，cells[14-20]）
                table.rows[row_idx].cells[14].merge(table.rows[row_idx].cells[20]).text = f"建蔽率：{data.get('建蔽率', '')}%"

                # 樑下（7欄，cells[21-27]）
                樑下 = data.get('樑下', '')
                table.rows[row_idx].cells[21].merge(table.rows[row_idx].cells[27]).text = f"樑下：約{樑下}米" if 樑下 else "樑下：約  米"
        
        # ========================================
        # 第 15 列：附屬建物（28欄版本，7+7+7+7）
        # ========================================
        附屬建物 = data.get('附屬建物', {})
        附屬建物坪數 = data.get('附屬建物坪數', '')

        if isinstance(附屬建物, dict) and 附屬建物:
            # 計算總坪數
            try:
                總坪數 = sum(float(坪數) for 坪數 in 附屬建物.values() if 坪數)
                附屬總計 = f"{總坪數:.2f}坪"
            except:
                附屬總計 = f"{附屬建物坪數}坪" if 附屬建物坪數 else "坪"

            # 各項明細
            附屬明細 = '、'.join([f"{類型}：{坪數}坪" for 類型, 坪數 in 附屬建物.items()])
            附屬文字 = f"附屬建物：{附屬總計}\n(含{附屬明細})"
        else:
            附屬文字 = f"附屬建物：{附屬建物坪數}坪"

        # 附屬建物（7欄，cells[0-6]）
        table.rows[14].cells[0].merge(table.rows[14].cells[6]).text = 附屬文字

        # 鄰近學校（7欄，cells[7-13]）
        table.rows[14].cells[7].merge(table.rows[14].cells[13]).text = f"鄰近學校：{data.get('鄰近學校', '')}"
        # 容積率（7欄，cells[14-20]）
        table.rows[14].cells[14].merge(table.rows[14].cells[20]).text = f"容積率：{data.get('容積率', '')}%"

        # 中脊（7欄，cells[21-27]）
        中脊 = data.get('中脊', '')
        table.rows[14].cells[21].merge(table.rows[14].cells[27]).text = f"中脊：約{中脊}米" if 中脊 else "中脊：約  米"

        # ========================================
        # 第 16 列：公共設施（28欄版本，7+14+7）
        # ========================================

        # 公共設施（7欄，cells[0-6]）
        table.rows[15].cells[0].merge(table.rows[15].cells[6]).text = f"公共設施：{data.get('公共設施坪數', '')}坪"

        # 所在商圈/工業區（14欄，cells[7-20]）
        table.rows[15].cells[7].merge(table.rows[15].cells[20]).text = f"所在商圈/工業區：{data.get('所在商圈', '')}"

        # 滴水（7欄，cells[21-27]）
        滴水 = data.get('滴水', '')
        table.rows[15].cells[21].merge(table.rows[15].cells[27]).text = f"滴水：約{滴水}米" if 滴水 else "滴水：約  米"
        
        # ========================================
        # 第 17 列：增建面積、都市土地使用分區、非都市土地使用地類別（28欄版本，7+9+12）
        # ========================================

        # 增建面積（7欄，cells[0-6]）
        table.rows[16].cells[0].merge(table.rows[16].cells[6]).text = f"增建面積：{data.get('增建面積', '')}坪"

        # 都市土地使用分區（9欄，cells[7-15]）
        table.rows[16].cells[7].merge(table.rows[16].cells[15]).text = f"都市土地使用分區：{data.get('都市土地使用分區', '')}"

        # 非都市土地使用地類別（12欄，cells[16-27]）
        table.rows[16].cells[16].merge(table.rows[16].cells[27]).text = f"非都市土地使用地類別：{data.get('非都市土地使用地類別', '')}"
        
        
        # ========================================
        # 第 18 列：訴求重點（左）、位置圖（右）- 標題列（28欄版本，14+14）
        # ========================================
        # 訴求重點標題（14欄，cells[0-13]，置中）
        訴求重點標題_cell = table.rows[17].cells[0].merge(table.rows[17].cells[13])
        訴求重點標題_cell.text = '訴求重點'
        訴求重點標題_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 位置圖標題（14欄，cells[14-27]，置中）
        位置圖標題_cell = table.rows[17].cells[14].merge(table.rows[17].cells[27])
        位置圖標題_cell.text = '位置圖'
        位置圖標題_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ========================================
        # 第 19 列：訴求重點內容（左）、位置圖內容（右）（28欄版本，14+14）
        # ========================================
        # 合併 18-19 列的內容區域
        # 左側：訴求重點 + 另顯示於後台內容
        訴求重點 = data.get('訴求重點', '')
        後台內容 = data.get('另顯示於後台內容', '')

        # 訴求重點內容（14欄，cells[0-13]，靠左靠上）
        訴求重點內容_cell = table.rows[18].cells[0].merge(table.rows[18].cells[13])
        訴求重點內容_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP  # 垂直靠上

        # 清除預設段落
        訴求重點內容_cell.text = ''

        # 處理訴求重點：每行之間加入段後間距，並設定懸掛縮排
        if 訴求重點.strip():
            lines = 訴求重點.strip().split('\n')
            for i, line in enumerate(lines):
                if i == 0:
                    # 第一行使用預設段落
                    p = 訴求重點內容_cell.paragraphs[0]
                    p.text = line
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    # 後續行：直接加入內容段落
                    p = 訴求重點內容_cell.add_paragraph()
                    p.text = line
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 設定懸掛縮排：第一行不縮排，後續行縮排（適用於「一、」「二、」這種格式）
                p.paragraph_format.left_indent = Cm(0.8)  # 左縮排 0.8cm
                p.paragraph_format.first_line_indent = Cm(-0.8)  # 首行凸排（負值）

                # 為每個段落設定段後間距（相當於半個空行）
                # 但最後一行不加間距
                if i < len(lines) - 1:
                    p.paragraph_format.space_after = Pt(6)  # 6pt 相當於半行的高度

        # 處理另顯示於後台內容
        if 後台內容.strip():
            # 加兩個空行分隔
            訴求重點內容_cell.add_paragraph()
            訴求重點內容_cell.add_paragraph()

            # 加入標題
            title_p = 訴求重點內容_cell.add_paragraph()
            title_p.text = '另顯示於後台內容：'
            title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 加入內容
            content_p = 訴求重點內容_cell.add_paragraph()
            content_p.text = 後台內容
            content_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 🔥 右側：位置圖
        # === 位置圖（第 19 列）===
        # 位置圖內容（14欄，cells[14-27]，置中）
        位置圖_cell = table.rows[18].cells[14].merge(table.rows[18].cells[27])
        # 🔥 設定位置圖儲存格垂直置中
        位置圖_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        緯度 = data.get('緯度', '')
        經度 = data.get('經度', '')
        地圖縮放 = data.get('地圖縮放', '16')
        地圖圖層 = data.get('地圖圖層', 'EMAP - 通用電子地圖(內政部)')

        if 緯度 and 經度:
            try:
                # 🔥🔥🔥 優先檢查是否有使用者手動截圖（從 4.其他相關 資料夾）🔥🔥🔥
                gui_map_path = None
                if get_screenshot_path and callable(get_screenshot_path):
                    gui_map_path = get_screenshot_path()  # 🔥 使用新的路徑函數
                map_image_path = None

                if gui_map_path and os.path.exists(gui_map_path):
                    # 驗證截圖檔案
                    try:
                        from PIL import Image
                        test_img = Image.open(gui_map_path)
                        print(f"[OK] 使用使用者截圖：{gui_map_path} ({test_img.size})")
                        test_img.close()
                        map_image_path = gui_map_path
                    except Exception as e:
                        print(f"[警告] 使用者截圖損壞：{e}，將自動生成")
                        map_image_path = None
                else:
                    print(f"[提示] 未找到截圖檔案：{gui_map_path}")
                    map_image_path = None
                
                # 🔥🔥🔥 如果沒有使用者截圖，自動生成（與截圖按鈕邏輯完全一致）🔥🔥🔥
                if not map_image_path:
                    print(f"自動生成報告地圖（縮放={地圖縮放}）...")
                    
                    # 確保縮放級別在合理範圍
                    try:
                        zoom_level = int(地圖縮放)
                        if zoom_level < 10:
                            zoom_level = 16
                        elif zoom_level > 20:
                            zoom_level = 18
                    except:
                        zoom_level = 16
                    
                    # 🔥 從圖層名稱提取圖層代碼（與截圖按鈕完全一致）
                    if 'EMAP5' in 地圖圖層:  # 🔥 先檢查 EMAP5
                        layer_code = 'EMAP5'
                    elif 'PHOTO_MIX' in 地圖圖層:
                        layer_code = 'PHOTO_MIX'
                    elif 'PHOTO2' in 地圖圖層:
                        layer_code = 'PHOTO2'
                    elif 'GOOGLE_SATELLITE' in 地圖圖層:
                        layer_code = 'GOOGLE_SATELLITE'
                    elif 'GOOGLE_HYBRID' in 地圖圖層:
                        layer_code = 'GOOGLE_HYBRID'
                    elif 'GOOGLE_STREET' in 地圖圖層:
                        layer_code = 'GOOGLE_STREET'
                    elif 'OSM_DE' in 地圖圖層:
                        layer_code = 'OSM_DE'
                    elif 'OSM' in 地圖圖層:
                        layer_code = 'OSM'
                    elif 'EMAP' in 地圖圖層:  # 🔥 EMAP 在後面
                        layer_code = 'EMAP'
                    else:
                        layer_code = 'EMAP'  # 🔥 預設為 EMAP
                    
                    print(f"使用圖層：{layer_code}, 縮放：{zoom_level}")
                    
                    # 生成地圖
                    map_image_path = generate_map_image(緯度, 經度, zoom_level, layer_code)

                    # 🔥 如果自動生成成功，同步儲存到 4.其他相關 資料夾
                    if map_image_path and os.path.exists(map_image_path) and gui_map_path:
                        try:
                            import shutil
                            shutil.copy2(map_image_path, gui_map_path)
                            print(f"[OK] 已同步儲存截圖至：{gui_map_path}")
                        except Exception as e:
                            print(f"[警告] 同步儲存失敗：{e}")
                
                # 🔥 插入地圖到 Word 文件
                if map_image_path and os.path.exists(map_image_path):
                    print("[OK] 插入地圖到 Word...")
                    
                    # 清空儲存格
                    位置圖_cell.text = ''
                    
                    # 插入地圖圖片
                    paragraph = 位置圖_cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    
                    # 調整圖片寬度
                    cell_width = available_width * (7 / 14)
                    run.add_picture(map_image_path, width=cell_width - Cm(0.2))
                    
                    print("[完成] 地圖已插入 Word", flush=True)
                    
                    # 🔥 清理臨時檔案（但保留 gui_map_screenshot.png）
                    if map_image_path != gui_map_path:
                        try:
                            os.remove(map_image_path)
                            print("[OK] 臨時地圖檔案已清理")
                        except:
                            pass
                else:
                    print("[警告] 地圖生成失敗，顯示座標文字")
                    位置圖_cell.text = f"座標：\n緯度 {緯度}\n經度 {經度}\n(無法生成地圖預覽)"
                    位置圖_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            except Exception as e:
                print(f"[錯誤] 插入地圖時發生錯誤：{e}")
                import traceback
                traceback.print_exc()
                
                位置圖_cell.text = f"座標：{緯度}, {經度}\n(地圖載入失敗)"
                位置圖_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            print("[警告] 沒有座標資料")
            位置圖_cell.text = '請設定位置座標'
            位置圖_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ========================================
        # 第 20 列：承辦人（28欄版本，7+7+14）
        # ========================================
        # 經紀人簽章（7欄，cells[0-6]）
        table.rows[19].cells[0].merge(table.rows[19].cells[6]).text = '經紀人簽章：'
        # 秘書（7欄，cells[7-13]）
        table.rows[19].cells[7].merge(table.rows[19].cells[13]).text = '秘書：'

        # 🔥 處理承辦人列表（新格式）或舊格式
        承辦人列表 = data.get('承辦人列表', [])
        if not 承辦人列表:
            # 相容舊版本（單一承辦人）
            舊承辦人 = data.get('承辦人', '')
            舊電話 = data.get('電話', '')
            if 舊承辦人 and 舊電話:
                承辦人文字 = f"承辦人：{舊承辦人}{舊電話}"
            elif 舊承辦人:
                承辦人文字 = f"承辦人：{舊承辦人}"
            elif 舊電話:
                承辦人文字 = f"承辦人：{舊電話}"
            else:
                承辦人文字 = "承辦人："
        else:
            # 新格式：組合承辦人列表（格式：姓名+電話）
            # 載入通訊錄取得電話
            # 🔥 臨時強制使用備用函數以獲取詳細日誌
            load_func = _fallback_load_contacts_from_excel
            # load_func = load_contacts_from_excel if load_contacts_from_excel else _fallback_load_contacts_from_excel
            contacts_data = load_func('通訊錄.xlsx')
            contacts_by_store = contacts_data.get('contacts_by_store', {})

            # 建立姓名到電話的對照表
            contacts_dict = {}
            for store_contacts in contacts_by_store.values():
                for contact in store_contacts:
                    contacts_dict[contact['name']] = contact['phone']

            承辦人項目 = []
            for name in 承辦人列表:
                phone = contacts_dict.get(name, '')
                if phone:
                    承辦人項目.append(f"{name}{phone}")
                else:
                    承辦人項目.append(name)

            # 🔥 每3個承辦人一行，第3個後面不加「、」
            if 承辦人項目:
                lines = []
                for i in range(0, len(承辦人項目), 3):
                    # 取3個承辦人為一組
                    group = 承辦人項目[i:i+3]
                    lines.append("、".join(group))
                承辦人文字 = "承辦人：" + "\n".join(lines)
            else:
                承辦人文字 = "承辦人："

        # 承辦人（14欄，cells[14-27]）
        承辦人_cell = table.rows[19].cells[14].merge(table.rows[19].cells[27])
        承辦人_cell.text = 承辦人文字
        # 🔥 設定分散對齊
        承辦人_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        
        # ========================================
        # 設定所有儲存格樣式（統一預設）
        # ========================================
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell)
                
                # 🔥 預設：垂直置中、水平靠左
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    for run in paragraph.runs:
                        run.font.name = 'PMingLiU'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
                        run.font.size = Pt(11)  # 🔥 加大字體
                        run.font.bold = True  # 🔥 加粗
                        # 🔥 預設全部為藍色
                        run.font.color.rgb = RGBColor(0, 112, 192)

        # ========================================
        # 🔥 將標題欄位改回黑色
        # ========================================
        標題欄位 = [
            (0, 0),   # 物件編號
            (0, 4),   # 契約編號
            (0, 8),   # 案名
            (1, 0),   # 土地標示
            (2, 0),   # 建物門牌
            (5, 0),   # 主建物
            (17, 0),  # 訴求重點
            (17, 7),  # 位置圖
        ]

        for row_idx, col_idx in 標題欄位:
            cell = table.rows[row_idx].cells[col_idx]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色

        # ========================================
        # 🔥 將所有包含「：」的標籤文字改回黑色
        # ========================================
        # 這會把「總價：」「土地單價：」等標題都改回黑色
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if '：' in text:
                        # 重新處理這個儲存格
                        paragraph.clear()

                        # 分割標題和內容
                        parts = text.split('：', 1)
                        if len(parts) == 2:
                            # 標題（黑色）
                            run1 = paragraph.add_run(parts[0] + '：')
                            run1.font.name = 'PMingLiU'
                            run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
                            run1.font.size = Pt(11)  # 🔥 加大字體
                            run1.font.bold = True  # 🔥 加粗
                            run1.font.color.rgb = RGBColor(0, 0, 0)  # 黑色

                            # 內容（藍色）
                            run2 = paragraph.add_run(parts[1])
                            run2.font.name = 'PMingLiU'
                            run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
                            run2.font.size = Pt(11)  # 🔥 加大字體
                            run2.font.bold = True  # 🔥 加粗
                            run2.font.color.rgb = RGBColor(0, 112, 192)  # 藍色

        # ========================================
        # 🔥 個別儲存格對齊調整（28欄版本）
        # ========================================
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # 1. 「主建物」三個字水平置中（第 5-13 列的第 0 欄，已合併）
        table.rows[5].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 2. 第 18 列（標題列）：訴求重點和位置圖標題都置中
        table.rows[17].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 訴求重點標題
        table.rows[17].cells[14].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 位置圖標題

        # 3. 第 19 列（內容列）：訴求重點靠左靠上，位置圖內容置中
        table.rows[18].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT  # 訴求重點內容靠左
        table.rows[18].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP  # 訴求重點內容靠上
        # 位置圖內容的置中已在插入圖片時設定

        # 4. 其他標題欄位置中
        table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 物件編號
        table.rows[0].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 契約編號
        table.rows[0].cells[16].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 案名
        table.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 土地標示
        table.rows[2].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 建物門牌

        # ========================================
        # 🔥 特殊字體大小調整（在所有統一格式設定之後）
        # ========================================

        # 1. 用途儲存格（第4列，cells[14-27]）改為10號字體
        用途_cell = table.rows[3].cells[14]
        for paragraph in 用途_cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)  # 🔥 改為10號字體

        # 2. 車位儲存格（第8列，cells[21-27]）改為10號字體
        車位_cell = table.rows[7].cells[21]
        for paragraph in 車位_cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)  # 🔥 改為10號字體

        # 3. 附屬建物儲存格（第15列，cells[0-6]）- 第二行改為9號字體
        附屬_cell = table.rows[14].cells[0]
        for paragraph in 附屬_cell.paragraphs:
            text = paragraph.text
            if '\n' in text:
                # 有換行，表示有明細資訊
                lines = text.split('\n')
                # 重新建立內容
                paragraph.clear()

                # 第一行：附屬建物：X.XX坪（11號字體）
                if '：' in lines[0]:
                    parts = lines[0].split('：', 1)
                    # 標題（黑色）
                    run1 = paragraph.add_run(parts[0] + '：')
                    run1.font.name = 'PMingLiU'
                    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
                    run1.font.size = Pt(11)
                    run1.font.bold = True
                    run1.font.color.rgb = RGBColor(0, 0, 0)

                    # 內容（藍色）
                    if len(parts) == 2:
                        run2 = paragraph.add_run(parts[1])
                        run2.font.name = 'PMingLiU'
                        run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
                        run2.font.size = Pt(11)
                        run2.font.bold = True
                        run2.font.color.rgb = RGBColor(0, 112, 192)

                # 第二行：(含陽台：X坪、雨遮：X坪)（9號字體，藍色）
                if len(lines) > 1:
                    run3 = paragraph.add_run('\n' + lines[1])
                    run3.font.name = 'PMingLiU'
                    run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
                    run3.font.size = Pt(9)  # 🔥 改為9號字體
                    run3.font.bold = True
                    run3.font.color.rgb = RGBColor(0, 112, 192)

        # ========================================
        # 🔥 添加浮水印（根據承辦人編號）
        # ========================================
        try:
            # 獲取承辦人列表中的第一個人
            承辦人列表 = data.get('承辦人列表', [])
            if not 承辦人列表:
                # 相容舊版本（單一承辦人）
                舊承辦人 = data.get('承辦人', '')
                if 舊承辦人:
                    承辦人列表 = [舊承辦人]

            if 承辦人列表:
                # 取第一個承辦人
                first_employee = 承辦人列表[0]

                # 載入通訊錄
                load_func = _fallback_load_contacts_from_excel
                contacts_data = load_func('通訊錄.xlsx')

                # 獲取員工編號
                employee_code = get_employee_code_from_contacts(first_employee, contacts_data)

                if employee_code:
                    # 格式化浮水印文字（加空格）
                    watermark_text = format_watermark_text(employee_code)
                    print(f"正在添加浮水印：{watermark_text}")

                    # 添加浮水印
                    add_watermark_to_word(doc, watermark_text)
                    print(f"✓ 浮水印添加成功")
                else:
                    print(f"警告：找不到員工 {first_employee} 的編號，跳過浮水印")
            else:
                print("警告：沒有承辦人資訊，跳過浮水印")
        except Exception as e:
            print(f"警告：添加浮水印時發生錯誤：{e}")
            import traceback
            traceback.print_exc()

        # 儲存文件
        doc.save(output_path)
        
    except Exception as e:
        print(f"生成售件報告失敗：{e}")
        import traceback
        traceback.print_exc()
        raise


def generate_word_report_rent(output_path, data, report_date):
    """生成 Word 格式報告（租件）- 完全重寫版本"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from PIL import Image
        import os
        import shutil

        doc = Document()

        # 設定頁面為 A4 直向
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        # 🔥 設定中文字體為新細明體（PMingLiU）
        style = doc.styles['Normal']
        style.font.name = 'PMingLiU'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
        style.font.size = Pt(11)  # 🔥 加大字體
        style.font.bold = True  # 🔥 加粗

        # === 標題區域（與售件相同格式）===
        from docx.oxml import OxmlElement

        # 建立一個 1 列 3 欄的表格
        header_table = doc.add_table(rows=1, cols=3)
        header_table.autofit = False
        header_table.allow_autofit = False

        # 設定欄寬
        header_table.columns[0].width = Cm(5)   # LOGO 欄
        header_table.columns[1].width = Cm(10)  # 標題欄
        header_table.columns[2].width = Cm(3)   # 日期欄

        # 隱藏表格邊框
        def remove_table_borders(table):
            """移除表格所有邊框"""
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tblBorders.append(border)
            tblPr.append(tblBorders)

        remove_table_borders(header_table)

        # 設定列高
        header_table.rows[0].height = Cm(1.5)

        # 左側：LOGO
        logo_cell = header_table.rows[0].cells[0]
        logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        logo_path = "logo-1.png"
        if os.path.exists(logo_path):
            run_logo = logo_para.add_run()
            run_logo.add_picture(logo_path, height=Cm(1.0))

        # 中間：標題
        title_cell = header_table.rows[0].cells[1]
        title_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        title_para = title_cell.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_title = title_para.add_run('物件個案調查表(租)')
        run_title.font.name = 'PMingLiU'
        run_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
        run_title.font.size = Pt(18)
        run_title.font.bold = True

        # 右側：日期
        date_cell = header_table.rows[0].cells[2]
        date_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        date_para = date_cell.paragraphs[0]
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run_date = date_para.add_run(f'製表日期：{report_date}')
        run_date.font.name = 'PMingLiU'
        run_date._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
        run_date.font.size = Pt(11)  # 🔥 加大字體
        run_date.font.bold = True  # 🔥 加粗

        # 建立主表格（21行14列）
        table = doc.add_table(rows=21, cols=14)
        table.style = 'Table Grid'

        # 🔥 設定表格填滿整頁
        # 計算可用寬度
        available_width = section.page_width - section.left_margin - section.right_margin

        # 設定表格寬度
        table.width = available_width

        # 🔥 精確設定每列高度
        列高設定 = {
            0: Cm(0.9),   # 物件編號、契約編號、案名
            1: Cm(0.9),   # 土地標示
            2: Cm(0.9),   # 建物門牌
            3: Cm(0.9),   # 型式、用途
            4: Cm(0.95),   # 租金、押金、租期（灰色列，原0.9+0.1）
            5: Cm(0.95),   # 總地坪、總建坪、增建面積（灰色列，原0.9+0.1）
            6: Cm(0.95),   # 樓層（灰色列，原1.0+0.1）
            7: Cm(0.95),   # 竣工日期、座向、建築結構（灰色列，原0.9+0.1）
            8: Cm(0.95),   # 調幅、整地期、管理費（灰色列，原0.9+0.1）
            9: Cm(0.95),   # 現況、電梯（灰色列，原0.9+0.1）
            10: Cm(0.95),  # 車位、車位型態、租約公證（灰色列，原0.9+0.1）
            11: Cm(0.95),  # 工廠登記、所在商圈（灰色列，原0.9+0.1）
            12: Cm(0.95),  # 建蔽率、容積率、滴水、中脊（灰色列，原0.9+0.1）
            13: Cm(0.95),  # 面臨道路、面寬、深度、鑰匙（灰色列，原0.9+0.1）
            14: Cm(0.95),  # 都市使用分區、非都市使用地類別（灰色列，原0.9+0.1）
            15: Cm(1.2),  # 訴求重點
            16: Cm(0.9),  # 適合行業別
            17: Cm(0.9),  # 起造人名義、位置圖標題
            18: Cm(0.95),  # 格局、位置圖內容（第1列）
            19: Cm(2.0),  # 稅賦、位置圖內容（第2列）
            20: Cm(0.9),  # 承辦人
        }

        # 套用列高設定
        for 列索引, 高度 in 列高設定.items():
            table.rows[列索引].height = 高度

        # 設定欄寬（第0欄較窄，其餘13欄平均分配剩餘寬度）
        first_col_width = Cm(1.5)  # 第0欄寬度（可調整此數值）
        other_col_width = (available_width - first_col_width) / 13  # 其餘13欄平均分配

        for row in table.rows:
            for cell_idx, cell in enumerate(row.cells):
                if cell_idx == 0:
                    cell.width = first_col_width
                else:
                    cell.width = other_col_width

        # ========================================
        # 第 1 列：物件編號、契約編號、案名
        # ========================================
        table.rows[0].cells[0].text = "物件\n編號"  # 🔥 兩字折行
        table.rows[0].cells[1].merge(table.rows[0].cells[3]).text = data.get('物件編號', '')
        table.rows[0].cells[4].text = "契約\n編號"  # 🔥 兩字折行
        table.rows[0].cells[5].merge(table.rows[0].cells[7]).text = data.get('契約編號', '')
        table.rows[0].cells[8].text = "案  名"
        table.rows[0].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 🔥 案名項目名稱置中
        table.rows[0].cells[9].merge(table.rows[0].cells[13]).text = data.get('案名', '')

        # ========================================
        # 第 2 列：土地標示
        # ========================================
        table.rows[1].cells[0].text = "土地\n標示"  # 🔥 兩字折行
        table.rows[1].cells[1].merge(table.rows[1].cells[13]).text = data.get('土地標示', '')

        # ========================================
        # 第 3 列：建物門牌
        # ========================================
        table.rows[2].cells[0].text = "建物\n門牌"  # 🔥 兩字折行
        table.rows[2].cells[1].merge(table.rows[2].cells[13]).text = data.get('建物門牌', '')

        # ========================================
        # 第 4 列：型式、用途（帶方框選項）
        # ========================================
        型式 = data.get('型式', '')
        型式選項 = ["土地", "透天", "大樓", "華廈", "公寓"]
        型式文字 = "型式：" + "  ".join([f"{'■' if opt == 型式 else '□'}{opt}" for opt in 型式選項])
        table.rows[3].cells[0].merge(table.rows[3].cells[5]).text = 型式文字  # 🔥 改為合併到第5欄（縮小型式區域）

        用途 = data.get('用途', '')
        用途選項 = ["建地", "工業地", "廠房", "店面", "辦公", "住家", "農地", "其他"]
        用途文字 = "用途：" + "  ".join([f"{'■' if opt == 用途 else '□'}{opt}" for opt in 用途選項])
        table.rows[3].cells[6].merge(table.rows[3].cells[13]).text = 用途文字  # 🔥 改為從第6欄開始（擴大用途區域）

        # ========================================
        # 第 5 列：租金、押金、租期（三個項目平均分配）
        # ========================================
        租金 = data.get('租金', '')
        押金 = data.get('押金', '')
        租期 = data.get('租期', '')
        租金_含稅 = data.get('租金_含稅', True)  # 預設為含稅

        # 根據含稅勾選狀態顯示租金
        租金_cell = table.rows[4].cells[0].merge(table.rows[4].cells[4])

        if 租金:
            if 租金_含稅:
                # 勾選含稅：顯示「含稅X萬元」
                租金_cell.text = f"租金：含稅 {租金} 萬元／月"
            else:
                # 未勾選含稅：檢查是否需要顯示含稅金額
                try:
                    租金_元 = float(租金) * 10000
                    # 只有當租金 >= 20,000 元時才顯示含稅金額
                    if 租金_元 >= 20000:
                        含稅金額_元 = calculate_rent_with_tax(租金)
                        if 含稅金額_元:
                            # 設定文字內容（兩行）
                            租金_cell.text = f"租金：未稅 {租金} 萬元／月"

                            # 新增第二行（含稅金額）
                            paragraph = 租金_cell.paragraphs[0]
                            paragraph.add_run(f"\n(含稅 {含稅金額_元:,} 元／月)")

                            # 設定第二行的字體大小為 10
                            for run in paragraph.runs:
                                if "(含稅" in run.text:
                                    run.font.size = Pt(10)
                        else:
                            租金_cell.text = f"租金：未稅 {租金} 萬元／月"
                    else:
                        # 低於 19,999 元，不顯示含稅金額
                        租金_cell.text = f"租金：未稅 {租金} 萬元／月"
                except:
                    租金_cell.text = f"租金：未稅 {租金} 萬元／月"
        else:
            租金_cell.text = "租金：含稅        萬元／月"
        table.rows[4].cells[5].merge(table.rows[4].cells[9]).text = f"押金：{押金}" if 押金 else "押金："
        table.rows[4].cells[10].merge(table.rows[4].cells[13]).text = f"租期：{租期} 年" if 租期 else "租期：      年"

        # ========================================
        # 第 6 列：總地坪、總建坪、增建面積
        # ========================================
        總地坪 = data.get('總地坪', '')
        總建坪 = data.get('總建坪', '')
        增建面積 = data.get('增建面積', '')

        table.rows[5].cells[0].merge(table.rows[5].cells[4]).text = f"總地坪：{總地坪} 坪" if 總地坪 else "總地坪："
        table.rows[5].cells[5].merge(table.rows[5].cells[9]).text = f"總建坪：{總建坪} 坪" if 總建坪 else "總建坪："
        table.rows[5].cells[10].merge(table.rows[5].cells[13]).text = f"增建面積：{增建面積} 坪" if 增建面積 else "增建面積："

        # ========================================
        # 第 7 列：樓層（所有樓層資料在同一列）
        # ========================================
        主建物明細 = data.get('主建物', {})
        附屬建物明細 = data.get('附屬建物', {})

        樓層文字列表 = []
        for 樓層名, 坪數 in 主建物明細.items():
            樓層文字列表.append(f"{樓層名}：{坪數}坪")
        for 類型, 坪數 in 附屬建物明細.items():
            樓層文字列表.append(f"{類型}：{坪數}坪")

        樓層文字 = "樓層：" + "， ".join(樓層文字列表) if 樓層文字列表 else "樓層："
        table.rows[6].cells[0].merge(table.rows[6].cells[13]).text = 樓層文字

        # ========================================
        # 第 8 列：竣工日期、座向、建築結構（三項平均分配）
        # ========================================
        竣工日期 = data.get('竣工日期', '')
        座向 = data.get('座向', '')
        建築結構 = data.get('建築結構', '')

        table.rows[7].cells[0].merge(table.rows[7].cells[4]).text = f"竣工日期：{竣工日期}"
        table.rows[7].cells[5].merge(table.rows[7].cells[9]).text = f"座向：{座向}"
        table.rows[7].cells[10].merge(table.rows[7].cells[13]).text = f"建築結構：{建築結構}"

        # ========================================
        # 第 9 列：調幅、整地期、管理費
        # ========================================
        調幅 = data.get('調幅', '')
        整地期 = data.get('整地期', '')
        管理費 = data.get('管理費', '')

        table.rows[8].cells[0].merge(table.rows[8].cells[4]).text = f"調幅：{調幅}" if 調幅 else "調幅："
        table.rows[8].cells[5].merge(table.rows[8].cells[9]).text = f"整地期：{整地期}" if 整地期 else "整地期："
        table.rows[8].cells[10].merge(table.rows[8].cells[13]).text = f"管理費：{管理費} 元/月" if 管理費 else "管理費："

        # ========================================
        # 第 10 列：現況、電梯
        # ========================================
        現況 = data.get('現況', '')
        現況選項 = ["空屋/空地", "出租中", "自用"]
        現況文字 = "現況：" + "  ".join([f"{'■' if opt in 現況 or (opt == '空屋/空地' and ('空屋' in 現況 or '空地' in 現況)) else '□'}{opt}" for opt in 現況選項])

        電梯數量 = data.get('電梯數量', '')

        table.rows[9].cells[0].merge(table.rows[9].cells[9]).text = 現況文字
        table.rows[9].cells[10].merge(table.rows[9].cells[13]).text = f"電梯：{電梯數量} 部" if 電梯數量 else "電梯：      部"

        # ========================================
        # 第 11 列：車位、車位型態、租約公證（同一列）
        # ========================================
        車位資料 = data.get('車位', {})
        車位有無 = 車位資料.get('有無', '')
        車位編號 = 車位資料.get('編號', '')
        車位型式 = 車位資料.get('型式', '') if isinstance(車位資料, dict) else ''
        車位型態文字 = "車位型態：" + "  ".join([f"{'■' if opt in 車位型式 else '□'}{opt}" for opt in ["平面", "機械"]])

        租約公證 = data.get('租約公證', '')
        租約公證文字 = "租約公證：" + "  ".join([f"{'■' if opt == 租約公證 else '□'}{opt}" for opt in ["需要", "不需要"]])

        table.rows[10].cells[0].merge(table.rows[10].cells[4]).text = f"車位：{車位有無}，樓層及編號：{車位編號}" if 車位有無 else "車位：無，樓層及編號："
        table.rows[10].cells[5].merge(table.rows[10].cells[9]).text = 車位型態文字
        table.rows[10].cells[10].merge(table.rows[10].cells[13]).text = 租約公證文字

        # ========================================
        # 第 12 列：工廠登記、鄰近學校、所在商圈/工業區
        # ========================================
        工廠登記 = data.get('工廠登記', '')
        工廠登記文字 = "工廠登記：" + "  ".join([f"{'■' if opt == 工廠登記 else '□'}{opt}" for opt in ["可", "不可"]])
        鄰近學校 = data.get('鄰近學校', '')
        所在商圈 = data.get('所在商圈', '')

        table.rows[11].cells[0].merge(table.rows[11].cells[4]).text = 工廠登記文字
        table.rows[11].cells[5].merge(table.rows[11].cells[8]).text = f"鄰近學校：{鄰近學校}"
        table.rows[11].cells[9].merge(table.rows[11].cells[13]).text = f"所在商圈/工業區：{所在商圈}"

        # ========================================
        # 第 13 列：建蔽率、容積率、滴水、中脊（順序修正）
        # ========================================
        建蔽率 = data.get('建蔽率', '')
        容積率 = data.get('容積率', '')
        滴水 = data.get('滴水', '')
        中脊 = data.get('中脊', '')

        table.rows[12].cells[0].merge(table.rows[12].cells[3]).text = f"建蔽率：{建蔽率}%" if 建蔽率 else "建蔽率："
        table.rows[12].cells[4].merge(table.rows[12].cells[7]).text = f"容積率：{容積率}%" if 容積率 else "容積率："
        # 🔥 滴水和中脊加上「約」字和「公尺」單位
        table.rows[12].cells[8].merge(table.rows[12].cells[10]).text = f"滴水：約{滴水}公尺" if 滴水 else "滴水："
        table.rows[12].cells[11].merge(table.rows[12].cells[13]).text = f"中脊：約{中脊}公尺" if 中脊 else "中脊："

        # ========================================
        # 第 14 列：面臨道路、面寬、深度、鑰匙
        # ========================================
        面臨道路 = data.get('面臨道路', '')
        面寬 = data.get('面寬', '')
        深度 = data.get('深度', '')
        鑰匙 = data.get('鑰匙', '')

        table.rows[13].cells[0].merge(table.rows[13].cells[3]).text = f"面臨道路：約{面臨道路}公尺" if 面臨道路 else "面臨道路："
        table.rows[13].cells[4].merge(table.rows[13].cells[7]).text = f"面寬：約{面寬}公尺" if 面寬 else "面寬："
        table.rows[13].cells[8].merge(table.rows[13].cells[10]).text = f"深度：約{深度}公尺" if 深度 else "深度："
        table.rows[13].cells[11].merge(table.rows[13].cells[13]).text = f"鑰匙：{鑰匙}"

        # ========================================
        # 第 15 列：都市使用分區、非都市使用地類別
        # ========================================
        都市分區 = data.get('都市土地使用分區', '')
        非都市類別 = data.get('非都市土地使用地類別', '')

        table.rows[14].cells[0].merge(table.rows[14].cells[4]).text = f"都市土地使用分區：{都市分區}"
        table.rows[14].cells[5].merge(table.rows[14].cells[13]).text = f"非都市土地使用地類別：{非都市類別}"

        # ========================================
        # 第 16 列：訴求重點（項目 + 內容分開）
        # ========================================
        table.rows[15].cells[0].text = '訴求\n重點'  # 🔥 兩字折行
        # 組合訴求重點 + 另顯示於後台內容
        訴求重點 = data.get('訴求重點', '')
        後台內容 = data.get('另顯示於後台內容', '')
        # 租件：訴求重點和後台內容都將換行符轉為空格
        訴求重點_single = 訴求重點.replace('\n', ' ')
        combined_content = 訴求重點_single
        if 後台內容.strip():  # 只有在有內容時才顯示標題和內容
            後台內容_single = 後台內容.replace('\n', ' ')
            combined_content += '\n另顯示於後台內容：' + 後台內容_single
        table.rows[15].cells[1].merge(table.rows[15].cells[13]).text = combined_content
        table.rows[15].cells[1].merge(table.rows[15].cells[13]).vertical_alignment = WD_ALIGN_VERTICAL.TOP
        table.rows[15].cells[1].merge(table.rows[15].cells[13]).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ========================================
        # 第 17 列：適合行業別（項目）、適合行業別（內容）
        # ========================================
        table.rows[16].cells[0].text = '適合行業別'  # 🔥 保持不折行（四個字）
        # 🔥 適合行業別：只顯示被選中的項目
        適合行業別_data = data.get('適合行業別', {})
        選中項目 = []

        if isinstance(適合行業別_data, str):
            # 相容舊版本（字串格式）
            選中項目_text = 適合行業別_data
        else:
            # 新版本（字典格式）
            行業選項 = ["住宅", "餐飲業", "美容業", "服飾業", "娛樂休閒", "電競電玩", "運動健身",
                        "百貨超市", "生活服務", "電器通訊", "大賣場", "電商業", "倉庫", "物流業"]

            for option in 行業選項:
                if 適合行業別_data.get(option, False):
                    選中項目.append(option)

            # 處理「其他」選項
            if 適合行業別_data.get("其他_checked", False):
                其他_text = 適合行業別_data.get("其他_text", "")
                if 其他_text.strip():
                    選中項目.append(f"其他：{其他_text}")
                else:
                    選中項目.append("其他")

            選中項目_text = "、".join(選中項目) if 選中項目 else ""

        table.rows[16].cells[1].merge(table.rows[16].cells[13]).text = 選中項目_text

        # ========================================
        # 第 18 列：起造人名義（項目）、起造人名義（內容）、位置圖（標題）
        # ========================================
        table.rows[17].cells[0].text = '起造人名義'  # 🔥 保持不折行（四個字）
        # 🔥 起造人名義：顯示選項格式
        起造人名義 = data.get('起造人名義', '')
        地主_checkbox = '■' if 起造人名義 == '地主' else '□'
        承租方_checkbox = '■' if 起造人名義 == '承租方' else '□'
        起造人名義_text = f'{地主_checkbox}地主  {承租方_checkbox}承租方'
        table.rows[17].cells[1].merge(table.rows[17].cells[6]).text = 起造人名義_text
        # 🔥 位置圖標題置中
        位置圖_title_cell = table.rows[17].cells[7].merge(table.rows[17].cells[13])
        位置圖_title_cell.text = '位置圖'
        位置圖_title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ========================================
        # 第 19 列：格局（項目）、格局（內容）、位置圖（內容 - 開始合併）
        # ========================================
        格局 = data.get('格局', {})
        格局選項_data = data.get('格局選項', [])
        格局選項 = ["開放式", "空地"]
        格局選項文字 = "  ".join([f"{'■' if opt in 格局選項_data else '□'}{opt}" for opt in 格局選項])
        # 🔥 格局資訊折行顯示，無論有無輸入都要呈現
        房 = 格局.get('房', '')
        廳 = 格局.get('廳', '')
        衛 = 格局.get('衛', '')
        格局文字 = f"\n{房}房{廳}廳{衛}衛"

        table.rows[18].cells[0].text = '格局'  # 🔥 改為「格局」
        table.rows[18].cells[1].merge(table.rows[18].cells[6]).text = 格局選項文字 + 格局文字

        # 位置圖（右側 - 將在第19-20列合併）
        位置圖_cell = table.rows[18].cells[7].merge(table.rows[19].cells[13])

        # 🔥 地圖插入邏輯（與售件相同）
        緯度 = data.get('緯度', '')
        經度 = data.get('經度', '')
        地圖縮放 = data.get('地圖縮放', '16')
        地圖圖層 = data.get('地圖圖層', 'EMAP - 通用電子地圖(內政部)')

        if 緯度 and 經度:
            try:
                # 🔥 使用新的路徑函數取得截圖路徑
                gui_map_path = get_screenshot_path()
                map_image_path = None

                if gui_map_path and os.path.exists(gui_map_path):
                    try:
                        test_img = Image.open(gui_map_path)
                        test_img.verify()
                        map_image_path = gui_map_path
                        print(f"[OK] 使用使用者截圖：{gui_map_path}")
                    except:
                        map_image_path = None
                        print(f"[警告] 截圖檔案損壞，將自動生成")
                else:
                    print(f"[提示] 未找到截圖檔案：{gui_map_path}")

                if not map_image_path:
                    print(f"自動生成報告地圖（縮放={地圖縮放}）...")
                    zoom_level = int(地圖縮放) if 地圖縮放 else 16

                    # 提取圖層代碼
                    if 'EMAP5' in 地圖圖層:
                        layer_code = 'EMAP5'
                    elif 'PHOTO_MIX' in 地圖圖層:
                        layer_code = 'PHOTO_MIX'
                    elif 'PHOTO2' in 地圖圖層:
                        layer_code = 'PHOTO2'
                    elif 'EMAP' in 地圖圖層:
                        layer_code = 'EMAP'
                    else:
                        layer_code = 'EMAP'

                    map_image_path = generate_map_image(緯度, 經度, zoom_level, layer_code)

                    # 🔥 同步儲存到 4.其他相關 資料夾
                    if map_image_path and os.path.exists(map_image_path) and gui_map_path:
                        try:
                            import shutil
                            shutil.copy2(map_image_path, gui_map_path)
                            print(f"[OK] 已同步儲存截圖至：{gui_map_path}")
                        except Exception as e:
                            print(f"[警告] 同步儲存失敗：{e}")

                if map_image_path and os.path.exists(map_image_path):
                    位置圖_cell.text = ''
                    paragraph = 位置圖_cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    cell_width = Cm(18) * (7 / 14)  # 直接計算，結果會是 Cm 物件
                    run.add_picture(map_image_path, width=cell_width - Cm(0.2))
                else:
                    位置圖_cell.text = f"座標：{緯度}, {經度}"
            except Exception as e:
                位置圖_cell.text = f"座標：{緯度}, {經度}\n(地圖載入失敗)"
        else:
            位置圖_cell.text = '請設定位置座標'

        # ========================================
        # 第 20 列：稅賦（項目）、稅賦（內容）、位置圖（內容 - 繼續合併）
        # ========================================
        稅賦 = data.get('稅賦', {})
        稅賦文字列 = []
        稅賦項目 = ["所得稅", "補充保費", "營業稅", "房屋稅", "地價稅"]
        for 項目 in 稅賦項目:
            負責方 = 稅賦.get(項目, '')
            稅賦文字列.append(f"{項目}：" + "  ".join([f"{'■' if opt == 負責方 else '□'}{opt}" for opt in ["出租方", "承租方"]]))
        稅賦文字 = "\n".join(稅賦文字列)

        table.rows[19].cells[0].text = '稅賦'
        table.rows[19].cells[1].merge(table.rows[19].cells[6]).text = 稅賦文字
        table.rows[19].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # 位置圖已經在第19列合併了

        # 🔥 設定位置圖標題置中（第 17 列的第 7 欄）
        table.rows[17].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ========================================
        # 第 21 列：經紀人簽章、秘書、承辦人（加大承辦人欄位）
        # ========================================
        table.rows[20].cells[0].merge(table.rows[20].cells[2]).text = '經紀人簽章：'
        table.rows[20].cells[3].merge(table.rows[20].cells[6]).text = '秘書：'

        # 🔥 處理承辦人列表（新格式）或舊格式
        承辦人列表 = data.get('承辦人列表', [])
        if not 承辦人列表:
            # 相容舊版本（單一承辦人）
            舊承辦人 = data.get('承辦人', '')
            舊電話 = data.get('電話', '')
            承辦人文字 = f'承辦人：{舊承辦人}  {舊電話}'
        else:
            # 新格式：組合承辦人列表（格式：姓名+電話）
            # 🔥 臨時強制使用備用函數以獲取詳細日誌
            load_func = _fallback_load_contacts_from_excel
            # load_func = load_contacts_from_excel if load_contacts_from_excel else _fallback_load_contacts_from_excel
            contacts_data = load_func('通訊錄.xlsx')
            contacts_by_store = contacts_data.get('contacts_by_store', {})

            # 建立姓名到電話的對照表
            contacts_dict = {}
            for store_contacts in contacts_by_store.values():
                for contact in store_contacts:
                    contacts_dict[contact['name']] = contact['phone']

            承辦人項目 = []
            for name in 承辦人列表:
                phone = contacts_dict.get(name, '')
                if phone:
                    承辦人項目.append(f"{name}{phone}")
                else:
                    承辦人項目.append(name)

            # 🔥 每3個承辦人一行，第3個後面不加「、」
            if 承辦人項目:
                lines = []
                for i in range(0, len(承辦人項目), 3):
                    # 取3個承辦人為一組
                    group = 承辦人項目[i:i+3]
                    lines.append("、".join(group))
                承辦人文字 = "承辦人：" + "\n".join(lines)
            else:
                承辦人文字 = "承辦人："

        承辦人_cell = table.rows[20].cells[7].merge(table.rows[20].cells[13])
        承辦人_cell.text = 承辦人文字
        # 🔥 設定分散對齊
        承辦人_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE

        # ========================================
        # 設定所有儲存格格式、對齊方式、顏色
        # ========================================
        # 需要水平置中的項目名稱列表（只有這些項目名稱才置中）
        centered_labels = ['物件\n編號', '契約\n編號', '案  名', '土地\n標示', '建物\n門牌', '訴求\n重點', '格局', '稅賦']
        # 需要水平靠左的項目名稱列表
        left_aligned_labels = ['適合行業別', '起造人名義']

        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                set_cell_border(cell)

                # 設定垂直對齊：訴求重點、格局、稅賦靠上，其他置中
                if row in [table.rows[15], table.rows[18], table.rows[19]]:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                else:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # 取得儲存格文字內容
                cell_text = cell.text.strip()
                is_label_cell = (cell_idx == 0)  # 第0欄是項目名稱欄

                # 設定字體和顏色
                for paragraph in cell.paragraphs:
                    # 設定水平對齊
                    if is_label_cell:
                        # 項目名稱欄的對齊方式
                        if cell_text in centered_labels:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif cell_text in left_aligned_labels:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # 其他不設定，保持預設（靠左）
                    else:
                        # 內容欄：特殊處理
                        if row_idx == 15:  # 訴求重點內容靠左
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif row_idx in [16, 17]:  # 適合行業別、起造人名義內容也靠左
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif row_idx in [0, 1, 2] and cell_idx in [1, 5, 9]:  # 前三列的內容（物件編號、契約編號、案名、土地標示、建物門牌）靠左
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # 其他保持預設（靠左）

                    # 🔥 處理同一儲存格內「項目名稱：內容」的顏色分離
                    full_text = paragraph.text
                    if '：' in full_text and not is_label_cell:
                        # 清除原有的runs
                        for run in paragraph.runs:
                            run._element.getparent().remove(run._element)

                        # 分割項目名稱和內容
                        parts = full_text.split('：', 1)
                        if len(parts) == 2:
                            # 項目名稱（黑色）
                            run1 = paragraph.add_run(parts[0] + '：')
                            run1.font.size = Pt(11)  # 🔥 加大字體
                            run1.font.bold = True  # 🔥 加粗
                            run1.font.name = 'PMingLiU'
                            run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
                            run1.font.color.rgb = RGBColor(0, 0, 0)  # 黑色

                            # 內容（藍色）
                            run2 = paragraph.add_run(parts[1])
                            run2.font.size = Pt(11)  # 🔥 加大字體
                            run2.font.bold = True  # 🔥 加粗
                            run2.font.name = 'PMingLiU'
                            run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
                            run2.font.color.rgb = RGBColor(0, 112, 192)  # 藍色
                        else:
                            # 沒有成功分割，全部設為黑色
                            run = paragraph.add_run(full_text)
                            run.font.size = Pt(11)  # 🔥 加大字體
                            run.font.bold = True  # 🔥 加粗
                            run.font.name = 'PMingLiU'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
                    else:
                        # 沒有冒號，或是項目名稱欄
                        for run in paragraph.runs:
                            run.font.size = Pt(11)  # 🔥 加大字體
                            run.font.bold = True  # 🔥 加粗
                            run.font.name = 'PMingLiU'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')

                            # 設定顏色
                            if is_label_cell:
                                # 項目名稱欄：黑色
                                pass  # 預設就是黑色
                            else:
                                # 內容欄且沒有冒號：全部藍色
                                run.font.color.rgb = RGBColor(0, 112, 192)  # 藍色

        # ========================================
        # 🔥 特殊字體大小調整（租件專用，在所有統一格式設定之後）
        # ========================================

        # 1. 適合行業別項目名稱儲存格（第17列，cells[0]）改為10號字體
        適合行業別_label_cell = table.rows[16].cells[0]
        for paragraph in 適合行業別_label_cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)  # 🔥 改為10號字體

        # 2. 起造人名義項目名稱儲存格（第18列，cells[0]）改為10號字體
        起造人名義_label_cell = table.rows[17].cells[0]
        for paragraph in 起造人名義_label_cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)  # 🔥 改為10號字體

        # 3. 用途儲存格（第4列，cells[6-13]）改為10號字體
        用途_cell = table.rows[3].cells[6]
        for paragraph in 用途_cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)  # 🔥 改為10號字體

        # 4. 型式儲存格（第4列，cells[0-5]）改為10號字體
        型式_cell = table.rows[3].cells[0]
        for paragraph in 型式_cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)  # 🔥 改為10號字體

        # 5. 租金含稅金額（第5列，cells[0-4]，包含「(含稅 XX,XXX 元／月)」的部分）改為10號字體
        租金_cell = table.rows[4].cells[0]
        for paragraph in 租金_cell.paragraphs:
            for run in paragraph.runs:
                if "(含稅" in run.text or "元／月)" in run.text:
                    run.font.size = Pt(10)  # 🔥 含稅金額改為10號字體

        # 6. 車位編號（第11列，cells[0-4]，「樓層及編號：」之後的部分）改為10號字體
        車位_cell = table.rows[10].cells[0]
        for paragraph in 車位_cell.paragraphs:
            full_text = paragraph.text
            # 如果包含「樓層及編號：」，則編號部分改為10號字體
            if "樓層及編號：" in full_text:
                # 清除現有的 runs
                for run in paragraph.runs[:]:
                    run._element.getparent().remove(run._element)

                # 重新分割文字
                parts = full_text.split("樓層及編號：")
                if len(parts) == 2:
                    # 前半部分（正常大小11）
                    run1 = paragraph.add_run(parts[0] + "樓層及編號：")
                    run1.font.size = Pt(11)
                    run1.font.bold = True
                    run1.font.name = 'PMingLiU'
                    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
                    run1.font.color.rgb = RGBColor(0, 112, 192)  # 藍色

                    # 後半部分（編號，10號字體）
                    if parts[1]:  # 如果有編號
                        run2 = paragraph.add_run(parts[1])
                        run2.font.size = Pt(10)  # 🔥 編號改為10號字體
                        run2.font.bold = True
                        run2.font.name = 'PMingLiU'
                        run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
                        run2.font.color.rgb = RGBColor(0, 112, 192)  # 藍色

        # ========================================
        # 🔥 添加浮水印（根據承辦人編號）
        # ========================================
        try:
            # 獲取承辦人列表中的第一個人
            承辦人列表 = data.get('承辦人列表', [])
            if not 承辦人列表:
                # 相容舊版本（單一承辦人）
                舊承辦人 = data.get('承辦人', '')
                if 舊承辦人:
                    承辦人列表 = [舊承辦人]

            if 承辦人列表:
                # 取第一個承辦人
                first_employee = 承辦人列表[0]

                # 載入通訊錄
                load_func = _fallback_load_contacts_from_excel
                contacts_data = load_func('通訊錄.xlsx')

                # 獲取員工編號
                employee_code = get_employee_code_from_contacts(first_employee, contacts_data)

                if employee_code:
                    # 格式化浮水印文字（加空格）
                    watermark_text = format_watermark_text(employee_code)
                    print(f"正在添加浮水印：{watermark_text}")

                    # 添加浮水印
                    add_watermark_to_word(doc, watermark_text)
                    print(f"✓ 浮水印添加成功")
                else:
                    print(f"警告：找不到員工 {first_employee} 的編號，跳過浮水印")
            else:
                print("警告：沒有承辦人資訊，跳過浮水印")
        except Exception as e:
            print(f"警告：添加浮水印時發生錯誤：{e}")
            import traceback
            traceback.print_exc()

        # 儲存文件
        doc.save(output_path)
        print(f"[完成] 租件報告已生成：{output_path}", flush=True)

    except Exception as e:
        print(f"生成租件報告失敗：{e}")
        import traceback
        traceback.print_exc()
        raise

def parse_land_lots_from_data(data):
    """
    從 data 字典中解析出地段地號列表

    參數:
        data: dict，包含 '土地標示' 欄位，例如：
              "高雄市大寮區義堂段358-21地號"
              "高雄市大寮區義堂段358-21、359地號"

    回傳:
        list of dict: [{'地段': '大寮區義堂段', '地號': '358-21'}, ...]

    注意：地段和地號必須完全匹配，包括母號和子號
          例如：358-21 和 358-22 是不同的地號
    """
    import re

    try:
        land_text = data.get('土地標示', '')
        if not land_text:
            print("[提示] 土地標示為空")
            return []

        print(f"[解析] 土地標示原文: '{land_text}'")
        result = []

        # 移除「地號」、「高雄市」、「台南市」等文字
        land_text = re.sub(r'地號$', '', land_text)
        land_text = re.sub(r'^(高雄市|台南市|臺南市)', '', land_text)
        land_text = land_text.strip()

        print(f"[解析] 清理後: '{land_text}'")

        # 分割地段和地號部分
        # 🔥 用 rsplit 從右邊切「段」字，避免地段名稱含多個段字（例如「大港段二小段」、「新興段四小段」）時切錯
        # 例如：「大寮區義堂段358-21、359」→ 地段='大寮區義堂段', 地號='358-21、359'
        #       「新興區新興段四小段 1451」→ 地段='新興區新興段四小段', 地號='1451'
        if '段' in land_text:
            parts = land_text.rsplit('段', 1)
            section = (parts[0] + '段').strip()
            lot_numbers_text = parts[1].strip() if len(parts) > 1 else ''

            print(f"[解析] 地段: '{section}'")
            print(f"[解析] 地號文字: '{lot_numbers_text}'")

            # 分割多個地號（用頓號、逗號、空格分隔）
            lot_numbers = re.split(r'[、,\s]+', lot_numbers_text)

            for lot_num in lot_numbers:
                lot_num = lot_num.strip()
                if lot_num:
                    result.append({
                        '地段': section,
                        '地號': lot_num
                    })
                    print(f"[解析] 新增地號: 地段='{section}' 地號='{lot_num}'")

        print(f"[解析] 完成，共解析出 {len(result)} 筆地段地號")
        return result

    except Exception as e:
        print(f"[警告] 解析土地標示失敗：{e}")
        import traceback
        traceback.print_exc()
        return []

def open_data_editor():
    """開啟資料編輯視窗"""

    # 🔥 強化版輔助函數：綁定 Entry 的焦點事件，解決游標消失問題
    def bind_entry_focus(entry_widget):
        """為 Entry 綁定焦點事件，提升輸入體驗"""
        def on_focus(event):
            try:
                widget = event.widget
                widget.focus_set()
                widget.icursor(tk.END)
                widget.selection_clear()
                widget._is_active = True
            except:
                pass

        def on_click(event):
            try:
                widget = event.widget
                widget.focus_set()
                widget._is_active = True
                # 不 return "break"，讓 tkinter 原生處理游標點擊定位
            except:
                pass

        def on_key(event):
            """當有按鍵輸入時，確保焦點存在"""
            try:
                widget = event.widget
                if not widget.focus_get() == widget:
                    widget.focus_set()
                    widget.icursor(tk.END)
            except:
                pass

        def on_focus_out(event):
            try:
                widget = event.widget
                widget._is_active = False
            except:
                pass

        entry_widget.bind('<FocusIn>', on_focus)
        entry_widget.bind('<Button-1>', on_click)
        entry_widget.bind('<KeyPress>', on_key)
        entry_widget.bind('<FocusOut>', on_focus_out)
        # 🔥 滑鼠右鍵文字選單（剪下/複製/貼上/全選/清除）
        entry_widget.bind('<Button-3>', _show_text_menu)
        return entry_widget

    # 🔥 宣告全域變數
    global map_widget, current_map_zoom, entries
    map_widget = None
    current_map_zoom = None

    # 🔥 先建立視窗，再檢查資料
    editor_window = tk.Toplevel(root)
    editor_window.title("物件資料編輯器")

    # 🔥 防呆：捲動表單時，滑鼠滾輪若滑過下拉選單(Combobox)會誤改它的選項值
    #    （例：履約保證 同意→不同意、空白→是），存檔後就跑掉了。
    #    讓所有 Combobox 不吃滾輪（值不再被滾輪改變；展開後的清單仍可正常捲動）。
    try:
        editor_window.bind_class("TCombobox", "<MouseWheel>", lambda e: "break")
    except Exception:
        pass

    # 🔥 文字輸入欄的滑鼠右鍵選單（復原/剪下/複製/貼上/全選/清除）+ 快速鍵提示
    # 復原：Text 用原生 undo；Entry/Combobox 無原生 undo，改用自訂快照堆疊。
    _text_menu = tk.Menu(editor_window, tearoff=0)

    def _tm_is_text(w):
        try:
            return w.winfo_class() == 'Text'
        except Exception:
            return False

    def _tm_get(w):
        try:
            return w.get('1.0', 'end-1c') if _tm_is_text(w) else w.get()
        except Exception:
            return None

    def _tm_set(w, val):
        try:
            if _tm_is_text(w):
                w.delete('1.0', 'end'); w.insert('1.0', val)
            else:
                w.delete(0, 'end'); w.insert(0, val)
        except Exception:
            pass

    def _tm_push_undo(w):
        """改變內容前，把目前值存入快照堆疊（Entry/Combobox 用；Text 走原生 undo）"""
        if _tm_is_text(w):
            return
        cur = _tm_get(w)
        if cur is None:
            return
        stk = getattr(w, '_undo_stack', None)
        if stk is None:
            stk = w._undo_stack = []
        if not stk or stk[-1] != cur:
            stk.append(cur)
            if len(stk) > 100:
                del stk[0]

    def _tm_undo(w):
        try:
            if _tm_is_text(w):
                try:
                    w.edit_undo()
                except Exception:
                    pass
                return
            stk = getattr(w, '_undo_stack', None)
            if stk:
                _tm_set(w, stk.pop())
        except Exception:
            pass

    def _text_menu_action(action):
        w = getattr(_text_menu, '_target', None)
        if not w:
            return
        try:
            if action == 'undo':
                _tm_undo(w)
            elif action == 'cut':
                _tm_push_undo(w); w.event_generate('<<Cut>>')
            elif action == 'copy':
                w.event_generate('<<Copy>>')
            elif action == 'paste':
                _tm_push_undo(w); w.event_generate('<<Paste>>')
            elif action == 'all':
                if _tm_is_text(w):
                    w.tag_add('sel', '1.0', 'end-1c')
                else:
                    w.selection_range(0, tk.END); w.icursor(tk.END)
            elif action == 'clear':
                _tm_push_undo(w)
                if _tm_is_text(w):
                    w.delete('1.0', tk.END)
                else:
                    w.delete(0, tk.END)
        except Exception:
            pass

    _text_menu.add_command(label="復原", accelerator="Ctrl+Z", command=lambda: _text_menu_action('undo'))
    _text_menu.add_separator()
    _text_menu.add_command(label="剪下", accelerator="Ctrl+X", command=lambda: _text_menu_action('cut'))
    _text_menu.add_command(label="複製", accelerator="Ctrl+C", command=lambda: _text_menu_action('copy'))
    _text_menu.add_command(label="貼上", accelerator="Ctrl+V", command=lambda: _text_menu_action('paste'))
    _text_menu.add_separator()
    _text_menu.add_command(label="全選", accelerator="Ctrl+A", command=lambda: _text_menu_action('all'))
    _text_menu.add_command(label="清除", command=lambda: _text_menu_action('clear'))

    # 會改變內容、唯讀欄位需停用的項目索引：復原0、剪下2、貼上4、清除7
    _editable_idx = (0, 2, 4, 7)

    def _show_text_menu(event):
        """在被點的文字欄上彈出右鍵選單"""
        w = event.widget
        try:
            w.focus_set()
            _text_menu._target = w
            try:
                editable = str(w.cget('state')) not in ('disabled', 'readonly')
            except Exception:
                editable = True
            for idx in _editable_idx:
                _text_menu.entryconfig(idx, state=(tk.NORMAL if editable else tk.DISABLED))
            _text_menu.tk_popup(event.x_root, event.y_root)
        finally:
            _text_menu.grab_release()

    def _ctrl_a_select_all(event):
        """Ctrl+A 改為全選（tkinter 預設是移到行首）"""
        w = event.widget
        try:
            if _tm_is_text(w):
                w.tag_add('sel', '1.0', 'end-1c')
            else:
                w.selection_range(0, tk.END); w.icursor(tk.END)
        except Exception:
            pass
        return 'break'

    def _ctrl_z_undo(event):
        _tm_undo(event.widget)
        return 'break'

    def _snapshot_undo(event):
        """進入欄位時存快照（Entry/Combobox 復原用）"""
        _tm_push_undo(event.widget)

    def _bind_text_menu_recursive(w):
        """遞迴把右鍵文字選單 + 快速鍵綁到所有可輸入文字的元件（Entry / Text / 可輸入 Combobox）"""
        try:
            cls = w.winfo_class()
            if cls in ('Entry', 'TEntry', 'Text', 'TCombobox'):
                w.bind('<Button-3>', _show_text_menu)      # 右鍵選單
                w.bind('<Control-a>', _ctrl_a_select_all)  # 全選
                w.bind('<Control-A>', _ctrl_a_select_all)
                w.bind('<Control-z>', _ctrl_z_undo)        # 復原
                w.bind('<Control-Z>', _ctrl_z_undo)
                if cls == 'Text':
                    try:
                        w.config(undo=True, autoseparators=True, maxundo=-1)  # 原生 undo
                    except Exception:
                        pass
                else:
                    w.bind('<FocusIn>', _snapshot_undo, add='+')  # 進欄位存快照供復原
        except Exception:
            pass
        try:
            for child in w.winfo_children():
                _bind_text_menu_recursive(child)
        except Exception:
            pass

    # （右鍵選單的綁定改在所有 UI 建好後、函式結尾統一遞迴掃描，較可靠）

    # 🔥 計算可用寬度和高度（佔滿左側螢幕）
    screen_width = editor_window.winfo_screenwidth()
    screen_height = editor_window.winfo_screenheight()
    main_window_width = root.winfo_width()
    main_window_x = root.winfo_x()
    
    # 🔥 編輯視窗佔滿主視窗左側的所有空間
    available_width = main_window_x - 10  # 扣除少許邊距
    available_height = screen_height - 70  # 扣除工作列和標題列
    
    # 🔥 設定視窗大小（佔滿左側）
    editor_window.geometry(f"{available_width}x{available_height}")
    
    # 🔥 置於螢幕左上角
    editor_x = 0
    editor_y = 0
    editor_window.geometry(f"+{editor_x}+{editor_y}")

    # 🔥 移除 transient 以恢復最小化按鈕
    # editor_window.transient(root)  # 移除：transient 會隱藏最小化按鈕
    # 🔥 移除 grab_set() 以避免輸入欄位無法輸入的問題
    # editor_window.grab_set()

    # 🔥 確保視窗有標準的標題欄按鈕（最小化、最大化、關閉）
    editor_window.attributes('-toolwindow', False)  # 確保不是工具視窗
    editor_window.resizable(True, True)  # 允許調整大小

    # 🔥 全局焦點保護機制 - 防止輸入時游標消失
    last_focused_widget = [None]  # 記錄最後有焦點的輸入欄位

    def protect_entry_focus(event=None):
        """保護輸入欄位的焦點不被搶走"""
        try:
            current_focus = editor_window.focus_get()
            # Entry 和 Text 都需要保護
            if isinstance(current_focus, (tk.Entry, tk.Text)):
                if hasattr(current_focus, '_is_active') and current_focus._is_active:
                    last_focused_widget[0] = current_focus
        except:
            pass

    def restore_entry_focus(event=None):
        """如果焦點被意外搶走，嘗試恢復"""
        try:
            current_focus = editor_window.focus_get()
            # 如果焦點不是 Entry 或 Text，但有記錄的活動欄位，恢復之
            if not isinstance(current_focus, (tk.Entry, tk.Text)) and last_focused_widget[0]:
                last_widget = last_focused_widget[0]
                if hasattr(last_widget, '_is_active') and last_widget._is_active:
                    last_widget.focus_set()
                    if isinstance(last_widget, tk.Entry):
                        last_widget.icursor(tk.END)
        except:
            pass

    # 定期檢查焦點狀態（每 100ms）
    def check_focus_periodically():
        protect_entry_focus()
        editor_window.after(100, check_focus_periodically)

    editor_window.after(500, check_focus_periodically)  # 延遲啟動，等視窗完全載入

    # ========================================
    # 🔥 再檢查資料是否存在
    # ========================================
    data_final_path = get_data_final_path()
    if not os.path.exists(data_final_path):
        # 如果沒有 data_final.json，檢查是否有已選擇的 JSON
        selected_json = get_selected_json_path() if get_selected_json_path else None
        if selected_json and os.path.exists(selected_json):
            result = show_large_yesno(
                "尚未轉換",
                "尚未轉換為物調資料\n\n是否要立即轉換已選擇的 JSON？"
            )
            if result:
                convert_transcript_to_final_data(selected_json)
                import time
                time.sleep(0.3)
            else:
                editor_window.destroy()  # 🔥 關閉視窗
                return
        else:
            result = show_large_yesno(
                "找不到資料",
                "尚未建立物調資料\n\n是否要選擇結構化 JSON 建立？"
            )
            if result:
                choose_json_file()
                editor_window.destroy()  # 🔥 關閉視窗
                return
            else:
                editor_window.destroy()  # 🔥 關閉視窗
                return
    
    # 載入現有資料
    try:
        print(f"正在讀取：{data_final_path}")

        with open(data_final_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f"成功讀取資料，資料內容：")
        print(f"土地標示：{data.get('土地標示', '無')}")
        print(f"建物門牌：{data.get('建物門牌', '無')}")
        print(f"總地坪：{data.get('總地坪', '無')}")

        # 🔥 顯示更詳細的資料載入訊息
        update_message("[資訊] 已載入物件資料：")
        update_message(f"  土地標示：{data.get('土地標示', '無')}")
        update_message(f"  建物門牌：{data.get('建物門牌', '無')}")
        update_message(f"  總地坪：{data.get('總地坪', '無')}")
        update_message(f"  總建坪：{data.get('總建坪', '無')}")

    except Exception as e:
        print(f"讀取失敗：{e}")
        messagebox.showwarning("警告", f"無法讀取 data_final.json：{e}")
        editor_window.destroy()  # 🔥 關閉視窗
        return

    # 🔥🔥🔥 載入中遮罩 🔥🔥🔥
    # 物件編輯器 UI 元件多、地圖瓦片非同步載入較慢，會逐步顯示半成品介面。
    # 用一個 topmost 的遮罩視窗蓋住整個編輯器，等所有 UI 建立完成才移除。
    # 放在「資料載入成功後」，避開前面資料檢查的提前 return（否則遮罩會變孤兒視窗）。
    loading_overlay = tk.Toplevel(editor_window)
    loading_overlay.overrideredirect(True)       # 無邊框
    loading_overlay.attributes('-topmost', True)  # 永遠在最上層，蓋住陸續建立的 UI
    loading_overlay.configure(bg='#2c3e50')
    try:
        loading_overlay.geometry(f"{available_width}x{available_height}+{editor_x}+{editor_y}")
    except Exception:
        pass
    tk.Label(
        loading_overlay,
        text="⏳ 載入中，請稍候...\n\n正在準備物件編輯介面與地圖",
        font=("Microsoft JhengHei", 18, "bold"),
        bg='#2c3e50', fg='white', justify='center'
    ).place(relx=0.5, rely=0.5, anchor='center')
    try:
        loading_overlay.update()
    except Exception:
        pass

    def _remove_loading_overlay():
        """移除載入中遮罩（可重複呼叫，安全）"""
        try:
            if loading_overlay.winfo_exists():
                loading_overlay.destroy()
        except Exception:
            pass

    # 防呆 1：若編輯器視窗被關閉，遮罩也跟著移除（避免變成孤兒視窗）
    def _on_editor_destroy(_e):
        if _e.widget is editor_window:
            _remove_loading_overlay()
    editor_window.bind("<Destroy>", _on_editor_destroy, add="+")
    # 防呆 2：最多 30 秒後強制移除（避免中途出錯遮罩永久卡住）
    editor_window.after(30000, _remove_loading_overlay)

    # 🔥 建立標準子資料夾結構
    try:
        if os.path.exists(get_data_json_path()):
            with open(get_data_json_path(), 'r', encoding='utf-8') as f:
                data_list = json.load(f)
                if data_list:
                    first_data = data_list[0]
                    area = first_data.get('area', '')
                    section = first_data.get('section', '')
                    lot_number = first_data.get('lot_number', '')
                    folder_name = get_work_folder(f"{area}{section}-{lot_number}")

                    # 定義子資料夾結構
                    subdirectories = [
                        "0.謄本",
                        "1.基本資料",
                        "2.照片",
                        "3.行情",
                        "4.其他相關",
                        "5.委託",
                        "6.成交"
                    ]

                    # 確保主資料夾存在
                    if not os.path.exists(folder_name):
                        os.makedirs(folder_name)
                        update_message(f"[建立] 主資料夾：{folder_name}")

                    # 建立子目錄
                    created_dirs = []
                    for subdir in subdirectories:
                        subdir_path = os.path.join(folder_name, subdir)
                        if not os.path.exists(subdir_path):
                            os.makedirs(subdir_path)
                            created_dirs.append(subdir)

                    # 建立 1.基本資料/png 子目錄
                    png_dir = os.path.join(folder_name, "1.基本資料", "png")
                    if not os.path.exists(png_dir):
                        os.makedirs(png_dir)
                        created_dirs.append("1.基本資料/png")

                    if created_dirs:
                        update_message(f"[建立] 子資料夾：{', '.join(created_dirs)}")
    except Exception as e:
        print(f"建立資料夾時發生錯誤：{e}")

    # 🔥🔥🔥 修正：從 data.json 讀取並補充座標資料 🔥🔥🔥
    try:
        if os.path.exists(get_data_json_path()):
            print("正在檢查 data.json 座標...")
            with open(get_data_json_path(), 'r', encoding='utf-8') as f:
                data_json_list = json.load(f)
                if data_json_list and len(data_json_list) > 0:
                    first_data = data_json_list[0]
                    coordinates = first_data.get('coordinates', '')
                    
                    if coordinates and ',' in coordinates:
                        parts = coordinates.split(',')
                        if len(parts) == 2:
                            lat = parts[0].strip()
                            lng = parts[1].strip()
                            
                            # 🔥 如果 data_final.json 中沒有緯經度，則從 data.json 補充
                            if not data.get('緯度') or not data.get('經度'):
                                data['緯度'] = lat
                                data['經度'] = lng
                                print(f"[OK] 從 data.json 補充座標：緯度 {lat}, 經度 {lng}")
                            else:
                                print(f"[OK] 已有座標：緯度 {data.get('緯度')}, 經度 {data.get('經度')}")
                            
                            # 🔥 確保地圖縮放和圖層有預設值
                            if not data.get('地圖縮放'):
                                data['地圖縮放'] = '16'
                                print("[OK] 設定預設縮放：16")
                            if not data.get('地圖圖層'):
                                data['地圖圖層'] = 'EMAP - 通用電子地圖(內政部)'
                                print("[OK] 設定預設圖層：EMAP")

                            # 🔥 自動載入都市計畫分區資料
                            print("正在檢查都市計畫分區資料...")
                            update_message("[資訊] 正在檢查都市計畫分區資料...")
                            try:
                                # 🔥 從 data 中解析地段地號
                                target_lots = parse_land_lots_from_data(data)
                                print(f"[檢查] 從土地標示解析出 {len(target_lots)} 筆地段地號")

                                planning_data = load_urban_planning_data(target_lots=target_lots)
                                if planning_data:
                                    zones_text = '、'.join(planning_data['zones'])
                                    building_coverage_text = '、'.join(planning_data.get('building_coverage_ratios', []))
                                    floor_area_text = '、'.join(planning_data.get('floor_area_ratios', []))

                                    # 🔥 顯示載入的資訊（不管是否已有資料）
                                    if zones_text:
                                        update_message(f"[OK] 都市計畫分區：{zones_text}")
                                    if building_coverage_text:
                                        update_message(f"[OK] 建蔽率：{building_coverage_text}%")
                                    if floor_area_text:
                                        update_message(f"[OK] 容積率：{floor_area_text}%")

                                    # 🔥 只在欄位不存在或為空時才自動填入（但若為「都市計畫外」則不帶入）
                                    if zones_text and zones_text != '都市計畫外' and (not data.get('都市土地使用分區') or str(data.get('都市土地使用分區', '')).strip() == ''):
                                        data['都市土地使用分區'] = zones_text
                                        print(f"[OK] 自動填入都市計畫分區：{zones_text}")
                                    elif zones_text == '都市計畫外':
                                        print(f"[略過] 都市計畫分區為「都市計畫外」，不自動填入")

                                    if building_coverage_text and (not data.get('建蔽率') or str(data.get('建蔽率', '')).strip() == ''):
                                        data['建蔽率'] = building_coverage_text
                                        print(f"[OK] 自動填入建蔽率：{building_coverage_text}")

                                    if floor_area_text and (not data.get('容積率') or str(data.get('容積率', '')).strip() == ''):
                                        data['容積率'] = floor_area_text
                                        print(f"[OK] 自動填入容積率：{floor_area_text}")
                                else:
                                    print("[提示] 未找到符合條件的都市計畫分區資料")
                                    update_message("[提示] 未找到符合條件的都市計畫分區資料")
                            except Exception as plan_e:
                                print(f"[提示] 載入都市計畫分區資料失敗：{plan_e}")
                                update_message(f"[提示] 載入都市計畫分區資料失敗：{plan_e}")
    except Exception as e:
        print(f"從 data.json 讀取座標失敗：{e}")
    
    # ========================================
    # 🔥 繼續建立 GUI 內容
    # ========================================
    # 建立主框架（含滾動條）
    main_frame = tk.Frame(editor_window)
    main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    
    # 滾動畫布
    canvas = tk.Canvas(main_frame)
    scrollbar = tk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
    scrollable_frame = tk.Frame(canvas)
    
    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
    )
    
    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)

    # 🔥 綁定滑鼠滾輪事件
    def _on_mousewheel(event):
        canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        # 🔥 移除 update_idletasks()，避免打斷輸入造成游標消失

    # 綁定到編輯器視窗，讓整個視窗都能滾動
    editor_window.bind("<MouseWheel>", _on_mousewheel)
    canvas.bind("<MouseWheel>", _on_mousewheel)
    scrollable_frame.bind("<MouseWheel>", _on_mousewheel)

    # 當滑鼠進入編輯器視窗時，只在沒有輸入欄位擁有焦點時才設定焦點（讓滾輪有效）
    def _on_enter(event):
        # 檢查當前焦點是否在 Entry 或 Text 欄位上（兩者都需保護，否則打字中游標會消失）
        focused = editor_window.focus_get()
        if focused is None or not isinstance(focused, (tk.Entry, tk.Text)):
            editor_window.focus_set()

    editor_window.bind("<Enter>", _on_enter)

    # 🔥 字型設定（放大）
    label_font = font.Font(family="Microsoft JhengHei", size=12, weight="bold")  # 從10改為12
    entry_font = font.Font(family="Microsoft JhengHei", size=12)  # 從10改為12
    section_font = font.Font(family="Microsoft JhengHei", size=14, weight="bold")  # 從12改為14
    
    # 儲存所有輸入欄位的字典
    entries = {}

    # 🔥 儲存所有區塊標題的列表（用於切換顏色）
    section_titles = []

    # === 建立各區塊的輸入欄位 ===

    def create_section(parent, title):
        """建立區塊標題"""
        frame = tk.Frame(parent, relief=tk.GROOVE, borderwidth=2)
        frame.pack(fill=tk.X, pady=5)  # 🔥 從 pady=10 改為 5，縮小上下間距

        title_label = tk.Label(frame, text=title, font=section_font,
                              bg="#E0E0E0", anchor="w", padx=10)
        title_label.pack(fill=tk.X, pady=3)  # 🔥 從 pady=5 改為 3，縮小標題高度

        # 🔥 將標題加入列表，以便後續改變顏色
        section_titles.append(title_label)

        content_frame = tk.Frame(frame)
        content_frame.pack(fill=tk.X, padx=10, pady=5)

        return content_frame
    
    def create_field(parent, label_text, key, data, width=40):
        """建立單一欄位"""
        row_frame = tk.Frame(parent)
        row_frame.pack(fill=tk.X, pady=2)  # 🔥 從 pady=3 改為 2，減少間距

        label = tk.Label(row_frame, text=label_text, font=label_font, width=30, anchor="e")  # 🔥 從 15 改為 20
        label.pack(side=tk.LEFT, padx=5)

        # 🔥 所有欄位都可編輯
        entry = tk.Entry(row_frame, font=entry_font, width=width)
        entry.pack(side=tk.LEFT, padx=5)

        # 填入現有資料
        value = data.get(key, "")
        if isinstance(value, dict):
            entry.insert(0, str(value))
        else:
            entry.insert(0, str(value))

        # 🔥 強化版焦點管理 - 解決游標消失問題
        def on_entry_focus(event):
            """當欄位獲得焦點時"""
            try:
                widget = event.widget
                # 標記為活動欄位，不強制移游標位置（讓使用者決定游標在哪）
                widget._is_active = True
            except:
                pass

        def on_entry_click(event):
            """當欄位被點擊時：只設焦點，不干預游標位置，讓 tkinter 自行處理點擊定位"""
            try:
                widget = event.widget
                widget.focus_set()
                widget._is_active = True
                # 不 return "break"，讓 tkinter 原生處理游標點擊定位
            except:
                pass

        def on_entry_key(event):
            """當有按鍵輸入時，確保焦點存在"""
            try:
                widget = event.widget
                if not widget.focus_get() == widget:
                    widget.focus_set()
            except:
                pass

        def on_focus_out(event):
            """當失去焦點時"""
            try:
                widget = event.widget
                widget._is_active = False
            except:
                pass

        entry.bind('<FocusIn>', on_entry_focus)
        entry.bind('<Button-1>', on_entry_click)
        entry.bind('<KeyPress>', on_entry_key)
        entry.bind('<FocusOut>', on_focus_out)

        entries[key] = entry
        return entry

    # === 交易類型區（重要）===
    transaction_frame = create_section(scrollable_frame, "🔄 交易類型")

    transaction_row = tk.Frame(transaction_frame)
    transaction_row.pack(fill=tk.X, pady=3)
    tk.Label(transaction_row, text="交易類型：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    transaction_var = tk.StringVar(value=data.get("交易類型", "售件"))

    # 使用按鈕取代單選鈕，配色與買賣/租賃按鈕相同
    def select_sale():
        transaction_var.set("售件")
        sale_btn.config(relief=tk.SUNKEN, bg='#70D070')  # 深綠色表示選中
        rent_btn.config(relief=tk.RAISED, bg='#FFB6C1')  # 恢復粉色

    def select_rent():
        transaction_var.set("租件")
        rent_btn.config(relief=tk.SUNKEN, bg='#FF96A1')  # 深粉色表示選中
        sale_btn.config(relief=tk.RAISED, bg='#90EE90')  # 恢復綠色

    # 創建按鈕
    sale_btn = tk.Button(transaction_row, text="售 件", command=select_sale,
                        font=("Microsoft JhengHei", 14, "bold"),
                        bg='#90EE90', width=10, height=1)
    sale_btn.pack(side=tk.LEFT, padx=10)

    rent_btn = tk.Button(transaction_row, text="租 件", command=select_rent,
                        font=("Microsoft JhengHei", 14, "bold"),
                        bg='#FFB6C1', width=10, height=1)
    rent_btn.pack(side=tk.LEFT, padx=10)

    # 根據初始值設定按鈕狀態
    if data.get("交易類型", "售件") == "售件":
        sale_btn.config(relief=tk.SUNKEN, bg='#70D070')
    else:
        rent_btn.config(relief=tk.SUNKEN, bg='#FF96A1')

    entries["交易類型"] = transaction_var

    # === 基本資訊區 ===
    basic_frame = create_section(scrollable_frame, "📋 基本資訊")
    create_field(basic_frame, "物件編號：", "物件編號", data, width=30)
    create_field(basic_frame, "契約編號：", "契約編號", data, width=30)
    create_field(basic_frame, "案名：", "案名", data, width=50)

    # === 地址資訊區 ===
    address_frame = create_section(scrollable_frame, "📍 地址資訊")
    create_field(address_frame, "土地標示：", "土地標示", data, width=60)
    create_field(address_frame, "建物門牌：", "建物門牌", data, width=60)

    # === 物件類型區 ===
    type_frame = create_section(scrollable_frame, "🏠 物件類型")

    # 型式（下拉選單）
    type_row = tk.Frame(type_frame)
    type_row.pack(fill=tk.X, pady=3)
    tk.Label(type_row, text="型式：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    type_var = tk.StringVar(value=data.get("型式", ""))

    import tkinter.ttk as ttk
    # 🔥 設定下拉選單字體（加大）
    combo_font = font.Font(family="Microsoft JhengHei", size=16)  # 比 entry_font 更大

    # 🔥 設定 ttk.Combobox 的下拉清單樣式（讓下拉選項也變大）
    style = ttk.Style()
    # 設定 Combobox 的下拉清單字體大小
    root.option_add('*TCombobox*Listbox.font', combo_font)
    root.option_add('*TCombobox*Listbox*Font', combo_font)

    type_combo = ttk.Combobox(type_row, textvariable=type_var, font=combo_font, width=15,
                                values=["土地", "透天", "大樓", "華廈", "公寓"])
    type_combo.pack(side=tk.LEFT, padx=5)
    entries["型式"] = type_var

    # 用途（下拉選單）
    use_row = tk.Frame(type_frame)
    use_row.pack(fill=tk.X, pady=3)
    tk.Label(use_row, text="用途：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    use_var = tk.StringVar(value=data.get("用途", ""))
    use_combo = ttk.Combobox(use_row, textvariable=use_var, font=combo_font, width=15,
                               values=["建地", "工業地", "廠房", "店面", "辦公", "住家", "農地", "其他"])
    use_combo.pack(side=tk.LEFT, padx=5)
    entries["用途"] = use_var

    # === 價格資訊區（重要）===
    price_frame = create_section(scrollable_frame, "💰 價格資訊（必填）")

    # 總價欄位（含建議區域的容器）
    total_price_container = tk.Frame(price_frame)
    total_price_container.pack(fill=tk.X, pady=3)

    total_price_row = tk.Frame(total_price_container)
    total_price_row.pack(fill=tk.X)
    tk.Label(total_price_row, text="總價（萬元）：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    total_price_entry = bind_entry_focus(tk.Entry(total_price_row, font=entry_font, width=20))  # 🔥 綁定焦點事件
    total_price_entry.pack(side=tk.LEFT, padx=5)
    total_price_entry.insert(0, str(data.get("總價", "")))
    entries["總價"] = total_price_entry

    # 🔥 中文金額顯示標籤
    total_price_chinese_label = tk.Label(total_price_row, text="",
                                         font=("Microsoft JhengHei", 10, "bold"),
                                         fg='#D32F2F', anchor='w')  # 紅色，醒目
    total_price_chinese_label.pack(side=tk.LEFT, padx=10)

    # 🔥 金額轉換函數
    def format_price_chinese(price_wan):
        """
        將萬元金額轉換為中文顯示格式（完整萬進制）

        單位系統（萬進制）：
        - 載 = 10^44 萬 = 1,0000,0000,0000,0000,0000,0000,0000,0000,0000,0000,0000 萬
        - 正 = 10^40 萬
        - 澗 = 10^36 萬
        - 溝 = 10^32 萬
        - 穰 = 10^28 萬
        - 秭 = 10^24 萬
        - 垓 = 10^20 萬
        - 京 = 10^16 萬
        - 兆 = 10^12 萬（註：台灣常用，10^8萬是簡化版）
        - 億 = 10^4 萬
        - 萬

        範例：
        - 153589 → 15億3589萬
        - 123456789012 → 12兆3456億7890萬（完整版）
        - 1234.5 → 1234.5萬（保留小數點後一位）
        """
        try:
            price_float = float(price_wan)

            if price_float == 0:
                return ""

            # 判斷是否有小數
            has_decimal = (price_float != int(price_float))

            # 定義各單位（從大到小）
            # 注意：台灣實務上「兆」常用 10^8 萬（1萬億）
            # 採用台灣實務慣例

            units = [
                ("載", 10**44),
                ("正", 10**40),
                ("澗", 10**36),
                ("溝", 10**32),
                ("穰", 10**28),
                ("秭", 10**24),
                ("垓", 10**20),
                ("京", 10**16),
                ("兆", 10**8),   # 台灣常用：1兆 = 1萬億 = 10^8萬
                ("億", 10**4),
                ("萬", 1)
            ]

            result = []
            remainder = price_float

            for i, (unit_name, unit_value) in enumerate(units):
                if remainder >= unit_value:
                    count = remainder / unit_value

                    # 最後一個單位保留小數（如果有的話）
                    if i == len(units) - 1 or remainder < units[i+1][1] if i+1 < len(units) else True:
                        if has_decimal and count < 10000:  # 小於下一個單位才顯示小數
                            result.append(f"{count:.1f}{unit_name}")
                            break
                        else:
                            count_int = int(count)
                            result.append(f"{count_int}{unit_name}")
                            break
                    else:
                        count_int = int(count)
                        remainder = remainder - (count_int * unit_value)
                        if count_int > 0:
                            result.append(f"{count_int}{unit_name}")

            return "".join(result) if result else ""

        except (ValueError, TypeError):
            return ""

    # 🔥 更新中文金額顯示
    def update_chinese_price(event=None):
        # 🔥 保存當前焦點
        focused_widget = editor_window.focus_get()

        price_str = total_price_entry.get().strip()
        if price_str:
            chinese_text = format_price_chinese(price_str)
            if chinese_text:
                total_price_chinese_label.config(text=f"= {chinese_text}")
            else:
                total_price_chinese_label.config(text="")
        else:
            total_price_chinese_label.config(text="")

        # 🔥 恢復焦點
        if focused_widget:
            focused_widget.focus_set()

    # 稍後會與 calculate_from_total_price 一起綁定

    # 🔥 建議總價顯示區域（初始隱藏）
    suggestion_frame = tk.Frame(total_price_container)
    suggestion_frame.pack(fill=tk.X, padx=(35+5, 5))  # 對齊輸入框

    suggestion_label = tk.Label(suggestion_frame, text="",
                                font=("Microsoft JhengHei", 9),
                                fg='#666666', anchor='w')
    suggestion_label.pack(side=tk.LEFT, padx=5)

    # 快速採用按鈕組
    adopt_frame = tk.Frame(suggestion_frame)
    adopt_frame.pack(side=tk.LEFT, padx=5)

    # === 售件專屬欄位 ===
    sale_fields_frame = tk.Frame(price_frame)
    sale_fields_frame.pack(fill=tk.X)

    create_field(sale_fields_frame, "含車位價（萬元）：", "含車位價格", data, width=20)

    # 🔥 土地單價（萬元/坪）- 可編輯，淺綠色背景
    land_unit_row = tk.Frame(sale_fields_frame)
    land_unit_row.pack(fill=tk.X, pady=3)
    tk.Label(land_unit_row, text="土地單價（萬元/坪）：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    land_unit_entry = tk.Entry(land_unit_row, font=entry_font, width=20,
                                bg='#E8F5E9', fg='#2E7D32')  # 淺綠色背景，深綠色文字
    land_unit_entry.pack(side=tk.LEFT, padx=5)
    land_unit_entry.insert(0, str(data.get("土地單價", "")))
    entries["土地單價"] = land_unit_entry
    tk.Label(land_unit_row, text="💡可輸入", font=("Microsoft JhengHei", 8), fg='#888888').pack(side=tk.LEFT)

    # ⭐ 顯示開關
    land_unit_display_var = tk.BooleanVar(value=data.get("土地單價_顯示", True))
    land_unit_display_cb = tk.Checkbutton(land_unit_row, text="顯示", variable=land_unit_display_var, font=label_font)
    land_unit_display_cb.pack(side=tk.LEFT, padx=10)
    entries["土地單價_顯示"] = land_unit_display_var

    # 🔥 建物單價（萬元/坪）- 可編輯，淺綠色背景
    building_unit_row = tk.Frame(sale_fields_frame)
    building_unit_row.pack(fill=tk.X, pady=3)
    tk.Label(building_unit_row, text="建物單價（萬元/坪）：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    building_unit_entry = tk.Entry(building_unit_row, font=entry_font, width=20,
                                    bg='#E8F5E9', fg='#2E7D32')  # 淺綠色背景，深綠色文字
    building_unit_entry.pack(side=tk.LEFT, padx=5)
    building_unit_entry.insert(0, str(data.get("建物單價", "")))
    entries["建物單價"] = building_unit_entry
    tk.Label(building_unit_row, text="💡可輸入", font=("Microsoft JhengHei", 8), fg='#888888').pack(side=tk.LEFT)

    # ⭐ 顯示開關
    building_unit_display_var = tk.BooleanVar(value=data.get("建物單價_顯示", True))
    building_unit_display_cb = tk.Checkbutton(building_unit_row, text="顯示", variable=building_unit_display_var, font=label_font)
    building_unit_display_cb.pack(side=tk.LEFT, padx=10)
    entries["建物單價_顯示"] = building_unit_display_var

    # === 租件專屬欄位 ===
    rent_fields_frame = tk.Frame(price_frame)
    rent_fields_frame.pack(fill=tk.X)

    # 租金欄位（帶含稅勾選框）
    rent_row = tk.Frame(rent_fields_frame)
    rent_row.pack(fill=tk.X, pady=3)
    tk.Label(rent_row, text="租金（萬元/月）：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    rent_entry = tk.Entry(rent_row, font=entry_font, width=20)
    rent_entry.pack(side=tk.LEFT, padx=5)
    rent_entry.insert(0, data.get("租金", ""))
    entries["租金"] = rent_entry

    # 含稅勾選框
    rent_tax_included_var = tk.BooleanVar(value=data.get("租金_含稅", True))  # 預設勾選
    rent_tax_cb = tk.Checkbutton(rent_row, text="含稅", variable=rent_tax_included_var, font=label_font)
    rent_tax_cb.pack(side=tk.LEFT, padx=10)
    entries["租金_含稅"] = rent_tax_included_var

    # 含稅金額顯示標籤（亮紅色，字體較大）
    rent_tax_amount_label = tk.Label(rent_row, text="", font=("Microsoft JhengHei", 13, "bold"), fg='#FF3333')
    rent_tax_amount_label.pack(side=tk.LEFT, padx=10)

    # 更新含稅金額顯示的函數
    def update_rent_tax_display(*args):
        try:
            rent_value = rent_entry.get().strip()
            is_tax_included = rent_tax_included_var.get()

            if rent_value and not is_tax_included:
                # 未勾選含稅時，檢查是否需要顯示含稅金額
                rent_value_yuan = float(rent_value) * 10000

                # 只有當租金 >= 20,000 元時才顯示含稅金額
                if rent_value_yuan >= 20000:
                    tax_included_yuan = calculate_rent_with_tax(rent_value)
                    if tax_included_yuan:
                        # 直接使用返回的元數（精確數字）
                        rent_tax_amount_label.config(text=f"含稅 {tax_included_yuan:,} 元")
                    else:
                        rent_tax_amount_label.config(text="")
                else:
                    # 低於 19,999 元，不顯示
                    rent_tax_amount_label.config(text="")
            else:
                # 勾選含稅或未輸入租金時，不顯示
                rent_tax_amount_label.config(text="")
        except:
            rent_tax_amount_label.config(text="")

    # 綁定事件：租金輸入變化時更新
    rent_entry.bind('<KeyRelease>', update_rent_tax_display)
    # 綁定事件：含稅勾選框變化時更新
    rent_tax_included_var.trace('w', update_rent_tax_display)
    # 初始化顯示
    update_rent_tax_display()

    # 押金欄位（下拉選單）
    deposit_row = tk.Frame(rent_fields_frame)
    deposit_row.pack(fill=tk.X, pady=3)
    tk.Label(deposit_row, text="押金：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    deposit_var = tk.StringVar(value=data.get("押金", "二個月"))
    deposit_combo = ttk.Combobox(deposit_row, textvariable=deposit_var, font=combo_font, width=20,
                                  values=["一個月", "二個月", "三個月", "月", "元"])
    deposit_combo.pack(side=tk.LEFT, padx=5)
    entries["押金"] = deposit_var

    create_field(rent_fields_frame, "租期（年）：", "租期", data, width=20)
    create_field(rent_fields_frame, "整地期：", "整地期", data, width=20)
    create_field(rent_fields_frame, "調幅：", "調幅", data, width=20)
    # 🔥 管理費移到 facility_frame（其他設施區），讓售件和租件都能使用
    
    # === 面積資訊區 ===
    area_frame = create_section(scrollable_frame, "📐 面積資訊")
    land_area_entry = create_field(area_frame, "總地坪：", "總地坪", data, width=20)

    main_building_entry = create_field(area_frame, "主建物坪數：", "主建物坪數", data, width=20)

    # 🔥 主建物各層明細（最多9層）
    floor_label_row = tk.Frame(area_frame)
    floor_label_row.pack(fill=tk.X, pady=(10, 3))
    tk.Label(floor_label_row, text="主建物各層明細：", font=("Microsoft JhengHei", 11, "bold"),
             width=30, anchor="e", fg='#1976D2').pack(side=tk.LEFT, padx=5)
    tk.Label(floor_label_row, text="（自動加總至主建物坪數）",
             font=("Microsoft JhengHei", 9), fg='#666666').pack(side=tk.LEFT)

    floor_entries = []
    主建物資料 = data.get("主建物", {})

    for i in range(9):
        floor_row = tk.Frame(area_frame)
        floor_row.pack(fill=tk.X, pady=2)

        # 層次名稱
        tk.Label(floor_row, text=f"  層次 {i+1}：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
        floor_name_entry = tk.Entry(floor_row, font=entry_font, width=15, bg='#FFF9C4')  # 淺黃色
        floor_name_entry.pack(side=tk.LEFT, padx=2)

        tk.Label(floor_row, text="坪數：", font=label_font).pack(side=tk.LEFT, padx=(10, 2))
        floor_area_entry = tk.Entry(floor_row, font=entry_font, width=10, bg='#FFF9C4')
        floor_area_entry.pack(side=tk.LEFT, padx=2)
        tk.Label(floor_row, text="坪", font=label_font, fg='#666666').pack(side=tk.LEFT)

        # 從 data 讀取現有資料
        if isinstance(主建物資料, dict):
            # 嘗試找到對應的樓層資料
            floor_data_found = False
            for floor_name, floor_area in 主建物資料.items():
                if not floor_data_found:  # 只取第 i 個
                    floor_name_entry.insert(0, str(floor_name))
                    floor_area_entry.insert(0, str(floor_area))
                    floor_data_found = True
                    主建物資料.pop(floor_name)  # 移除已使用的資料
                    break

        floor_entries.append((floor_name_entry, floor_area_entry))
        entries[f"主建物_樓層{i+1}_名稱"] = floor_name_entry
        entries[f"主建物_樓層{i+1}_坪數"] = floor_area_entry

    附屬建物_entry = create_field(area_frame, "附屬建物坪數：", "附屬建物坪數", data, width=20)

    # 🔥 附屬建物明細（最多9項）
    attached_label_row = tk.Frame(area_frame)
    attached_label_row.pack(fill=tk.X, pady=(10, 3))
    tk.Label(attached_label_row, text="附屬建物明細：", font=("Microsoft JhengHei", 11, "bold"),
             width=30, anchor="e", fg='#1976D2').pack(side=tk.LEFT, padx=5)
    tk.Label(attached_label_row, text="（自動加總至附屬建物坪數）",
             font=("Microsoft JhengHei", 9), fg='#666666').pack(side=tk.LEFT)

    attached_entries = []
    附屬建物資料 = data.get("附屬建物", {})

    for i in range(9):
        attached_row = tk.Frame(area_frame)
        attached_row.pack(fill=tk.X, pady=2)

        # 附屬建物類型
        tk.Label(attached_row, text=f"  項目 {i+1}：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
        attached_type_entry = tk.Entry(attached_row, font=entry_font, width=15, bg='#E1F5FE')  # 淺藍色
        attached_type_entry.pack(side=tk.LEFT, padx=2)

        tk.Label(attached_row, text="坪數：", font=label_font).pack(side=tk.LEFT, padx=(10, 2))
        attached_area_entry = tk.Entry(attached_row, font=entry_font, width=10, bg='#E1F5FE')
        attached_area_entry.pack(side=tk.LEFT, padx=2)
        tk.Label(attached_row, text="坪", font=label_font, fg='#666666').pack(side=tk.LEFT)

        # 從 data 讀取現有資料
        if isinstance(附屬建物資料, dict):
            # 嘗試找到對應的附屬建物資料
            attached_data_found = False
            for attached_type, attached_area in list(附屬建物資料.items()):
                if not attached_data_found:  # 只取第 i 個
                    attached_type_entry.insert(0, str(attached_type))
                    attached_area_entry.insert(0, str(attached_area))
                    attached_data_found = True
                    附屬建物資料.pop(attached_type)  # 移除已使用的資料
                    break

        attached_entries.append((attached_type_entry, attached_area_entry))
        entries[f"附屬建物_項目{i+1}_類型"] = attached_type_entry
        entries[f"附屬建物_項目{i+1}_坪數"] = attached_area_entry

    公共設施_entry = create_field(area_frame, "公共設施坪數：", "公共設施坪數", data, width=20)
    增建面積_entry = create_field(area_frame, "增建面積：", "增建面積", data, width=20)
    total_building_entry = create_field(area_frame, "總建坪：", "總建坪", data, width=20)

    # 🔥 自動計算總建坪（各項加總）
    def calculate_total_building_area(event=None):
        """自動加總：主建物各層 + 附屬建物 + 公共設施 + 增建面積 = 總建坪"""
        # 🔥 保存當前焦點
        focused_widget = editor_window.focus_get()

        try:
            total = 0.0

            # 加總主建物各層
            main_building_total = 0.0
            for floor_name_entry, floor_area_entry in floor_entries:
                area_str = floor_area_entry.get().strip()
                if area_str:
                    try:
                        main_building_total += float(area_str)
                    except ValueError:
                        pass

            # 更新主建物坪數
            main_building_entry.delete(0, tk.END)
            if main_building_total > 0:
                main_building_entry.insert(0, f"{main_building_total:.2f}")
            total += main_building_total

            # 🔥 加總附屬建物各項
            attached_total = 0.0
            for attached_type_entry, attached_area_entry in attached_entries:
                area_str = attached_area_entry.get().strip()
                if area_str:
                    try:
                        attached_total += float(area_str)
                    except ValueError:
                        pass

            # 🔥 更新附屬建物坪數
            附屬建物_entry.delete(0, tk.END)
            if attached_total > 0:
                附屬建物_entry.insert(0, f"{attached_total:.2f}")
            total += attached_total

            # 加上公共設施
            公設 = 公共設施_entry.get().strip()
            if 公設:
                try:
                    total += float(公設)
                except ValueError:
                    pass

            # 加上增建面積
            增建 = 增建面積_entry.get().strip()
            if 增建:
                try:
                    total += float(增建)
                except ValueError:
                    pass

            # 更新總建坪
            total_building_entry.delete(0, tk.END)
            if total > 0:
                total_building_entry.insert(0, f"{total:.2f}")

        except Exception:
            pass

        # 🔥 恢復焦點
        if focused_widget:
            focused_widget.focus_set()

    # 🔥 自動計算函數（雙向計算）
    calculating = [False]  # 防止循環觸發的旗標

    def calculate_from_total_price(event=None):
        """從總價計算單價"""
        if calculating[0]:
            return
        calculating[0] = True

        # 🔥 保存當前焦點
        focused_widget = editor_window.focus_get()

        try:
            total_price_str = total_price_entry.get().strip()

            # 🔥 如果總價為空，清空單價和建議
            if not total_price_str:
                land_unit_entry.delete(0, tk.END)
                building_unit_entry.delete(0, tk.END)
                suggestion_label.config(text="")
                # 清空按鈕
                for widget in adopt_frame.winfo_children():
                    widget.destroy()
                calculating[0] = False
                # 🔥 恢復焦點
                if focused_widget:
                    focused_widget.focus_set()
                return

            total_price = float(total_price_str)

            # 計算土地單價
            land_area_str = land_area_entry.get().strip()
            if land_area_str:
                try:
                    land_area = float(land_area_str)
                    if land_area > 0:
                        land_unit_price = total_price / land_area
                        land_unit_entry.delete(0, tk.END)
                        land_unit_entry.insert(0, f"{land_unit_price:.2f}")
                except ValueError:
                    pass

            # 計算建物單價（需考慮車位價格和坪數）
            building_area_str = total_building_entry.get().strip()
            if not building_area_str:
                building_area_str = main_building_entry.get().strip()

            if building_area_str:
                try:
                    building_area = float(building_area_str)
                    if building_area > 0:
                        # 🔥 檢查是否有含車位價格
                        price_for_building = total_price  # 預設使用總價
                        area_for_building = building_area  # 預設使用總建坪

                        # 檢查含車位價（從 entries 中取得）
                        parking_price_entry = entries.get("含車位價格")
                        if parking_price_entry:
                            parking_price_str = parking_price_entry.get().strip()
                            if parking_price_str:
                                try:
                                    parking_price = float(parking_price_str)
                                    # 🔥 從總價扣除車位價
                                    price_for_building = total_price - parking_price

                                    # 🔥 檢查是否有車位坪數（從公共設施中找停車位相關）
                                    # 如果公共設施中有車位資料，也要扣除車位坪數
                                    # 這部分較複雜，暫時不處理坪數扣除
                                    # TODO: 未來可以加入車位坪數的判斷邏輯

                                except ValueError:
                                    pass

                        if area_for_building > 0 and price_for_building > 0:
                            building_unit_price = price_for_building / area_for_building
                            building_unit_entry.delete(0, tk.END)
                            building_unit_entry.insert(0, f"{building_unit_price:.2f}")
                except ValueError:
                    pass

            # 清空建議（因為是從總價算出來的）
            suggestion_label.config(text="")
            for widget in adopt_frame.winfo_children():
                widget.destroy()

        except ValueError:
            # 總價無效，清空單價
            land_unit_entry.delete(0, tk.END)
            building_unit_entry.delete(0, tk.END)
            suggestion_label.config(text="")
            for widget in adopt_frame.winfo_children():
                widget.destroy()

        calculating[0] = False

        # 🔥 恢復焦點
        if focused_widget:
            focused_widget.focus_set()

    def calculate_from_unit_prices(event=None):
        """從單價計算建議總價"""
        if calculating[0]:
            return
        calculating[0] = True

        # 🔥 保存當前焦點
        focused_widget = editor_window.focus_get()

        try:
            # 清空舊的按鈕
            for widget in adopt_frame.winfo_children():
                widget.destroy()

            suggestions = []

            # 從土地單價推算
            land_unit_str = land_unit_entry.get().strip()
            land_area_str = land_area_entry.get().strip()
            if land_unit_str and land_area_str:
                try:
                    land_unit = float(land_unit_str)
                    land_area = float(land_area_str)
                    if land_unit > 0 and land_area > 0:
                        suggested_price = land_unit * land_area
                        suggestions.append(("土地", suggested_price))
                except ValueError:
                    pass

            # 從建物單價推算
            building_unit_str = building_unit_entry.get().strip()
            building_area_str = total_building_entry.get().strip()
            if not building_area_str:
                building_area_str = main_building_entry.get().strip()

            if building_unit_str and building_area_str:
                try:
                    building_unit = float(building_unit_str)
                    building_area = float(building_area_str)
                    if building_unit > 0 and building_area > 0:
                        suggested_price = building_unit * building_area
                        suggestions.append(("建物", suggested_price))
                except ValueError:
                    pass

            # 顯示建議
            if suggestions:
                suggestion_texts = []
                for source, price in suggestions:
                    suggestion_texts.append(f"{source}: {price:.2f} 萬")

                suggestion_label.config(text="💡 建議總價：" + " | ".join(suggestion_texts))

                # 🔥 建立快速採用按鈕
                def create_adopt_button(value, label):
                    btn = tk.Button(adopt_frame, text=label,
                                   font=("Microsoft JhengHei", 8),
                                   bg='#4CAF50', fg='white',
                                   padx=5, pady=2,
                                   command=lambda: adopt_price(value))
                    btn.pack(side=tk.LEFT, padx=2)

                # 使用第一個建議值（如果有土地單價就用土地，否則用建物）
                primary_price = suggestions[0][1]

                # 原始值
                create_adopt_button(primary_price, f"採用 {primary_price:.2f}")

                # 整數
                if primary_price != int(primary_price):
                    create_adopt_button(round(primary_price), f"整數 {round(primary_price)}")

                # 無條件進位到十位
                rounded_up = int((primary_price + 9) // 10 * 10)
                if rounded_up != primary_price:
                    create_adopt_button(rounded_up, f"進位 {rounded_up}")

            else:
                suggestion_label.config(text="")

        except Exception as e:
            suggestion_label.config(text="")

        calculating[0] = False

        # 🔥 恢復焦點
        if focused_widget:
            focused_widget.focus_set()

    def adopt_price(price):
        """採用建議價格"""
        total_price_entry.delete(0, tk.END)
        total_price_entry.insert(0, str(int(price) if price == int(price) else price))
        update_chinese_price()  # 更新中文顯示
        calculate_from_total_price()  # 重新計算單價

    # 🔥 綁定事件
    # 總價變更 → 更新中文顯示 + 計算單價
    def on_total_price_change(event=None):
        update_chinese_price(event)
        calculate_from_total_price(event)

    total_price_entry.bind('<KeyRelease>', on_total_price_change)
    land_area_entry.bind('<KeyRelease>', calculate_from_total_price)
    total_building_entry.bind('<KeyRelease>', calculate_from_total_price)

    # 🔥 含車位價變更 → 重新計算建物單價
    parking_price_entry = entries.get("含車位價格")
    if parking_price_entry:
        parking_price_entry.bind('<KeyRelease>', calculate_from_total_price)

    # 單價變更 → 計算建議總價
    land_unit_entry.bind('<KeyRelease>', calculate_from_unit_prices)
    building_unit_entry.bind('<KeyRelease>', calculate_from_unit_prices)

    # 🔥 主建物各層、附屬建物各項、公共設施、增建面積變更 → 自動加總總建坪 → 重新計算單價
    def on_area_change(event=None):
        calculate_total_building_area()  # 先加總
        calculate_from_total_price()  # 再計算單價

    # 綁定主建物各層
    for floor_name_entry, floor_area_entry in floor_entries:
        floor_area_entry.bind('<KeyRelease>', on_area_change)

    # 🔥 綁定附屬建物各項
    for attached_type_entry, attached_area_entry in attached_entries:
        attached_area_entry.bind('<KeyRelease>', on_area_change)

    公共設施_entry.bind('<KeyRelease>', on_area_change)
    增建面積_entry.bind('<KeyRelease>', on_area_change)

    # 🔥 初始計算
    calculate_total_building_area()  # 先計算總建坪
    update_chinese_price()  # 初始顯示中文金額
    if data.get("總價", ""):
        calculate_from_total_price()
    elif data.get("土地單價", "") or data.get("建物單價", ""):
        calculate_from_unit_prices()

    # === 交易類型切換功能（需要在車位相關變數定義後才能完整定義）===
    # 🔥 先不定義函數，等到所有相關變數都定義完成後再定義

    # === 土地所有權型態區 ===
    ownership_frame = create_section(scrollable_frame, "📋 土地所有權型態")

    ownership_row = tk.Frame(ownership_frame)
    ownership_row.pack(fill=tk.X, pady=3)
    tk.Label(ownership_row, text="土地所有權型態：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)

    # 取得現有的所有權型態
    現有型態 = data.get('所有權型態', '')

    # 單獨所有勾選框
    單獨所有_var = tk.BooleanVar(value='單獨所有' in 現有型態)
    單獨所有_check = tk.Checkbutton(ownership_row, text="單獨所有", variable=單獨所有_var, font=entry_font)
    單獨所有_check.pack(side=tk.LEFT, padx=5)
    entries["所有權型態_單獨所有"] = 單獨所有_var

    # 分別共有勾選框
    分別共有_var = tk.BooleanVar(value='分別共有' in 現有型態)
    分別共有_check = tk.Checkbutton(ownership_row, text="分別共有", variable=分別共有_var, font=entry_font)
    分別共有_check.pack(side=tk.LEFT, padx=5)
    entries["所有權型態_分別共有"] = 分別共有_var

    # 公同共有勾選框
    公同共有_var = tk.BooleanVar(value='公同共有' in 現有型態)
    公同共有_check = tk.Checkbutton(ownership_row, text="公同共有", variable=公同共有_var, font=entry_font)
    公同共有_check.pack(side=tk.LEFT, padx=5)
    entries["所有權型態_公同共有"] = 公同共有_var

    # 🔥 新增：地上權勾選框
    地上權_var = tk.BooleanVar(value='地上權' in 現有型態)
    地上權_check = tk.Checkbutton(ownership_row, text="地上權", variable=地上權_var, font=entry_font)
    地上權_check.pack(side=tk.LEFT, padx=5)
    entries["所有權型態_地上權"] = 地上權_var

    # === 🔥 自動讀取土地增值稅 ===
    auto_load_tax_data = {}
    all_found_files = []  # 記錄所有找到的檔案
    expected_path = None  # 預期的檔案路徑

    try:
        # 🔥 從 data.json 讀取地段資訊，找到對應的資料夾
        tax_file_found = None

        if os.path.exists(get_data_json_path()):
            with open(get_data_json_path(), 'r', encoding='utf-8') as f:
                data_list = json.load(f)

            if data_list and len(data_list) > 0:
                first_data = data_list[0]

                # 取得地段資訊
                area = first_data.get('area', '')  # 例如：湖內區
                section = first_data.get('section', '')  # 例如：大湖段
                lot_number = first_data.get('lot_number', '')  # 例如：2589-6

                if area and section and lot_number:
                    # 🔥 推算資料夾名稱（與 land.py 的邏輯完全一致）
                    folder_name = get_work_folder(f"{area}{section}-{lot_number}")
                    other_folder = os.path.join(folder_name, "4.其他相關")

                    # 🔥 搜尋所有符合模式的稅額檔案（財政部_tax_result.json 或 高雄市_tax_result.json）
                    import glob

                    # 先嘗試搜尋該資料夾下的所有稅額檔案
                    if os.path.exists(other_folder):
                        all_tax_files = glob.glob(os.path.join(other_folder, "*tax_result.json"))

                        if all_tax_files:
                            all_found_files = all_tax_files
                            if len(all_tax_files) == 1:
                                tax_file_found = all_tax_files[0]
                        else:
                            # 未找到稅額檔案，靜默處理
                            pass
                    else:
                        # 🔥 資料夾不存在，跳過自動載入（避免全域搜尋造成延遲）
                        pass

        # 🔥 定義讀取稅額的函數
        def load_tax_data_from_file(file_path):
            """從指定檔案讀取稅額資料"""
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    tax_data = json.load(f)

                calc_date = tax_data.get('計算日期', '')
                自用稅額 = tax_data.get('自用住宅稅額', 0)
                一般稅額 = tax_data.get('一般土地稅額', 0)

                print(f"[OK] 成功讀取稅額：")
                print(f"  檔案：{file_path}")
                print(f"  計算日期：{calc_date}")
                print(f"  自用稅額：{自用稅額:,} 元")
                print(f"  一般稅額：{一般稅額:,} 元")

                # 將稅額資料寫入 data 字典（包括 0 也要顯示）
                if 自用稅額 is not None:
                    data['自用增值稅'] = f"{自用稅額:,}"
                if 一般稅額 is not None:
                    data['一般增值稅'] = f"{一般稅額:,}"

                update_message(f"[OK] 已載入土地增值稅資料（{calc_date}）")
                update_message(f"     來源：{file_path}")
                update_message(f"     自用增值稅：{自用稅額:,} 元")
                update_message(f"     一般增值稅：{一般稅額:,} 元")

                # 更新編輯器中的欄位顯示
                if '自用增值稅' in entries:
                    entries['自用增值稅'].delete(0, tk.END)
                    entries['自用增值稅'].insert(0, data.get('自用增值稅', ''))
                if '一般增值稅' in entries:
                    entries['一般增值稅'].delete(0, tk.END)
                    entries['一般增值稅'].insert(0, data.get('一般增值稅', ''))

                return True
            except Exception as e:
                update_message(f"[錯誤] 讀取檔案失敗：{e}")
                return False

        # 🔥 定義比較兩個稅額檔案的函數
        def load_and_compare_tax_files(finance_file, kaohsiung_file):
            """載入並比較財政部與高雄市的稅額資料"""
            try:
                # 讀取財政部資料
                with open(finance_file, 'r', encoding='utf-8') as f:
                    finance_data = json.load(f)

                finance_calc_date = finance_data.get('計算日期', '')
                finance_自用 = finance_data.get('自用住宅稅額', 0)
                finance_一般 = finance_data.get('一般土地稅額', 0)

                # 讀取高雄市資料
                with open(kaohsiung_file, 'r', encoding='utf-8') as f:
                    kaohsiung_data = json.load(f)

                kaohsiung_calc_date = kaohsiung_data.get('計算日期', '')
                kaohsiung_自用 = kaohsiung_data.get('自用住宅稅額', 0)
                kaohsiung_一般 = kaohsiung_data.get('一般土地稅額', 0)

                # 🔥 優先使用財政部版本
                data['自用增值稅'] = f"{finance_自用:,}"
                data['一般增值稅'] = f"{finance_一般:,}"

                # 更新編輯器中的欄位顯示
                if '自用增值稅' in entries:
                    entries['自用增值稅'].delete(0, tk.END)
                    entries['自用增值稅'].insert(0, data.get('自用增值稅', ''))
                if '一般增值稅' in entries:
                    entries['一般增值稅'].delete(0, tk.END)
                    entries['一般增值稅'].insert(0, data.get('一般增值稅', ''))

                # 🔥 顯示比較結果
                update_message(f"[OK] 已載入土地增值稅資料（使用財政部版本）")
                update_message(f"")
                update_message(f"  財政部版本：")
                update_message(f"    自用住宅稅額：{finance_自用:>12,} 元")
                update_message(f"    一般土地稅額：{finance_一般:>12,} 元")
                update_message(f"")
                update_message(f"  高雄市版本：")
                update_message(f"    自用住宅稅額：{kaohsiung_自用:>12,} 元")
                update_message(f"    一般土地稅額：{kaohsiung_一般:>12,} 元")

                # 🔥 計算差異
                自用差異 = abs(finance_自用 - kaohsiung_自用)
                一般差異 = abs(finance_一般 - kaohsiung_一般)

                if 自用差異 > 0 or 一般差異 > 0:
                    update_message(f"")
                    update_message(f"  ⚠️ 稅額差異：")
                    if 自用差異 > 0:
                        update_message(f"     自用住宅稅額差異：{自用差異:,} 元")
                    if 一般差異 > 0:
                        update_message(f"     一般土地稅額差異：{一般差異:,} 元")
                else:
                    update_message(f"")
                    update_message(f"  ✓ 兩個版本的稅額完全一致")

                update_message(f"")

                # 🔥 檢查是否需要20年提示
                check_and_show_20_year_warning(finance_data)

                return True

            except Exception as e:
                update_message(f"[錯誤] 比較稅額檔案失敗：{e}")
                return False

        # 🔥 定義檢查20年提示的函數
        def check_and_show_20_year_warning(tax_data):
            """檢查前次移轉日期是否超過20年，並顯示提示"""
            try:
                from datetime import datetime

                # 從 tax_data 的詳細資料中取得原地價年月
                原地價年月 = tax_data.get('原地價年月', '')

                if not 原地價年月:
                    return

                # 解析日期格式：例如 "94年1月" 或 "094年01月"
                import re
                match = re.search(r'(\d+)年(\d+)月', 原地價年月)
                if not match:
                    return

                year_roc = int(match.group(1))
                month = int(match.group(2))

                # 轉換為西元年
                year_ad = year_roc + 1911

                # 建立日期物件（使用每月1日）
                transfer_date = datetime(year_ad, month, 1)
                current_date = datetime.now()

                # 計算年份差異
                years_diff = (current_date - transfer_date).days / 365.25

                if years_diff > 20:
                    # 🔥 顯示亮黃色提示
                    update_message(f"")
                    update_message(f"  🟡 ⚠️ 前次移轉日期距今超過20年，已達 {int(years_diff)} 年（原地價年月：{原地價年月}）")
                    update_message(f"  🟡 ⚠️ 建議以財政部估算之稅額為主")
                    update_message(f"")

            except Exception as e:
                print(f"[Debug] 檢查20年提示失敗：{e}")

        # 🔥 定義手動選擇檔案的函數
        def manual_select_tax_file():
            """手動選擇 tax_result.json 檔案"""
            from tkinter import filedialog

            # 🔥 計算預設開啟路徑（基於 data.json 第一筆資料）
            default_dir = "."
            if os.path.exists(get_data_json_path()):
                try:
                    with open(get_data_json_path(), 'r', encoding='utf-8') as f:
                        data_list = json.load(f)
                    if data_list and len(data_list) > 0:
                        first_data = data_list[0]
                        area = first_data.get('area', '')
                        section = first_data.get('section', '')
                        lot_number = first_data.get('lot_number', '')

                        if area and section and lot_number:
                            # 🔥 使用 get_work_folder 確保在主程式目錄下建立資料夾
                            folder_name = get_work_folder(f"{area}{section}-{lot_number}")
                            other_folder = os.path.join(folder_name, "4.其他相關")
                            if os.path.exists(other_folder):
                                default_dir = other_folder
                                print(f"[資訊] 預設開啟路徑：{default_dir}")
                except Exception as e:
                    print(f"[警告] 無法取得預設路徑：{e}")

            file_path = filedialog.askopenfilename(
                title="選擇增值稅結果檔案 (tax_result.json)",
                filetypes=[("JSON files", "tax_result.json"), ("All files", "*.*")],
                initialdir=default_dir
            )
            if file_path:
                load_tax_data_from_file(file_path)

        # 🔥 讀取稅額資料
        if tax_file_found and os.path.exists(tax_file_found):
            load_tax_data_from_file(tax_file_found)
        elif len(all_found_files) > 1:
            # 找到多個檔案，優先選擇財政部版本
            finance_file = None
            kaohsiung_file = None

            for file in all_found_files:
                filename = os.path.basename(file)
                if '財政部' in filename:
                    finance_file = file
                elif '高雄市' in filename:
                    kaohsiung_file = file

            if finance_file and kaohsiung_file:
                # 🔥 兩個版本都有，進行比較
                load_and_compare_tax_files(finance_file, kaohsiung_file)
            elif finance_file:
                # 只有財政部版本
                load_tax_data_from_file(finance_file)
            elif kaohsiung_file:
                # 只有高雄市版本
                load_tax_data_from_file(kaohsiung_file)
            else:
                # 🔥 找到的檔案不符合命名規則，讓使用者手動選擇
                update_message("[提示] 找到多個稅額檔案，但無法識別版本：")
                for file in all_found_files:
                    update_message(f"     - {os.path.basename(file)}")
                update_message("        請點擊下方「手動選擇增值稅結果」按鈕")
        else:
            print("[警告] 未找到 tax_result.json")
            update_message("[警告] 未找到 tax_result.json 檔案")
            if expected_path:
                update_message(f"[提示] 預期路徑：{expected_path}")
            update_message("[建議] 請先執行「土地增值稅」程式計算稅額")
            update_message("        或點擊下方「手動選擇增值稅結果」按鈕")

    except Exception as e:
        error_msg = f"[錯誤] 讀取稅額失敗：{e}"
        print(error_msg)
        update_message(error_msg)
        import traceback
        traceback.print_exc()


    # === 💰 土地增值稅資訊（新區塊，售件專用）===
    land_value_tax_frame = create_section(scrollable_frame, "💰 土地增值稅資訊（售件專用）")

    # 🔥 公告現值也加入千分位
    if '公告現值' in data and data['公告現值']:
        try:
            # 先移除可能已有的千分位和空白
            cleaned_value = str(data['公告現值']).replace(',', '').replace(' ', '').strip()
            if cleaned_value and cleaned_value.replace('.', '').isdigit():
                # 轉為浮點數後格式化（保留小數）
                公告現值_num = float(cleaned_value)
                if '.' in cleaned_value:
                    # 有小數點，保留小數
                    data['公告現值'] = f"{公告現值_num:,.2f}"
                else:
                    # 無小數點，只顯示整數
                    data['公告現值'] = f"{int(公告現值_num):,}"
        except:
            pass  # 格式化失敗就保持原樣

    create_field(land_value_tax_frame, "公告現值（元/平方公尺）：", "公告現值", data, width=20)

    # 🔥 自動帶入稅額（如果有的話，加入千分位）
    if auto_load_tax_data:
        if '自用增值稅' not in data or not data['自用增值稅']:
            自用稅額 = auto_load_tax_data.get('自用住宅稅額', 0)
            if 自用稅額:
                data['自用增值稅'] = f"{int(自用稅額):,}"
            else:
                data['自用增值稅'] = ''

        if '一般增值稅' not in data or not data['一般增值稅']:
            一般稅額 = auto_load_tax_data.get('一般土地稅額', 0)
            if 一般稅額:
                data['一般增值稅'] = f"{int(一般稅額):,}"
            else:
                data['一般增值稅'] = ''

    create_field(land_value_tax_frame, "自用增值稅（元）：", "自用增值稅", data, width=20)
    create_field(land_value_tax_frame, "一般增值稅（元）：", "一般增值稅", data, width=20)

    # 🔥 手動選擇增值稅結果按鈕（置中）
    manual_tax_button_row = tk.Frame(land_value_tax_frame)
    manual_tax_button_row.pack(fill=tk.X, pady=10)
    tk.Button(
        manual_tax_button_row,
        text="📂 手動選擇增值稅結果",
        command=manual_select_tax_file,
        font=("微軟正黑體", 11),
        bg='#2196F3',
        fg='white',
        padx=20,
        pady=8,
        cursor="hand2"
    ).pack(anchor=tk.CENTER)

    # === 貸款資訊（售件專用）===
    loan_frame = create_section(scrollable_frame, "🏦 貸款資訊（售件專用）")
    create_field(loan_frame, "貸款銀行：", "貸款銀行", data, width=40)
    create_field(loan_frame, "設定金額：", "設定金額", data, width=20)

    # === 建物資訊區 ===
    building_frame = create_section(scrollable_frame, "🏗️ 建物資訊")
    create_field(building_frame, "竣工日期：", "竣工日期", data, width=30)
    create_field(building_frame, "建築結構：", "建築結構", data, width=30)
    
    # 🔥 新增：建物層數（地上/地下）
    floor_row = tk.Frame(building_frame)
    floor_row.pack(fill=tk.X, pady=3)
    tk.Label(floor_row, text="建物層數：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)

    # 🔥 新增：所有權型態（複選框）
    # ownership_row = tk.Frame(building_frame)
    # ownership_row.pack(fill=tk.X, pady=3)
    # tk.Label(ownership_row, text="所有權型態：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)

    

    tk.Label(floor_row, text="地上", font=entry_font).pack(side=tk.LEFT, padx=2)
    floor_above_entry = tk.Entry(floor_row, font=entry_font, width=5)
    floor_above_entry.insert(0, data.get('地上層數', ''))
    floor_above_entry.pack(side=tk.LEFT, padx=2)
    entries["地上層數"] = floor_above_entry
    tk.Label(floor_row, text="層", font=entry_font).pack(side=tk.LEFT, padx=2)

    tk.Label(floor_row, text="地下", font=entry_font).pack(side=tk.LEFT, padx=(15, 2))
    floor_below_entry = tk.Entry(floor_row, font=entry_font, width=5)
    floor_below_entry.insert(0, data.get('地下層數', ''))
    floor_below_entry.pack(side=tk.LEFT, padx=2)
    entries["地下層數"] = floor_below_entry
    tk.Label(floor_row, text="層", font=entry_font).pack(side=tk.LEFT, padx=2)

    # 座向（下拉選單）
    orient_row = tk.Frame(building_frame)
    orient_row.pack(fill=tk.X, pady=3)
    tk.Label(orient_row, text="座向：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    orient_var = tk.StringVar(value=data.get("座向", ""))
    orient_combo = ttk.Combobox(orient_row, textvariable=orient_var, font=combo_font, width=20,
                                  values=["座東朝西", "座西朝東", "座南朝北", "座北朝南", 
                                         "座東北朝西南", "座西南朝東北", "座西北朝東南", "座東南朝西北"])
    orient_combo.pack(side=tk.LEFT, padx=5)
    entries["座向"] = orient_var
    
    # === 格局資訊 ===
    layout_frame = create_section(scrollable_frame, "🚪 格局配置")

    # 隔局複選框（租件專用）
    layout_type_row = tk.Frame(layout_frame)
    layout_type_row.pack(fill=tk.X, pady=3)
    tk.Label(layout_type_row, text="格局：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)

    隔局_data = data.get("格局選項", [])
    隔局_vars = {}

    for option in ["開放式", "空地"]:
        var = tk.BooleanVar(value=(option in 隔局_data))
        cb = tk.Checkbutton(layout_type_row, text=option, variable=var, font=entry_font)
        cb.pack(side=tk.LEFT, padx=10)
        隔局_vars[option] = var

    entries["隔局"] = 隔局_vars

    # 格局（房廳衛）
    layout_row = tk.Frame(layout_frame)
    layout_row.pack(fill=tk.X, pady=3)

    tk.Label(layout_row, text="格局：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)

    layout_data = data.get("格局", {})

    # 🔥 房（輸入框在前）
    room_entry = tk.Entry(layout_row, font=entry_font, width=5)
    room_entry.insert(0, layout_data.get("房", ""))
    room_entry.pack(side=tk.LEFT, padx=2)
    entries["格局_房"] = room_entry
    tk.Label(layout_row, text="房", font=entry_font).pack(side=tk.LEFT, padx=2)

    # 🔥 廳（輸入框在前）
    living_entry = tk.Entry(layout_row, font=entry_font, width=5)
    living_entry.insert(0, layout_data.get("廳", ""))
    living_entry.pack(side=tk.LEFT, padx=2)
    entries["格局_廳"] = living_entry
    tk.Label(layout_row, text="廳", font=entry_font).pack(side=tk.LEFT, padx=2)

    # 🔥 衛（輸入框在前）
    bath_entry = tk.Entry(layout_row, font=entry_font, width=5)
    bath_entry.insert(0, layout_data.get("衛", ""))
    bath_entry.pack(side=tk.LEFT, padx=2)
    entries["格局_衛"] = bath_entry
    tk.Label(layout_row, text="衛", font=entry_font).pack(side=tk.LEFT, padx=2)

    # === 稅賦配置（租件專用）===
    tax_frame = create_section(scrollable_frame, "💰 稅賦配置（租件專用）")

    # ⭐ 確保 tax_data 是字典類型
    tax_data = data.get("稅賦", {})
    if not isinstance(tax_data, dict):
        tax_data = {}

    tax_items = ["所得稅", "補充保費", "營業稅", "房屋稅", "地價稅"]
    tax_vars = {}

    for tax_item in tax_items:
        tax_row = tk.Frame(tax_frame)
        tax_row.pack(fill=tk.X, pady=3)

        tk.Label(tax_row, text=f"{tax_item}：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)

        # ⭐ 確保預設值（如果沒有設定或為空，預設為「清除選擇」）
        current_value = tax_data.get(tax_item, "")
        # ⭐ 強制檢查：只有「出租方」或「承租方」才是有效值，其他一律設為「清除選擇」
        if current_value not in ["出租方", "承租方"]:
            current_value = "清除選擇"

        print(f"[稅賦配置] {tax_item}: 初始值='{current_value}'")  # ⭐ 除錯訊息

        # ⭐ 先創建變數，暫不設定值
        tax_var = tk.StringVar()

        landlord_radio = tk.Radiobutton(tax_row, text="出租方", variable=tax_var,
                                        value="出租方", font=combo_font)
        landlord_radio.pack(side=tk.LEFT, padx=10)

        tenant_radio = tk.Radiobutton(tax_row, text="承租方", variable=tax_var,
                                      value="承租方", font=combo_font)
        tenant_radio.pack(side=tk.LEFT, padx=10)

        # ⭐⭐⭐ 「清除選擇」使用特殊值，而不是空字串
        clear_radio = tk.Radiobutton(tax_row, text="清除選擇", variable=tax_var,
                                     value="清除選擇", font=combo_font)
        clear_radio.pack(side=tk.LEFT, padx=10)

        # ⭐ 在創建完所有 Radiobutton 後，才設定初始值
        tax_var.set(current_value)

        tax_vars[tax_item] = tax_var

    entries["稅賦"] = tax_vars

    # === 物件現況區 ===
    status_frame = create_section(scrollable_frame, "📦 物件現況")

    # 現況（下拉選單）
    status_row = tk.Frame(status_frame)
    status_row.pack(fill=tk.X, pady=3)
    tk.Label(status_row, text="現況：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    status_var = tk.StringVar(value=data.get("現況", ""))
    status_combo = ttk.Combobox(status_row, textvariable=status_var, font=combo_font, width=15,
                                values=["空屋", "空地", "自用", "出租中", "出借中", "其他"])
    status_combo.pack(side=tk.LEFT, padx=5)
    entries["現況"] = status_var

    # 🔥 條件輸入框（出租至/其他說明）
    tk.Label(status_row, text="", font=entry_font, width=2).pack(side=tk.LEFT)  # 間隔
    status_detail_label = tk.Label(status_row, text="", font=entry_font, width=8)
    status_detail_label.pack(side=tk.LEFT, padx=2)

    status_detail_entry = tk.Entry(status_row, font=entry_font, width=20, state='disabled')
    status_detail_entry.pack(side=tk.LEFT, padx=2)
    entries["現況說明"] = status_detail_entry

    # 🔥 根據現有資料填入
    現況說明 = data.get('現況說明', '')
    if 現況說明:
        status_detail_entry.config(state='normal')
        status_detail_entry.insert(0, 現況說明)
        if status_var.get() == "出租中":
            status_detail_label.config(text="出租至")
        elif status_var.get() == "其他":
            status_detail_label.config(text="說明：")

    # 🔥 監聽現況變化，動態顯示/隱藏輸入框
    def on_status_change(*args):
        current_status = status_var.get()
        if current_status == "出租中":
            status_detail_label.config(text="出租至")
            status_detail_entry.config(state='normal')
        elif current_status == "其他":
            status_detail_label.config(text="說明：")
            status_detail_entry.config(state='normal')
        else:
            status_detail_label.config(text="")
            status_detail_entry.config(state='disabled')
            status_detail_entry.delete(0, tk.END)

    status_var.trace('w', on_status_change)
    on_status_change()  # 初始化狀態

    # 鑰匙（改為下拉選單，預設「未填」）
    key_row = tk.Frame(status_frame)
    key_row.pack(fill=tk.X, pady=3)
    tk.Label(key_row, text="鑰匙：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    key_var = tk.StringVar(value=data.get("鑰匙", "未填"))  # 🔥 預設「未填」
    key_combo = ttk.Combobox(key_row, textvariable=key_var, font=combo_font, width=15,
                            values=["未填", "有", "無"])
    key_combo.pack(side=tk.LEFT, padx=5)
    entries["鑰匙"] = key_var
    
    # === 車位資訊 ===
    parking_frame = create_section(scrollable_frame, "🚗 車位資訊")

    parking_data = data.get("車位", {})

    # 🔥 車位有無（改為下拉選單，根據交易類型動態變化）
    park_row1 = tk.Frame(parking_frame)
    park_row1.pack(fill=tk.X, pady=3)
    tk.Label(park_row1, text="車位：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)

    # 🔥 根據目前的車位值決定初始值
    current_parking = parking_data.get("有無", "無")
    if current_parking in ["有，公設內", "有，獨立產權"]:
        park_initial_value = "有" if transaction_var.get() == "租件" else current_parking
    else:
        park_initial_value = current_parking

    park_var = tk.StringVar(value=park_initial_value)
    park_combo = ttk.Combobox(park_row1, textvariable=park_var, font=combo_font, width=15,
                                values=["有，公設內", "有，獨立產權", "無"])  # 預設為售件選項
    park_combo.pack(side=tk.LEFT, padx=5)
    entries["車位_有無"] = park_var

    # 車位型式（可複選）- 只有選「有」時才啟用
    park_row2 = tk.Frame(parking_frame)
    park_row2.pack(fill=tk.X, pady=3)
    tk.Label(park_row2, text="車位型式：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)

    現有型式 = parking_data.get("型式", "")

    平面_var = tk.BooleanVar(value='平面' in 現有型式)
    平面_check = tk.Checkbutton(park_row2, text="平面", variable=平面_var, font=entry_font)
    平面_check.pack(side=tk.LEFT, padx=5)
    entries["車位_型式_平面"] = 平面_var

    機械_var = tk.BooleanVar(value='機械' in 現有型式)
    機械_check = tk.Checkbutton(park_row2, text="機械", variable=機械_var, font=entry_font)
    機械_check.pack(side=tk.LEFT, padx=5)
    entries["車位_型式_機械"] = 機械_var

    # 車位編號
    park_row3 = tk.Frame(parking_frame)
    park_row3.pack(fill=tk.X, pady=3)
    tk.Label(park_row3, text="車位編號：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    park_num_entry = tk.Entry(park_row3, font=entry_font, width=30)
    park_num_entry.insert(0, parking_data.get("編號", ""))
    park_num_entry.pack(side=tk.LEFT, padx=5)
    entries["車位_編號"] = park_num_entry

    # 🔥 當車位選「無」時，禁用型式和編號
    def on_parking_change(*args):
        current_value = park_var.get()
        if current_value == "無":
            平面_var.set(False)
            機械_var.set(False)
            平面_check.config(state='disabled')
            機械_check.config(state='disabled')
            park_num_entry.config(state='disabled')
        else:
            平面_check.config(state='normal')
            機械_check.config(state='normal')
            park_num_entry.config(state='normal')

    park_var.trace('w', on_parking_change)
    on_parking_change()  # 初始化狀態
    
    # === 其他設施 ===
    facility_frame = create_section(scrollable_frame, "🏢 其他設施")

    # 電梯（改為條件輸入）
    elev_row = tk.Frame(facility_frame)
    elev_row.pack(fill=tk.X, pady=3)
    tk.Label(elev_row, text="電梯：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    elev_var = tk.StringVar(value=data.get("電梯", ""))
    elev_combo = ttk.Combobox(elev_row, textvariable=elev_var, font=combo_font, width=10,
                                values=["有", "無"])
    elev_combo.pack(side=tk.LEFT, padx=5)
    entries["電梯"] = elev_var

    # 🔥 電梯數量
    tk.Label(elev_row, text="數量：", font=entry_font).pack(side=tk.LEFT, padx=5)
    elev_count_entry = tk.Entry(elev_row, font=entry_font, width=5)
    elev_count_entry.insert(0, data.get("電梯數量", ""))
    elev_count_entry.pack(side=tk.LEFT, padx=2)
    entries["電梯數量"] = elev_count_entry
    tk.Label(elev_row, text="部", font=entry_font).pack(side=tk.LEFT, padx=2)

    # 🔥 一層幾戶
    tk.Label(elev_row, text="一層", font=entry_font).pack(side=tk.LEFT, padx=(15, 2))
    elev_households_entry = tk.Entry(elev_row, font=entry_font, width=5)
    elev_households_entry.insert(0, data.get("一層幾戶", ""))
    elev_households_entry.pack(side=tk.LEFT, padx=2)
    entries["一層幾戶"] = elev_households_entry
    tk.Label(elev_row, text="戶", font=entry_font).pack(side=tk.LEFT, padx=2)

    # 當電梯選「無」時，禁用相關輸入
    def on_elevator_change(*args):
        if elev_var.get() == "無":
            elev_count_entry.config(state='disabled')
            elev_households_entry.config(state='disabled')
            elev_count_entry.delete(0, tk.END)
            elev_households_entry.delete(0, tk.END)
        else:
            elev_count_entry.config(state='normal')
            elev_households_entry.config(state='normal')

    elev_var.trace('w', on_elevator_change)
    on_elevator_change()

    # 🔥 管理費欄位（售件和租件共用）
    管理費_row = tk.Frame(facility_frame)
    管理費_row.pack(fill=tk.X, pady=2)
    tk.Label(管理費_row, text="管理費（元/月）：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    管理費_entry = tk.Entry(管理費_row, font=entry_font, width=20)
    管理費_entry.pack(side=tk.LEFT, padx=5)
    管理費_entry.insert(0, data.get("管理費", ""))
    entries["管理費"] = 管理費_entry

    # 🔥 新增：履約保證（改為下拉選單，售件專用）
    guarantee_row = tk.Frame(facility_frame)
    guarantee_row.pack(fill=tk.X, pady=3)
    tk.Label(guarantee_row, text="履約保證：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    guarantee_var = tk.StringVar(value=data.get("履約保證", ""))
    guarantee_combo = ttk.Combobox(guarantee_row, textvariable=guarantee_var, font=combo_font, width=15,
                                    values=["同意", "不同意"])
    guarantee_combo.pack(side=tk.LEFT, padx=5)
    entries["履約保證"] = guarantee_var

    # 🔥 新增：廠房保存登記
    factory_reg_row = tk.Frame(facility_frame)
    factory_reg_row.pack(fill=tk.X, pady=3)
    tk.Label(factory_reg_row, text="廠房保存登記：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    factory_reg_var = tk.StringVar(value=data.get("廠房保存登記", ""))
    factory_reg_combo = ttk.Combobox(factory_reg_row, textvariable=factory_reg_var, font=combo_font, width=15,
                                    values=["是", "否"])
    factory_reg_combo.pack(side=tk.LEFT, padx=5)
    entries["廠房保存登記"] = factory_reg_var

    # 🔥 新增：工廠登記（可/不可）
    factory_row = tk.Frame(facility_frame)
    factory_row.pack(fill=tk.X, pady=3)
    tk.Label(factory_row, text="工廠登記：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    factory_var = tk.StringVar(value=data.get("工廠登記", ""))
    factory_combo = ttk.Combobox(factory_row, textvariable=factory_var, font=combo_font, width=15,
                                values=["可", "不可"])
    factory_combo.pack(side=tk.LEFT, padx=5)
    entries["工廠登記"] = factory_var

    # 🔥 新增：分割銷售
    division_row = tk.Frame(facility_frame)
    division_row.pack(fill=tk.X, pady=3)
    tk.Label(division_row, text="分割銷售：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    division_var = tk.StringVar(value=data.get("分割銷售", ""))
    division_combo = ttk.Combobox(division_row, textvariable=division_var, font=combo_font, width=15,
                                    values=["是", "否"])
    division_combo.pack(side=tk.LEFT, padx=5)
    entries["分割銷售"] = division_var

    # 🔥 新增：天車
    crane_row = tk.Frame(facility_frame)
    crane_row.pack(fill=tk.X, pady=3)
    tk.Label(crane_row, text="天車：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    crane_var = tk.StringVar(value=data.get("天車", ""))
    crane_combo = ttk.Combobox(crane_row, textvariable=crane_var, font=combo_font, width=15,
                                values=["有", "無"])
    crane_combo.pack(side=tk.LEFT, padx=5)
    entries["天車"] = crane_var

    # 🔥 新增：樑下（售件專用）
    樑下_row = tk.Frame(facility_frame)
    樑下_row.pack(fill=tk.X, pady=3)
    tk.Label(樑下_row, text="樑下（米）：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    樑下_entry = tk.Entry(樑下_row, font=entry_font, width=20)
    樑下_entry.insert(0, data.get("樑下", ""))
    樑下_entry.pack(side=tk.LEFT, padx=5)
    entries["樑下"] = 樑下_entry

    # 🔥 中脊欄位（需要自動加入約字和公尺單位）
    中脊_row = tk.Frame(facility_frame)
    中脊_row.pack(fill=tk.X, pady=3)
    tk.Label(中脊_row, text="中脊（米）：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    中脊_entry = tk.Entry(中脊_row, font=entry_font, width=20)
    中脊_entry.insert(0, data.get("中脊", ""))
    中脊_entry.pack(side=tk.LEFT, padx=5)
    entries["中脊"] = 中脊_entry

    # 🔥 滴水欄位（需要自動加入約字和公尺單位）
    滴水_row = tk.Frame(facility_frame)
    滴水_row.pack(fill=tk.X, pady=3)
    tk.Label(滴水_row, text="滴水（米）：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    滴水_entry = tk.Entry(滴水_row, font=entry_font, width=20)
    滴水_entry.insert(0, data.get("滴水", ""))
    滴水_entry.pack(side=tk.LEFT, padx=5)
    entries["滴水"] = 滴水_entry

    # 🔥 新增：查詢學區按鈕（獨立一行）
    query_btn_row = tk.Frame(facility_frame)
    query_btn_row.pack(fill=tk.X, pady=3)

    # 學區檔轉換功能
    def open_school_district_converter():
        """開啟學區檔轉換對話框"""
        dialog = tk.Toplevel(root)
        dialog.title("學區檔轉換")
        dialog.geometry("800x600")
        dialog.resizable(False, False)

        # 置中顯示
        dialog.transient(root)
        dialog.grab_set()

        # 縣市選擇變數
        city_var = tk.StringVar(value="高雄市")

        # 檔案路徑/URL變數（115學年度）
        elementary_pdf_path = tk.StringVar(value="https://orgws.kcg.gov.tw/001/KcgOrgUploadFiles/223/relfile/0/70200/5bc55e74-6414-49b5-ad07-6d97c540554a.pdf")
        junior_pdf_path = tk.StringVar(value="https://orgws.kcg.gov.tw/001/KcgOrgUploadFiles/225/RelFile/0/10487/ff8b78bf-fbb0-4fab-9daa-b7fea2db65d2.pdf")

        # 標題
        tk.Label(
            dialog,
            text="學區檔案轉換為 JSON 格式",
            font=("微軟正黑體", 18, "bold"),
            fg="#FF9800"
        ).pack(pady=10)

        # 縣市選擇區域
        city_frame = tk.LabelFrame(dialog, text="選擇縣市", font=("微軟正黑體", 14), padx=15, pady=10)
        city_frame.pack(fill=tk.X, padx=30, pady=8)

        tk.Radiobutton(
            city_frame,
            text="高雄市",
            variable=city_var,
            value="高雄市",
            font=("微軟正黑體", 13)
        ).pack(anchor=tk.W, pady=3)

        tk.Radiobutton(
            city_frame,
            text="台南市",
            variable=city_var,
            value="台南市",
            font=("微軟正黑體", 13)
        ).pack(anchor=tk.W, pady=3)

        tk.Radiobutton(
            city_frame,
            text="屏東縣",
            variable=city_var,
            value="屏東縣",
            font=("微軟正黑體", 13)
        ).pack(anchor=tk.W, pady=3)

        # 檔案選擇區域
        file_frame = tk.LabelFrame(dialog, text="PDF 來源（支援本地檔案或 URL）", font=("微軟正黑體", 14), padx=15, pady=10)
        file_frame.pack(fill=tk.X, padx=30, pady=8)

        # 說明文字
        tk.Label(
            file_frame,
            text="可輸入 URL 或點擊「瀏覽」選擇本地檔案",
            font=("微軟正黑體", 11),
            fg="#666"
        ).pack(anchor=tk.W, pady=(0, 10))

        # 國小學區檔案
        elementary_row = tk.Frame(file_frame)
        elementary_row.pack(fill=tk.X, pady=8)

        tk.Label(elementary_row, text="國小學區:", font=("微軟正黑體", 13), width=10, anchor=tk.W).pack(side=tk.LEFT)
        tk.Entry(elementary_row, textvariable=elementary_pdf_path, width=40, font=("微軟正黑體", 11)).pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        tk.Button(
            elementary_row,
            text="瀏覽",
            command=lambda: elementary_pdf_path.set(
                filedialog.askopenfilename(
                    title="選擇國小學區 PDF 檔案",
                    filetypes=[("PDF 檔案", "*.pdf"), ("所有檔案", "*.*")]
                )
            ),
            font=("微軟正黑體", 11),
            padx=10,
            pady=3
        ).pack(side=tk.LEFT)

        # 國中學區檔案
        junior_row = tk.Frame(file_frame)
        junior_row.pack(fill=tk.X, pady=8)

        tk.Label(junior_row, text="國中學區:", font=("微軟正黑體", 13), width=10, anchor=tk.W).pack(side=tk.LEFT)
        tk.Entry(junior_row, textvariable=junior_pdf_path, width=40, font=("微軟正黑體", 11)).pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        tk.Button(
            junior_row,
            text="瀏覽",
            command=lambda: junior_pdf_path.set(
                filedialog.askopenfilename(
                    title="選擇國中學區 PDF 檔案",
                    filetypes=[("PDF 檔案", "*.pdf"), ("所有檔案", "*.*")]
                )
            ),
            font=("微軟正黑體", 11),
            padx=10,
            pady=3
        ).pack(side=tk.LEFT)

        # 執行轉換函數
        def execute_conversion():
            # 🔥 確保 messagebox 在巢狀函數中可用
            from tkinter import messagebox

            city = city_var.get()
            elem_source = elementary_pdf_path.get().strip()
            jun_source = junior_pdf_path.get().strip()

            if not elem_source or not jun_source:
                messagebox.showwarning("缺少來源", "請輸入國小和國中學區的 PDF 檔案路徑或 URL！")
                return

            dialog.destroy()

            # 🔥 轉換前先記下舊版 JSON 的學校/村里數（用於轉換後驗證對照）
            from base_dir_helper import get_base_dir
            _base_dir = get_base_dir()
            if city == '高雄市':
                _old_jun_path = os.path.join(_base_dir, '高雄市_國中學區資料.json')
                _old_elem_path = os.path.join(_base_dir, '高雄市_國小學區資料.json')
            elif city == '台南市':
                _old_jun_path = os.path.join(_base_dir, '台南市國中學區資料.json')
                _old_elem_path = os.path.join(_base_dir, '台南市國小學區資料.json')
            elif city == '屏東縣':
                _old_jun_path = os.path.join(_base_dir, '屏東縣國中學區資料.json')
                _old_elem_path = os.path.join(_base_dir, '屏東縣國小學區資料.json')
            else:
                _old_jun_path = _old_elem_path = None

            def _count_json_file(path):
                """讀 JSON，回傳 (學校數, 村里總數)；不存在或失敗回傳 None"""
                if not path or not os.path.exists(path):
                    return None
                try:
                    with open(path, 'r', encoding='utf-8') as f:
                        d = json.load(f)
                    schools = len(d)
                    villages = sum(len(v) for s in d.values() for v in s.values())
                    return (schools, villages)
                except Exception:
                    return None

            _old_jun_counts = _count_json_file(_old_jun_path)
            _old_elem_counts = _count_json_file(_old_elem_path)

            try:
                import subprocess
                import urllib.request
                import urllib.error

                update_message(f"[資訊] 開始 {city} 學區檔轉換...")

                # 處理國小學區 PDF
                if elem_source.startswith('http://') or elem_source.startswith('https://'):
                    update_message(f"[下載] 正在下載國小學區 PDF...")
                    try:
                        elem_path = '001.國小學區來源檔.pdf'
                        # 🔥 使用 requests 庫 + 完整瀏覽器 headers + 忽略 SSL 驗證
                        # 高雄市政府網站需要這些設定才能正常下載
                        import requests
                        import warnings
                        warnings.filterwarnings('ignore', message='Unverified HTTPS request')

                        headers = {
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7',
                            'Accept-Language': 'zh-TW,zh;q=0.9,en-US;q=0.8,en;q=0.7',
                            'Accept-Encoding': 'gzip, deflate, br',
                            'Connection': 'keep-alive',
                            'Upgrade-Insecure-Requests': '1',
                        }

                        update_message(f"[DEBUG] 國小學區 URL: {elem_source}")
                        response = _get_with_ssl_fallback(elem_source, headers=headers, timeout=60)
                        update_message(f"[DEBUG] HTTP 狀態碼: {response.status_code}")

                        if response.status_code == 200:
                            with open(elem_path, 'wb') as out_file:
                                out_file.write(response.content)
                            update_message(f"[成功] 國小學區 PDF 下載完成 (大小: {len(response.content)} bytes)")
                        else:
                            raise Exception(f"HTTP 錯誤: {response.status_code}")
                    except Exception as e:
                        messagebox.showerror("下載失敗", f"下載國小學區 PDF 失敗！\n\n{e}")
                        return
                else:
                    # 智能路徑處理：支援絕對路徑、相對路徑
                    elem_path = elem_source

                    # 如果不是絕對路徑或檔案不存在，嘗試多個位置
                    if not os.path.isabs(elem_source) or not os.path.exists(elem_source):
                        # 嘗試 1: 轉換為絕對路徑（相對於當前工作目錄）
                        candidate1 = os.path.abspath(elem_source)
                        # 嘗試 2: 相對於主程式開發目錄
                        from base_dir_helper import get_base_dir
                        candidate2 = os.path.join(get_base_dir(), elem_source)

                        if os.path.exists(candidate1):
                            elem_path = candidate1
                        elif os.path.exists(candidate2):
                            elem_path = candidate2
                        else:
                            messagebox.showerror("檔案不存在",
                                f"國小學區 PDF 檔案不存在！\n\n已嘗試查找：\n1. {elem_source}\n2. {candidate1}\n3. {candidate2}")
                            return

                    update_message(f"[資訊] 使用國小學區檔案: {elem_path}")

                # 處理國中學區 PDF
                if jun_source.startswith('http://') or jun_source.startswith('https://'):
                    update_message(f"[下載] 正在下載國中學區 PDF...")
                    try:
                        jun_path = '001.國中學區來源檔.pdf'
                        # 🔥 使用 requests 庫 + 完整瀏覽器 headers + 忽略 SSL 驗證
                        # 高雄市政府網站需要這些設定才能正常下載
                        import requests
                        import warnings
                        warnings.filterwarnings('ignore', message='Unverified HTTPS request')

                        headers = {
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7',
                            'Accept-Language': 'zh-TW,zh;q=0.9,en-US;q=0.8,en;q=0.7',
                            'Accept-Encoding': 'gzip, deflate, br',
                            'Connection': 'keep-alive',
                            'Upgrade-Insecure-Requests': '1',
                        }

                        update_message(f"[DEBUG] 國中學區 URL: {jun_source}")
                        response = _get_with_ssl_fallback(jun_source, headers=headers, timeout=60)
                        update_message(f"[DEBUG] HTTP 狀態碼: {response.status_code}")

                        if response.status_code == 200:
                            with open(jun_path, 'wb') as out_file:
                                out_file.write(response.content)
                            update_message(f"[成功] 國中學區 PDF 下載完成 (大小: {len(response.content)} bytes)")
                        else:
                            raise Exception(f"HTTP 錯誤: {response.status_code}")
                    except Exception as e:
                        messagebox.showerror("下載失敗", f"下載國中學區 PDF 失敗！\n\n{e}")
                        return
                else:
                    # 智能路徑處理：支援絕對路徑、相對路徑
                    jun_path = jun_source

                    # 如果不是絕對路徑或檔案不存在，嘗試多個位置
                    if not os.path.isabs(jun_source) or not os.path.exists(jun_source):
                        # 嘗試 1: 轉換為絕對路徑（相對於當前工作目錄）
                        candidate1 = os.path.abspath(jun_source)
                        # 嘗試 2: 相對於主程式開發目錄
                        from base_dir_helper import get_base_dir
                        candidate2 = os.path.join(get_base_dir(), jun_source)

                        if os.path.exists(candidate1):
                            jun_path = candidate1
                        elif os.path.exists(candidate2):
                            jun_path = candidate2
                        else:
                            messagebox.showerror("檔案不存在",
                                f"國中學區 PDF 檔案不存在！\n\n已嘗試查找：\n1. {jun_source}\n2. {candidate1}\n3. {candidate2}")
                            return

                    update_message(f"[資訊] 使用國中學區檔案: {jun_path}")

                update_message(f"[資訊] 正在解析 PDF 檔案，請稍候...")

                # 根據城市選擇對應的轉換程式
                if city == '高雄市':
                    script_name = 'build_school_district_data.py'
                    output_prefix = ''
                elif city == '台南市':
                    script_name = 'build_tainan_school_district_data.py'
                    output_prefix = '台南市'
                elif city == '屏東縣':
                    script_name = 'build_pingtung_school_district_data.py'
                    output_prefix = '屏東縣'
                else:
                    messagebox.showerror("錯誤", f"不支援的城市：{city}")
                    return

                # 🔥 直接導入並執行轉換函數，避免啟動新的程式視窗
                try:
                    if city == '高雄市':
                        update_message(f"[DEBUG] 開始導入 build_school_district_data...")
                        import build_school_district_data
                        update_message(f"[DEBUG] 導入成功，模組路徑: {build_school_district_data.__file__}")
                        # 重新導入以確保使用最新版本
                        import importlib
                        update_message(f"[DEBUG] 重新載入模組...")
                        importlib.reload(build_school_district_data)
                        update_message(f"[DEBUG] 重新載入成功")

                        # 執行轉換
                        update_message(f"[DEBUG] 開始解析國中學區 PDF...")
                        junior_data = build_school_district_data.build_school_district_data('001.國中學區來源檔.pdf', '國中')
                        update_message(f"[DEBUG] 國中學區解析完成，共 {len(junior_data)} 所學校")
                        junior_path = build_school_district_data.save_to_json(junior_data, '高雄市_國中學區資料.json')
                        update_message(f"[DEBUG] 國中學區 JSON 儲存完成: {junior_path}")

                        update_message(f"[DEBUG] 開始解析國小學區 PDF...")
                        elementary_data = build_school_district_data.build_school_district_data('001.國小學區來源檔.pdf', '國小')
                        update_message(f"[DEBUG] 國小學區解析完成，共 {len(elementary_data)} 所學校")
                        elementary_path = build_school_district_data.save_to_json(elementary_data, '高雄市_國小學區資料.json')
                        update_message(f"[DEBUG] 國小學區 JSON 儲存完成: {elementary_path}")

                        success = True
                    elif city == '台南市':
                        import build_tainan_school_district_data
                        import importlib
                        importlib.reload(build_tainan_school_district_data)

                        # 取得主程式目錄（JSON檔案輸出位置）
                        from base_dir_helper import get_base_dir
                        base_dir = get_base_dir()

                        # 使用使用者指定的 PDF 檔案路徑（elem_path 和 jun_path）
                        # 執行轉換
                        elementary_data = build_tainan_school_district_data.build_tainan_school_district_data(elem_path, '國小')
                        elementary_path = build_tainan_school_district_data.save_json(elementary_data, os.path.join(base_dir, '台南市國小學區資料.json'))

                        junior_data = build_tainan_school_district_data.build_tainan_school_district_data(jun_path, '國中')
                        junior_path = build_tainan_school_district_data.save_json(junior_data, os.path.join(base_dir, '台南市國中學區資料.json'))

                        success = True
                    elif city == '屏東縣':
                        import build_pingtung_school_district_data
                        import importlib
                        importlib.reload(build_pingtung_school_district_data)

                        # 取得主程式目錄（JSON檔案輸出位置）
                        from base_dir_helper import get_base_dir
                        base_dir = get_base_dir()

                        # 使用使用者指定的 PDF 檔案路徑（elem_path 和 jun_path）
                        # 執行轉換
                        elementary_data = build_pingtung_school_district_data.build_pingtung_school_district_data(elem_path, '國小')
                        elementary_path = build_pingtung_school_district_data.save_json(elementary_data, os.path.join(base_dir, '屏東縣國小學區資料.json'))

                        junior_data = build_pingtung_school_district_data.build_pingtung_school_district_data(jun_path, '國中')
                        junior_path = build_pingtung_school_district_data.save_json(junior_data, os.path.join(base_dir, '屏東縣國中學區資料.json'))

                        success = True
                    else:
                        success = False

                except Exception as e:
                    success = False
                    error_msg = str(e)
                    import traceback
                    error_detail = traceback.format_exc()

                if success:
                    # 🔥 使用轉換函數返回的完整路徑
                    junior_file = junior_path
                    elementary_file = elementary_path

                    update_message(f"[成功] {city} 學區檔轉換完成！")
                    update_message(f"[資訊] 已生成：")
                    update_message(f"  - {junior_file}")
                    update_message(f"  - {elementary_file}")

                    # 🔥 驗證：對照舊版筆數，警告大幅減少的情況
                    def _count_data_dict(d):
                        sch = len(d) if d else 0
                        vil = sum(len(v) for s in d.values() for v in s.values()) if d else 0
                        return (sch, vil)

                    def _validate(label, old, new):
                        n_school, n_village = new
                        if old is None:
                            return f"{label}: {n_school} 校 / {n_village} 村里 (首次建立)", False
                        o_school, o_village = old
                        sch_d = n_school - o_school
                        vil_d = n_village - o_village
                        msg = f"{label}: {n_school} 校 ({sch_d:+d}) / {n_village} 村里 ({vil_d:+d})"
                        # 警告：村里數減少超過 5% 或 10 個（取較大者）
                        threshold = max(10, int(o_village * 0.05))
                        is_warning = vil_d < -threshold
                        return msg, is_warning

                    _new_jun_counts = _count_data_dict(junior_data)
                    _new_elem_counts = _count_data_dict(elementary_data)
                    _jun_msg, _jun_warn = _validate("國中", _old_jun_counts, _new_jun_counts)
                    _elem_msg, _elem_warn = _validate("國小", _old_elem_counts, _new_elem_counts)
                    update_message(f"[驗證] {_jun_msg}{'  ⚠ 大幅減少' if _jun_warn else ''}")
                    update_message(f"[驗證] {_elem_msg}{'  ⚠ 大幅減少' if _elem_warn else ''}")
                    _has_warning = _jun_warn or _elem_warn

                    # 建立自訂的完成訊息對話框（字體較大）
                    success_dialog = tk.Toplevel(root)
                    success_dialog.title("轉換完成")
                    success_dialog.transient(root)
                    success_dialog.grab_set()

                    # 🔥 設定最小尺寸，允許調整大小（驗證資訊需要較高空間）
                    dialog_width = 700
                    dialog_height = 540
                    success_dialog.minsize(500, 400)
                    success_dialog.resizable(True, True)

                    # 🔥 置中顯示（先更新以獲取正確尺寸）
                    success_dialog.update_idletasks()
                    screen_width = success_dialog.winfo_screenwidth()
                    screen_height = success_dialog.winfo_screenheight()
                    x = (screen_width // 2) - (dialog_width // 2)
                    y = (screen_height // 2) - (dialog_height // 2)
                    success_dialog.geometry(f"{dialog_width}x{dialog_height}+{x}+{y}")

                    # 🔥 確保彈窗顯示在最前面
                    success_dialog.attributes('-topmost', True)
                    success_dialog.lift()
                    success_dialog.focus_force()

                    # 🔥 主框架，使用 pack 填滿
                    main_frame = tk.Frame(success_dialog)
                    main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=15)

                    # 圖示標籤
                    tk.Label(
                        main_frame,
                        text="ℹ",
                        font=("微軟正黑體", 40),
                        fg="#2196F3"
                    ).pack(pady=(10, 10))

                    # 標題文字
                    tk.Label(
                        main_frame,
                        text=f"{city} 學區檔已成功轉換為 JSON 檔案！",
                        font=("微軟正黑體", 14, "bold"),
                        fg="#333"
                    ).pack(pady=5)

                    # 🔥 驗證資訊（學校數、村里數對照）
                    validation_text = f"{_jun_msg}\n{_elem_msg}"
                    validation_color = "#d32f2f" if _has_warning else "#2e7d32"
                    tk.Label(
                        main_frame,
                        text="📊 資料驗證：",
                        font=("微軟正黑體", 11, "bold"),
                        fg="#333"
                    ).pack(pady=(8, 2))
                    tk.Label(
                        main_frame,
                        text=validation_text,
                        font=("Consolas", 10),
                        fg=validation_color,
                        justify=tk.LEFT
                    ).pack(pady=(0, 5))
                    if _has_warning:
                        tk.Label(
                            main_frame,
                            text="⚠ 村里數較舊版明顯減少，請確認 PDF 內容是否正確",
                            font=("微軟正黑體", 10, "bold"),
                            fg="#d32f2f",
                            wraplength=600
                        ).pack(pady=(0, 8))

                    # 生成檔案資訊
                    tk.Label(
                        main_frame,
                        text="生成檔案：",
                        font=("微軟正黑體", 12),
                        fg="#666"
                    ).pack(pady=(8, 5))

                    # 🔥 顯示完整路徑（使用 wraplength 自動換行）
                    files_text = f"{junior_file}\n\n{elementary_file}"
                    files_label = tk.Label(
                        main_frame,
                        text=files_text,
                        font=("微軟正黑體", 10),
                        fg="#666",
                        justify=tk.LEFT,
                        wraplength=600
                    )
                    files_label.pack(pady=5, fill=tk.X)

                    # 🔥 底部框架（固定確定按鈕位置）
                    bottom_frame = tk.Frame(main_frame)
                    bottom_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(15, 10))

                    # 確定按鈕：關閉對話框（不再自動執行查詢附近學區）
                    ok_btn = tk.Button(
                        bottom_frame,
                        text="確定",
                        command=success_dialog.destroy,
                        font=("微軟正黑體", 12),
                        bg="#2196F3",
                        fg="white",
                        padx=30,
                        pady=8,
                        cursor="hand2"
                    )
                    ok_btn.pack()

                    # 🔥 綁定 Enter 鍵關閉對話框
                    success_dialog.bind('<Return>', lambda e: success_dialog.destroy())
                else:
                    # 🔥 使用 try-except 中捕獲的錯誤訊息
                    if 'error_msg' not in locals():
                        error_msg = "未知錯誤"
                    if 'error_detail' not in locals():
                        error_detail = ""

                    update_message(f"[錯誤] {city} 轉換失敗：{error_msg}")
                    update_message(f"[詳細] {error_detail}")
                    messagebox.showerror("轉換失敗", f"{city} 學區檔轉換失敗！\n\n錯誤訊息：\n{error_msg}")

            except Exception as e:
                update_message(f"[錯誤] {city} 學區檔轉換發生異常：{e}")
                messagebox.showerror("執行錯誤", f"執行 {city} 學區檔轉換時發生錯誤：\n\n{e}")

        # 按鈕區域
        button_frame = tk.Frame(dialog)
        button_frame.pack(pady=20)

        tk.Button(
            button_frame,
            text="執行轉換",
            command=execute_conversion,
            font=("微軟正黑體", 14, "bold"),
            bg='#4CAF50',
            fg='white',
            padx=30,
            pady=10,
            cursor="hand2"
        ).pack(side=tk.LEFT, padx=15)

        tk.Button(
            button_frame,
            text="取消",
            command=dialog.destroy,
            font=("微軟正黑體", 14),
            padx=30,
            pady=10,
            cursor="hand2"
        ).pack(side=tk.LEFT, padx=15)

    # 🔥 helper：找 PDF 中學校所在的頁次
    def _find_school_pages(pdf_path, schools):
        """回傳 {school: page_number}"""
        if not pdf_path or not schools:
            return {}
        results = {}
        import os
        if not os.path.exists(pdf_path):
            return {}
        # 🔥 依序嘗試多個 PDF 庫，確保打包版(dist-0 有 pypdf)與開發環境(有 PyPDF2/fitz)都能讀
        pages_text = None
        try:
            from pypdf import PdfReader
            pages_text = [(p.extract_text() or '') for p in PdfReader(pdf_path).pages]
        except Exception:
            try:
                from PyPDF2 import PdfReader
                pages_text = [(p.extract_text() or '') for p in PdfReader(pdf_path).pages]
            except Exception:
                try:
                    import fitz
                    _doc = fitz.open(pdf_path)
                    pages_text = [_doc[i].get_text() for i in range(len(_doc))]
                    _doc.close()
                except Exception as _e:
                    update_message(f"[警告] 讀 PDF 失敗 {pdf_path}：{_e}")
                    return {}
        for i, text in enumerate(pages_text):
            for sch in schools:
                if sch not in results and sch in text:
                    results[sch] = i + 1
        return results

    def _find_pdf_path(filename):
        """在 BASE_DIR 跟上一層找 PDF"""
        import os
        candidates = [
            os.path.join(BASE_DIR, filename),
            os.path.join(os.path.dirname(BASE_DIR), filename),
        ]
        for p in candidates:
            if os.path.exists(p):
                return p
        return None

    # 🔥 各縣市對應的 115 PDF 檔名
    _CITY_115_PDF = {
        '高雄市': ('115年國小學區.pdf', '115年國中學區.pdf'),
        '台南市': ('115年台南市國小學區.pdf', '115年台南市國中學區.pdf'),
        '屏東縣': ('115年屏東縣國小學區.pdf', '115年屏東縣國中學區.pdf'),
    }

    # 🔥 整合版彈窗：一次顯示查詢結果 + 提供 115 PDF 比對按鈕（三縣市通用）
    def _show_kh_result_with_compare(village, neighbor, school_names, school_text, city='高雄市'):
        """整合查詢結果與 115 PDF 對照功能"""
        import os
        import tkinter as tk

        elem_schools = [s for s in school_names if '國小' in s]
        jun_schools = [s for s in school_names if '國中' in s]

        elem_pdf_name, jun_pdf_name = _CITY_115_PDF.get(city, ('115年國小學區.pdf', '115年國中學區.pdf'))
        elem_pdf = _find_pdf_path(elem_pdf_name)
        jun_pdf = _find_pdf_path(jun_pdf_name)

        elem_pages = _find_school_pages(elem_pdf, elem_schools) if elem_pdf else {}
        jun_pages = _find_school_pages(jun_pdf, jun_schools) if jun_pdf else {}

        dlg = tk.Toplevel(root)
        dlg.title(f"學區查詢成功（114 學年資料 - {city}）")
        dlg.geometry("680x560")
        dlg.update_idletasks()
        x = (dlg.winfo_screenwidth() - 680) // 2
        y = (dlg.winfo_screenheight() - 560) // 2
        dlg.geometry(f"+{x}+{y}")

        # 標題
        tk.Label(dlg, text=f"學區查詢成功（114 學年資料）",
                 font=("微軟正黑體", 16, "bold"), fg="#1976D2",
                 pady=15).pack()

        # 查詢結果
        result_frame = tk.Frame(dlg)
        result_frame.pack(fill=tk.X, padx=25, pady=5)
        info_text = f"村里：{village}　|　鄰：第{neighbor}鄰\n\n"
        if elem_schools:
            info_text += f"★ 國小：{ '、'.join(elem_schools)}\n"
        if jun_schools:
            info_text += f"★ 國中：{'、'.join(jun_schools)}\n"
        info_text += f"\n已自動填入「鄰近學校」欄位"
        tk.Label(result_frame, text=info_text,
                 font=("微軟正黑體", 12), justify=tk.LEFT,
                 anchor='w').pack(fill=tk.X)

        # 警示提醒
        warn_frame = tk.Frame(dlg, bg='#FFF3CD', relief=tk.SOLID, borderwidth=1)
        warn_frame.pack(fill=tk.X, padx=25, pady=15)
        tk.Label(warn_frame,
                 text="💡 此結果為 114 學年資料\n   建議按下方按鈕開啟 115 學年 PDF 比對\n   先複製里名 → 開啟 PDF → 用 Ctrl+F 搜尋里名\n   若有差異，請手動修改右側「鄰近學校」欄位",
                 font=("微軟正黑體", 11), bg='#FFF3CD', fg='#856404',
                 justify=tk.LEFT, anchor='w',
                 padx=12, pady=10).pack(fill=tk.X)

        # 🔥 複製里名按鈕（複製後可在 PDF 用 Ctrl+F 搜尋）
        copy_frame = tk.Frame(dlg)
        copy_frame.pack(pady=(0, 8))

        def _copy_village():
            try:
                dlg.clipboard_clear()
                dlg.clipboard_append(village)
                dlg.update()
                update_message(f"📋 已複製「{village}」到剪貼簿，可在 PDF 按 Ctrl+F 搜尋")
            except Exception as _e:
                update_message(f"⚠ 複製失敗：{_e}")

        tk.Button(copy_frame, text=f"📋 複製里名「{village}」",
                  command=_copy_village,
                  font=("微軟正黑體", 11, "bold"),
                  bg='#9C27B0', fg='white',
                  width=30, height=2,
                  cursor='hand2').pack()

        # PDF 比對按鈕區
        btn_frame = tk.Frame(dlg)
        btn_frame.pack(pady=10)

        def _open_pdf(pdf_path, label, pages):
            if not pdf_path:
                update_message(f"⚠ 找不到 {label} PDF 檔")
                return
            try:
                os.startfile(pdf_path)
                update_message(f"📄 已開啟 {label}（請用 Ctrl+F 搜尋里名）")
                for sch, page in pages.items():
                    update_message(f"   • {sch} 在第 {page} 頁")
            except Exception as _e:
                update_message(f"⚠ 無法開啟 PDF：{_e}")

        # 國小 PDF 按鈕
        if elem_schools:
            page_hint = '、'.join(f"{s}(第{p}頁)" for s, p in elem_pages.items())
            elem_btn_text = f"📄 開啟 115 國小 PDF\n{page_hint or '(請用 Ctrl+F 搜尋里名)'}"
            tk.Button(btn_frame, text=elem_btn_text,
                      command=lambda: _open_pdf(elem_pdf, '115 國小', elem_pages),
                      font=("微軟正黑體", 11),
                      bg='#4CAF50', fg='white',
                      width=30, height=3,
                      cursor='hand2',
                      justify=tk.LEFT).pack(side=tk.LEFT, padx=8)

        # 國中 PDF 按鈕
        if jun_schools:
            page_hint = '、'.join(f"{s}(第{p}頁)" for s, p in jun_pages.items())
            jun_btn_text = f"📄 開啟 115 國中 PDF\n{page_hint or '(請用 Ctrl+F 搜尋里名)'}"
            tk.Button(btn_frame, text=jun_btn_text,
                      command=lambda: _open_pdf(jun_pdf, '115 國中', jun_pages),
                      font=("微軟正黑體", 11),
                      bg='#FF9800', fg='white',
                      width=30, height=3,
                      cursor='hand2',
                      justify=tk.LEFT).pack(side=tk.LEFT, padx=8)

        # 完成按鈕
        tk.Button(dlg, text="完成（已比對 OR 不需比對）",
                  command=dlg.destroy,
                  font=("微軟正黑體", 12, "bold"),
                  bg='#2196F3', fg='white',
                  width=25, height=2,
                  cursor='hand2').pack(pady=15)

        dlg.transient(root)
        try:
            dlg.grab_set()
        except Exception:
            pass
        dlg.bind('<Escape>', lambda e: dlg.destroy())
        dlg.wait_window()

    # 查詢附近學區按鈕
    # 🔥 用 list 當「執行中旗標」，避免雙擊或自動觸發+手動按鈕同時開兩個 Chrome
    _query_running = [False]

    def query_nearby_schools():
        """查詢附近學區功能 - 先查村里，再查學區"""
        # 🔥 檢查是否已在執行中
        if _query_running[0]:
            update_message("[資訊] 查詢附近學區已在執行中，請稍候...")
            return
        _query_running[0] = True
        try:
            update_message("[資訊] 開始查詢附近學區...")

            # 取得建物門牌
            address = data.get('建物門牌', '')

            if not address:
                update_message("[錯誤] 無法取得建物門牌資料")
                show_large_message("錯誤", "無法取得建物門牌資料\n\n請先填寫「建物門牌」欄位")
                return

            # 解析門牌地址
            import re

            # 🔥 正規化：去掉重複的「縣市」前綴（曾見門牌被存成「高雄市高雄市鼓山區…」，
            #    會讓下方正則把第二個『高雄市』誤當成區、街道也跟著錯位）
            _pm = re.match(r'(.+?[縣市])', address)
            if _pm:
                _pfx = _pm.group(1)
                while address.startswith(_pfx + _pfx):
                    address = address[len(_pfx):]

            # 嘗試解析：縣市、區鄉鎮、路街、門牌號碼
            # 範例：高雄市左營區重惠街１０６號七樓
            # 🔥 區鄉鎮用 (?![鄉鎮市區]) 負向前瞻：非貪婪原本會在區名內部的行政字停下，
            #    如「前鎮區」停在『鎮』→ 區只抓到「前鎮」、剩下的『區』被推進路街變「區光華二路」
            #    →戶政司查街道查無資料。台南「新市區」也是同款(停在『市』)。
            #    加前瞻讓它一定吃到最後一個行政字(區/鄉/鎮/市)後才收尾。
            address_pattern = r'(.+?[縣市])(.+?[鄉鎮市區](?![鄉鎮市區]))(.+?[路街巷弄])(.+)'
            match = re.match(address_pattern, address)

            if not match:
                update_message("[錯誤] 無法解析建物門牌格式")
                show_large_message("錯誤", f"無法解析建物門牌格式\n\n門牌：{address}\n\n請確認門牌格式正確（例如：高雄市左營區重惠街106號）")
                return

            city = match.group(1)  # 高雄市
            district = match.group(2)  # 左營區
            road = match.group(3)  # 重惠街
            number_part = match.group(4)  # １０６號七樓

            # 提取門牌號碼（支援更細緻的格式：段、巷、弄、號、之、樓）
            # 範例：106號7樓之1、3段100號12樓、1巷2號3樓、5弄10號之2
            # 先處理包含"段"的情況
            if '段' in number_part:
                # 範例：3段100號12樓之1 -> 提取 3段100號
                segment_pattern = r'([０-９0-9]+段[０-９0-9]+號?(?:之[０-９0-9]+)?)'
                segment_match = re.match(segment_pattern, number_part)
                number = segment_match.group(1) if segment_match else number_part.split('樓')[0]
            else:
                # 一般情況：巷、弄、號、之
                number_pattern = r'([０-９0-9]+[巷弄]?[０-９0-9]*號?(?:之[０-９0-9]+)?)'
                number_match = re.match(number_pattern, number_part)
                if number_match:
                    number = number_match.group(1)
                else:
                    # 如果沒有匹配到，嘗試簡單提取數字
                    simple_match = re.match(r'([０-９0-9]+)', number_part)
                    number = simple_match.group(1) if simple_match else ""

            update_message(f"[解析] 縣市：{city}")
            update_message(f"[解析] 區鄉鎮：{district}")
            update_message(f"[解析] 路街：{road}")
            update_message(f"[解析] 完整門牌：{number_part}")
            update_message(f"[解析] 查詢用門牌：{number}")

            # 🔥 使用 Selenium 查詢村里
            update_message("[步驟 1] 使用內政部戶政司查詢村里資訊...")
            update_message("[提示] 即將開啟瀏覽器視窗，請手動輸入驗證碼")

            try:
                # 🔥 使用 python_embedded 的 Python 執行查詢（避免模組衝突）
                import subprocess
                import json
                from base_dir_helper import get_base_dir
                import os

                base_dir = get_base_dir()
                python_exe = os.path.join(base_dir, 'python_embedded', 'python.exe')
                wrapper_script = os.path.join(base_dir, '_internal', 'ris_query_wrapper.py')

                # 檢查文件是否存在
                if not os.path.exists(python_exe):
                    update_message(f"[錯誤] 找不到 python_embedded: {python_exe}")
                    show_large_message("錯誤", f"找不到 python_embedded\n{python_exe}")
                    return

                if not os.path.exists(wrapper_script):
                    update_message(f"[錯誤] 找不到查詢腳本: {wrapper_script}")
                    show_large_message("錯誤", f"找不到查詢腳本\n{wrapper_script}")
                    return

                update_message(f"[提示] 使用 python_embedded 執行查詢")
                update_message(f"[提示] Chrome 瀏覽器即將開啟，請點擊 Chrome 視窗進行操作")

                # 在新執行緒中執行查詢（避免阻塞 GUI）
                def run_query():
                    try:
                        # 調用 python_embedded 執行查詢
                        cmd = [python_exe, wrapper_script, city, district, road, number, number_part]

                        process = subprocess.Popen(
                            cmd,
                            stdout=subprocess.PIPE,
                            stderr=subprocess.PIPE,
                            text=True,
                            encoding='utf-8',
                            creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0
                        )

                        # 讀取輸出
                        result_json = None
                        capture = False
                        result_lines = []

                        for line in process.stdout:
                            line = line.rstrip()
                            if line == "===RESULT_START===":
                                capture = True
                                continue
                            elif line == "===RESULT_END===":
                                capture = False
                                result_json = ''.join(result_lines)
                                continue

                            if capture:
                                result_lines.append(line)
                            else:
                                # 輸出一般訊息
                                update_message(line)

                        process.wait()

                        # 讀取錯誤輸出
                        stderr_output = process.stderr.read()
                        if stderr_output:
                            update_message(f"[錯誤輸出] {stderr_output}")

                        # 檢查返回碼
                        if process.returncode != 0:
                            update_message(f"[錯誤] 程序返回碼：{process.returncode}")

                        # 解析結果
                        if result_json:
                            result = json.loads(result_json)
                        else:
                            result = None

                        if result:
                            village = result.get('村里', '')
                            neighbor = result.get('鄰', '')

                            update_message(f"[OK] 查詢到村里：{village}")
                            update_message(f"[OK] 查詢到鄰：第{neighbor}鄰")

                            # 步驟 2：查詢學區（支援高雄市、台南市、屏東縣）
                            # 標準化縣市名稱
                            if city == '高雄' or city == '高雄市':
                                city_normalized = '高雄市'
                            elif city == '台南' or city == '台南市' or city == '臺南' or city == '臺南市':
                                city_normalized = '台南市'
                            elif city == '屏東' or city == '屏東縣':
                                city_normalized = '屏東縣'
                            else:
                                city_normalized = city

                            if city_normalized in ['高雄市', '台南市', '屏東縣']:
                                try:
                                    from school_district_lookup import query_schools

                                    update_message(f"\n[資訊] 開始查詢{city_normalized}學區...")
                                    # 傳入縣市、村里、鄰、行政區
                                    # 🔥 一定要帶 district：同名里會跨區（例：竹西里 前鎮區 vs 路竹區），
                                    #    不帶行政區會把別區同名里的學校也撈進來（蔡文國小/路竹高中）。
                                    schools_result = query_schools(city_normalized, village, neighbor=neighbor, district=district, log_callback=update_message)

                                    # 組合學校名稱
                                    school_names = []
                                    if schools_result.get('國小'):
                                        for school in schools_result['國小']:
                                            school_names.append(school)
                                    if schools_result.get('國中'):
                                        for school in schools_result['國中']:
                                            school_names.append(school)

                                    if school_names:
                                        # 填入「鄰近學校」欄位
                                        school_text = '、'.join(school_names)
                                        school_entry.delete(0, tk.END)
                                        school_entry.insert(0, school_text)
                                        data["鄰近學校"] = school_text

                                        update_message(f"\n✓ 查詢完成！")
                                        update_message(f"[OK] 已填入鄰近學校：{school_text}")

                                        # 🔥 只有高雄市需要比對對話框（用 114 JSON + 115 PDF 提示）
                                        # 台南、屏東已直接用 115 學年 JSON，無需比對
                                        if city_normalized == '高雄市':
                                            _show_kh_result_with_compare(village, neighbor, school_names, school_text, city_normalized)
                                        else:
                                            show_large_message(
                                                "學區查詢成功",
                                                f"查詢結果：\n\n"
                                                f"村里：{village}\n"
                                                f"鄰：第{neighbor}鄰\n\n"
                                                f"學區：\n" + '\n'.join(school_names) + "\n\n"
                                                f"已自動填入「鄰近學校」欄位"
                                            )
                                    else:
                                        update_message("[提示] 未查詢到學區資料")
                                        # 根據縣市顯示不同的提示
                                        if city_normalized == '高雄市':
                                            website_tip = "請手動至高雄市教育局網站查詢"
                                        elif city_normalized == '台南市':
                                            website_tip = "請手動至台南市教育局網站查詢"
                                        elif city_normalized == '屏東縣':
                                            website_tip = "請手動至屏東縣教育處網站查詢"
                                        else:
                                            website_tip = "請手動至各縣市教育局網站查詢"

                                        show_large_message(
                                            "查詢完成",
                                            f"村里：{village}\n"
                                            f"鄰：第{neighbor}鄰\n\n"
                                            f"⚠ 學區資料中查無「{village}」的學校資訊。\n"
                                            f"可能是近期「行政區域調整」產生的新里名，\n"
                                            f"學區對照表尚未更新到這個里。\n\n"
                                            f"{website_tip}（用「{village}」查詢）"
                                        )

                                except ImportError:
                                    update_message("[錯誤] 學區查詢模組載入失敗")
                                    show_large_message(
                                        "錯誤",
                                        f"村里：{village}\n"
                                        f"鄰：第{neighbor}鄰\n\n"
                                        f"學區查詢模組載入失敗\n"
                                        f"請確認 school_district_query.py 檔案存在"
                                    )
                                except Exception as e:
                                    update_message(f"[錯誤] 學區查詢失敗：{e}")
                                    # 根據縣市顯示不同的提示
                                    if city_normalized == '高雄市':
                                        website_tip = "請手動至高雄市教育局網站查詢"
                                    elif city_normalized == '台南市':
                                        website_tip = "請手動至台南市教育局網站查詢"
                                    elif city_normalized == '屏東縣':
                                        website_tip = "請手動至屏東縣教育處網站查詢"
                                    else:
                                        website_tip = "請手動至各縣市教育局網站查詢"

                                    show_large_message(
                                        "查詢完成",
                                        f"村里：{village}\n"
                                        f"鄰：第{neighbor}鄰\n\n"
                                        f"學區查詢失敗：{e}\n"
                                        f"{website_tip}"
                                    )
                            else:
                                # 非支援縣市，僅顯示村里資訊
                                update_message(f"[提示] 目前僅支援高雄市、台南市、屏東縣學區查詢")
                                show_large_message(
                                    "查詢完成",
                                    f"村里：{village}\n"
                                    f"鄰：第{neighbor}鄰\n\n"
                                    f"目前僅支援高雄市、台南市、屏東縣學區自動查詢\n"
                                    f"請至 {city} 政府教育局網站\n"
                                    f"查詢「{village}」的國小、國中學區"
                                )
                        else:
                            update_message("[提示] 查詢失敗或已取消")

                    except Exception as e:
                        update_message(f"[錯誤] 查詢失敗：{e}")
                        show_large_message("錯誤", f"查詢失敗\n\n錯誤訊息：{e}")

                # 在新執行緒中執行（避免阻塞 GUI）
                import threading
                query_thread = threading.Thread(target=run_query, daemon=True)
                query_thread.start()

            except Exception as e:
                update_message(f"[錯誤] 啟動查詢失敗：{e}")
                show_large_message("錯誤", f"啟動查詢失敗\n\n錯誤訊息：{e}")

        except Exception as e:
            update_message(f"[錯誤] 查詢失敗：{e}")
            show_large_message("錯誤", f"查詢失敗\n\n錯誤訊息：{e}")
            import traceback
            traceback.print_exc()
        finally:
            # 🔥 不論成功或失敗，都要釋放執行中旗標
            _query_running[0] = False

    # 查詢按鈕（學區檔轉換已移除：三縣市 JSON 已內建）
    tk.Button(
        query_btn_row,
        text="🏫 查詢附近學區",
        command=query_nearby_schools,
        font=("微軟正黑體", 13),
        bg='#FF9800',
        fg='white',
        padx=15,
        pady=8,
        cursor="hand2"
    ).pack(side=tk.LEFT, padx=(250, 5))  # 左側留白對齊

    # 鄰近學校欄位（新的一行）
    school_row = tk.Frame(facility_frame)
    school_row.pack(fill=tk.X, pady=3)
    tk.Label(school_row, text="鄰近學校：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    school_entry = tk.Entry(school_row, font=entry_font, width=30)
    school_entry.insert(0, data.get("鄰近學校", ""))
    school_entry.pack(side=tk.LEFT, padx=5)
    entries["鄰近學校"] = school_entry
    bind_entry_focus(school_entry)

    # 🔥 新增：所在商圈/工業區
    create_field(facility_frame, "所在商圈/工業區：", "所在商圈", data, width=30)

    # === 土地使用資訊 ===
    land_use_frame = create_section(scrollable_frame, "🗺️ 土地使用資訊")
    # create_field(land_use_frame, "使用分區：", "使用分區", data, width=30)
    create_field(land_use_frame, "建蔽率（%）：", "建蔽率", data, width=20)
    create_field(land_use_frame, "容積率（%）：", "容積率", data, width=20)
    create_field(land_use_frame, "面臨道路（米）：", "面臨道路", data, width=20)
    create_field(land_use_frame, "面寬（米）：", "面寬", data, width=20)
    create_field(land_use_frame, "深度（米）：", "深度", data, width=20)
    # 增建面積已移至面積資訊區，此處不再重複定義

    # 🔥 新增：都市土地使用分區
    create_field(land_use_frame, "都市土地使用分區：", "都市土地使用分區", data, width=30)

    # 🔥 新增：非都市土地使用地類別
    create_field(land_use_frame, "非都市土地使用地類別：", "非都市土地使用地類別", data, width=30)

    # 🔥 新增：讀取都市計畫資料按鈕
    urban_planning_button_row = tk.Frame(land_use_frame)
    urban_planning_button_row.pack(fill=tk.X, pady=5)

    def load_urban_planning_button_click():
        """按鈕點擊事件：讀取都市計畫分區資料（手動選擇檔案）"""
        from tkinter import filedialog
        import re

        # 🔥 建立預設開啟路徑：地段號資料夾下的「4.其他相關」
        default_dir = "."
        try:
            if os.path.exists(get_data_json_path()):
                with open(get_data_json_path(), 'r', encoding='utf-8') as f:
                    data_list = json.load(f)
                if data_list and len(data_list) > 0:
                    first_data = data_list[0]
                    area = first_data.get('area', '')
                    section = first_data.get('section', '')
                    lot_number = first_data.get('lot_number', '')

                    if area and section and lot_number:
                        # 🔥 使用 get_work_folder 確保在主程式目錄下建立資料夾
                        folder_name = get_work_folder(f"{area}{section}-{lot_number}")
                        other_folder = os.path.join(folder_name, "4.其他相關")
                        if os.path.exists(other_folder):
                            default_dir = other_folder
                            print(f"[資訊] 預設開啟路徑：{default_dir}")
        except Exception as e:
            print(f"[警告] 無法取得預設路徑：{e}")

        # 🔥 讓使用者手動選擇都市計畫資料彙整.json檔案
        json_path = filedialog.askopenfilename(
            title="選擇都市計畫資料彙整.json檔案",
            filetypes=[("JSON檔案", "*.json"), ("所有檔案", "*.*")],
            initialdir=default_dir
        )

        if not json_path:
            print("[提示] 使用者取消選擇檔案")
            return

        print(f"[OK] 使用者選擇檔案：{json_path}")

        try:
            # 讀取 JSON 檔案
            with open(json_path, 'r', encoding='utf-8') as f:
                combined_data = json.load(f)

            # 取得兩類資料
            zone_data_list = combined_data.get('分區查詢資料', [])
            urban_plan_list = combined_data.get('都市計畫圖街廓資訊', [])

            print(f"[OK] 讀取到 {len(zone_data_list)} 筆分區查詢資料")
            print(f"[OK] 讀取到 {len(urban_plan_list)} 筆都市計畫圖街廓資訊")

            # 🔥 手動選擇檔案時，不做地段地號比對，直接載入所有資料
            print(f"[提示] 手動選擇檔案，將載入檔案中的所有資料")

            # 處理所有資料（分區查詢資料 + 都市計畫圖街廓資訊）
            all_data = zone_data_list + urban_plan_list
            all_zones = set()
            all_building_coverage = set()
            all_floor_area = set()
            all_details = []

            for data_dict in all_data:
                try:
                    # 過濾空資料
                    json_section = data_dict.get('地段', '').strip()
                    json_lot = data_dict.get('地號', '').strip()

                    if not json_section and not json_lot:
                        print(f"     [跳過] 空資料")
                        continue

                    print(f"     [✓ 載入] 地段='{json_section}' 地號='{json_lot}'")

                    # 收集使用分區
                    zone = data_dict.get('使用分區', '') or data_dict.get('使用分區完整名稱', '') or data_dict.get('使用分區別名', '')
                    if zone:
                        all_zones.add(zone)
                        print(f"        使用分區={zone}")

                    # 收集建蔽率
                    building_coverage = data_dict.get('建蔽率', '')
                    if building_coverage:
                        cleaned_value = re.sub(r'[^\d.]', '', str(building_coverage))
                        if cleaned_value:
                            all_building_coverage.add(cleaned_value)
                            print(f"        建蔽率={cleaned_value}")

                    # 收集容積率
                    floor_area = data_dict.get('容積率', '')
                    if floor_area:
                        cleaned_value = re.sub(r'[^\d.]', '', str(floor_area))
                        if cleaned_value:
                            all_floor_area.add(cleaned_value)
                            print(f"        容積率={cleaned_value}")

                    all_details.append(data_dict)

                except Exception as e:
                    print(f"[警告] 處理資料時失敗：{e}")
                    continue

            if not all_details:
                show_large_message("提示", "找不到符合條件的都市計畫資料\n\n可能原因：\n1. 土地標示中的地段地號與JSON檔案不符\n2. JSON檔案格式不正確")
                return

            # 整理結果
            planning_data = {
                'zones': sorted(list(all_zones)),
                'building_coverage_ratios': sorted(list(all_building_coverage), key=lambda x: float(x) if x else 0),
                'floor_area_ratios': sorted(list(all_floor_area), key=lambda x: float(x) if x else 0),
                'details': all_details
            }

            print(f"[OK] 整理完成（共 {len(all_details)} 筆匹配資料）")
            print(f"     使用分區 ({len(planning_data['zones'])} 種): {', '.join(planning_data['zones'])}")
            print(f"     建蔽率 ({len(planning_data['building_coverage_ratios'])} 種): {', '.join(planning_data['building_coverage_ratios'])}")
            print(f"     容積率 ({len(planning_data['floor_area_ratios'])} 種): {', '.join(planning_data['floor_area_ratios'])}")

        except Exception as e:
            show_large_message("錯誤", f"讀取檔案失敗：{e}")
            print(f"[錯誤] 讀取都市計畫資料失敗：{e}")
            import traceback
            traceback.print_exc()
            return

        if planning_data:
            # 將使用分區（去重後用頓號分隔）填入「都市土地使用分區」欄位
            # 🔥 若為「都市計畫外」則不帶入
            zones_text = '、'.join(planning_data['zones'])
            if zones_text and zones_text != '都市計畫外' and '都市土地使用分區' in entries:
                entries['都市土地使用分區'].delete(0, tk.END)
                entries['都市土地使用分區'].insert(0, zones_text)

            # 🔥 將建蔽率（去重後用頓號分隔）填入「建蔽率」欄位
            building_coverage_text = '、'.join(planning_data.get('building_coverage_ratios', []))
            if building_coverage_text and '建蔽率' in entries:
                entries['建蔽率'].delete(0, tk.END)
                entries['建蔽率'].insert(0, building_coverage_text)

            # 🔥 將容積率（去重後用頓號分隔）填入「容積率」欄位
            floor_area_text = '、'.join(planning_data.get('floor_area_ratios', []))
            if floor_area_text and '容積率' in entries:
                entries['容積率'].delete(0, tk.END)
                entries['容積率'].insert(0, floor_area_text)

            # 🔥 顯示成功訊息（包含當前土地標示和詳細比對資訊）
            # 從 data 中解析地段地號，用於比對
            target_lots = parse_land_lots_from_data(data)

            message = "=" * 60 + "\n"
            message += "📋 當前物件的土地標示\n"
            message += "=" * 60 + "\n"
            message += f"土地標示：{data.get('土地標示', '(無)')}\n"
            if target_lots:
                message += f"解析結果：\n"
                for idx, lot in enumerate(target_lots, 1):
                    message += f"  {idx}. 地段='{lot.get('地段', '')}' 地號='{lot.get('地號', '')}'\n"
            else:
                message += "解析結果：(無法解析)\n"
            message += "\n"

            message += "=" * 60 + "\n"
            message += "📊 載入的都市計畫資料\n"
            message += "=" * 60 + "\n"
            message += f"使用分區（已去重）：{zones_text}\n"
            message += f"建蔽率（已去重）：{building_coverage_text}\n"
            message += f"容積率（已去重）：{floor_area_text}\n\n"
            message += f"共載入 {len(planning_data['details'])} 筆都市計畫資料\n\n"

            for idx, detail in enumerate(planning_data['details'], 1):
                json_section = detail.get('地段', '')
                json_lot = detail.get('地號', '')

                # 🔥 檢查是否匹配
                is_matched = False
                if target_lots:
                    for lot in target_lots:
                        target_section = lot.get('地段', '')
                        target_lot = lot.get('地號', '')
                        # 簡單的標準化比對（移除前導零）
                        if target_section == json_section:
                            target_lot_normalized = target_lot.lstrip('0').replace('-0000', '').replace('-0', '-') or '0'
                            json_lot_normalized = json_lot.lstrip('0').replace('-0000', '').replace('-0', '-') or '0'
                            if target_lot_normalized == json_lot_normalized:
                                is_matched = True
                                break

                match_icon = "✓ 匹配" if is_matched else "○ 未匹配"
                message += f"【{match_icon}】資料 {idx}：\n"
                message += f"  地段：{json_section}\n"
                message += f"  地號：{json_lot}\n"
                message += f"  使用分區：{detail.get('使用分區', '') or detail.get('使用分區完整名稱', '') or detail.get('使用分區別名', '')}\n"
                if detail.get('建蔽率', ''):
                    message += f"  建蔽率：{detail.get('建蔽率', '')}\n"
                if detail.get('容積率', ''):
                    message += f"  容積率：{detail.get('容積率', '')}\n"
                if detail.get('土地面積', ''):
                    message += f"  土地面積：{detail.get('土地面積', '')}\n"
                if detail.get('公告現值', ''):
                    message += f"  公告現值：{detail.get('公告現值', '')}\n"
                if detail.get('公告地價', ''):
                    message += f"  公告地價：{detail.get('公告地價', '')}\n"
                if detail.get('附註', ''):
                    message += f"  附註：{detail.get('附註', '')}\n"
                message += "\n"

            show_large_message("讀取成功", message)
        else:
            show_large_message("提示", "找不到都市計畫資料\n\n請確認：\n1. 已執行「高雄都市計畫」程式\n2. JSON 檔案位於「4.其他相關」資料夾\n3. 檔案格式：分區查詢資料-*.json 或 都市計畫圖街廓資訊-*.json")

    # 置中按鈕
    tk.Label(urban_planning_button_row, text="", width=30).pack(side=tk.LEFT)
    tk.Button(
        urban_planning_button_row,
        text="讀取都市計畫分區資料",
        font=("微軟正黑體", 10),
        command=load_urban_planning_button_click,
        bg="#4A90E2",
        fg="white",
        padx=10,
        pady=5
    ).pack(side=tk.LEFT, padx=5)
  
    # === 訴求重點（文字框）===
    selling_frame = create_section(scrollable_frame, "✨ 訴求重點")
    
    text_row = tk.Frame(selling_frame)
    text_row.pack(fill=tk.X, pady=3)
    tk.Label(text_row, text="訴求重點：", font=label_font, width=15, anchor="ne").pack(side=tk.LEFT, padx=5)
    
    selling_text = tk.Text(text_row, font=entry_font, width=60, height=5, wrap=tk.WORD)
    selling_text.insert("1.0", data.get("訴求重點", ""))
    selling_text.pack(side=tk.LEFT, padx=5)
    selling_text.bind('<Button-3>', _show_text_menu)  # 🔥 右鍵文字選單
    entries["訴求重點"] = selling_text

    # === ✨ Google AI 生成行銷文（依物件資料自動撰寫，帶回上方訴求重點）===
    def _ai_gather_fields():
        """擷取對行銷有用的欄位，回傳 (是否租件, 資料文字)。"""
        is_rent = (transaction_var.get() == "租件")

        def val(k):
            w = entries.get(k)
            if w is None or isinstance(w, dict):
                return ""
            try:
                if isinstance(w, tk.Text):
                    return w.get("1.0", "end-1c").strip()
                return str(w.get()).strip()
            except Exception:
                return ""

        def is_on(k):
            w = entries.get(k)
            try:
                return bool(w.get())
            except Exception:
                return False

        lines = []
        # 格局（房/廳/衛 合併）
        g = ""
        for f in ("房", "廳", "衛"):
            v = val("格局_" + f)
            if v and v != "0":
                g += f"{v}{f}"
        if g:
            lines.append(f"格局：{g}")
        # 車位
        if val("車位_有無"):
            types = []
            if is_on("車位_型式_平面"):
                types.append("平面")
            if is_on("車位_型式_機械"):
                types.append("機械")
            cp = f"車位：{val('車位_有無')}"
            if types:
                cp += "（" + "、".join(types) + "）"
            lines.append(cp)
        # 一般有用欄位（白名單，排除編號/稅賦/經緯度/後台等無關行銷者）
        COMMON = ["案名", "建物門牌", "土地標示", "所在商圈", "用途", "使用分區",
                  "都市土地使用分區", "非都市土地使用地類別", "總建坪", "主建物坪數",
                  "附屬建物坪數", "公共設施坪數", "總地坪", "增建面積", "地上層數",
                  "地下層數", "電梯", "電梯數量", "建築結構", "竣工日期", "座向",
                  "面寬", "深度", "面臨道路", "現況", "現況說明", "鄰近學校",
                  "管理費", "工廠登記", "廠房保存登記", "天車", "型式"]
        SALE = ["總價", "含車位價格", "土地單價", "建物單價"]
        RENT = ["租金", "押金", "租期", "調幅", "履約保證", "租約公證", "整地期"]
        # 為易誤解的欄位加上更清楚的標籤，避免 AI 把土地/建物搞混或誤判預售
        LABELS = {
            "總地坪": "土地坪數（地號基地總面積，區分所有建物本戶為持份）",
            "總建坪": "建物坪數（總建坪）",
            "主建物坪數": "主建物坪數",
            "竣工日期": "建築完成日期（已完工成屋）",
        }
        for k in COMMON + (RENT if is_rent else SALE):
            v = val(k)
            if v and v not in ("無", "沒有", "0", "False", "0.0"):
                lines.append(f"{LABELS.get(k, k)}：{v}")
        return is_rent, "\n".join(lines)

    def _ask_gemini_key(current=""):
        """大字級對話框輸入 Gemini 金鑰，含可點擊自動開啟的免費申請連結。回傳金鑰或 None。"""
        import webbrowser
        parent = selling_text.winfo_toplevel()
        win = tk.Toplevel(parent)
        win.title("設定／更換 Google Gemini API 金鑰")
        win.configure(bg="white")
        win.resizable(False, False)
        win.transient(parent)
        result = {"key": None}
        F = ("Microsoft JhengHei", 13)
        FB = ("Microsoft JhengHei", 13, "bold")

        frm = tk.Frame(win, bg="white")
        frm.pack(fill="both", expand=True, padx=26, pady=22)
        tk.Label(frm, text="請貼上 Google Gemini API 金鑰",
                 font=FB, bg="white", justify="left").pack(anchor="w")
        tk.Label(frm, text="（設定一次會自動記住；若免費額度用完，可在這裡改貼另一個 Google 帳號的金鑰繼續使用）",
                 font=F, bg="white", fg="#555", justify="left", wraplength=560).pack(anchor="w", pady=(0, 10))

        link_row = tk.Frame(frm, bg="white")
        link_row.pack(anchor="w")
        tk.Label(link_row, text="免費申請（點我開啟）：", font=F, bg="white").pack(side="left")
        url = "https://aistudio.google.com/apikey"
        link = tk.Label(link_row, text=url, font=("Microsoft JhengHei", 13, "underline"),
                        fg="#1A73E8", bg="white", cursor="hand2")
        link.pack(side="left")
        link.bind("<Button-1>", lambda e: webbrowser.open(url))
        tk.Label(frm, text="登入後按「Create API key」，複製整串貼到下面：",
                 font=("Microsoft JhengHei", 11), fg="#888", bg="white",
                 justify="left").pack(anchor="w", pady=(8, 2))

        ent = tk.Entry(frm, font=F, width=48, relief="solid", bd=1)
        ent.pack(fill="x", pady=(2, 6), ipady=4)
        if current:
            ent.insert(0, current)  # 更換金鑰時預填現有值
        ent.focus_set()

        btn_row = tk.Frame(frm, bg="white")
        btn_row.pack(fill="x", pady=(12, 0))

        def _ok():
            result["key"] = ent.get().strip()
            win.destroy()

        def _cancel():
            result["key"] = None
            win.destroy()

        tk.Button(btn_row, text="確定", font=FB, bg="#4285F4", fg="white", width=8,
                  relief="flat", cursor="hand2", padx=10, pady=4, command=_ok).pack(side="right", padx=4)
        tk.Button(btn_row, text="取消", font=F, bg="#E0E0E0", fg="black", width=8,
                  relief="flat", cursor="hand2", padx=10, pady=4, command=_cancel).pack(side="right", padx=4)
        ent.bind("<Return>", lambda e: _ok())
        win.bind("<Escape>", lambda e: _cancel())

        win.update_idletasks()
        w, h = win.winfo_reqwidth(), win.winfo_reqheight()
        try:
            x = parent.winfo_rootx() + (parent.winfo_width() - w) // 2
            y = parent.winfo_rooty() + (parent.winfo_height() - h) // 2
        except Exception:
            x, y = 200, 200
        win.geometry(f"+{max(0, x)}+{max(0, y)}")
        win.grab_set()
        win.wait_window()
        return result["key"]

    def _edit_prompt_dialog(initial):
        """顯示可編輯的提示詞；按「送出」回傳文字、按「取消」回傳 None。"""
        parent = selling_text.winfo_toplevel()
        win = tk.Toplevel(parent)
        win.title("確認／修改提示詞")
        win.configure(bg="white")
        win.transient(parent)
        result = {"text": None}
        F = ("Microsoft JhengHei", 12)
        FB = ("Microsoft JhengHei", 12, "bold")

        frm = tk.Frame(win, bg="white")
        frm.pack(fill="both", expand=True, padx=20, pady=16)
        tk.Label(frm, text="以下是要送給 Google AI 的提示詞，可自行修改後再送出：",
                 font=FB, bg="white", justify="left").pack(anchor="w", pady=(0, 8))

        txt_wrap = tk.Frame(frm, bg="white")
        txt_wrap.pack(fill="both", expand=True)
        scroll = tk.Scrollbar(txt_wrap)
        scroll.pack(side="right", fill="y")
        txt = tk.Text(txt_wrap, font=F, width=74, height=20, wrap="word",
                      relief="solid", bd=1, yscrollcommand=scroll.set)
        txt.pack(side="left", fill="both", expand=True)
        scroll.config(command=txt.yview)
        txt.insert("1.0", initial)
        txt.focus_set()

        btn_row = tk.Frame(frm, bg="white")
        btn_row.pack(fill="x", pady=(12, 0))

        def _send():
            result["text"] = txt.get("1.0", "end-1c").strip()
            win.destroy()

        def _cancel():
            result["text"] = None
            win.destroy()

        tk.Button(btn_row, text="✨ 送出給 AI 生成", font=FB, bg="#4285F4", fg="white",
                  relief="flat", cursor="hand2", padx=14, pady=5, command=_send).pack(side="right", padx=4)
        tk.Button(btn_row, text="取消", font=F, bg="#E0E0E0", fg="black",
                  relief="flat", cursor="hand2", padx=14, pady=5, command=_cancel).pack(side="right", padx=4)
        win.bind("<Escape>", lambda e: _cancel())

        win.update_idletasks()
        w, h = win.winfo_reqwidth(), win.winfo_reqheight()
        try:
            x = parent.winfo_rootx() + (parent.winfo_width() - w) // 2
            y = parent.winfo_rooty() + (parent.winfo_height() - h) // 2
        except Exception:
            x, y = 150, 100
        win.geometry(f"+{max(0, x)}+{max(0, y)}")
        win.grab_set()
        win.wait_window()
        return result["text"]

    def _ai_generate():
        cfg = load_config() or {}
        api_key = (cfg.get("gemini_api_key") or "").strip()
        if not api_key:
            api_key = (_ask_gemini_key() or "").strip()
            if not api_key:
                return
            cfg["gemini_api_key"] = api_key
            try:
                save_config(cfg, show_message=False)
            except Exception:
                pass
        model = (cfg.get("gemini_model") or "gemini-2.5-flash").strip()

        is_rent, data_text = _ai_gather_fields()
        if not data_text.strip():
            messagebox.showwarning("沒有資料", "請先填寫物件欄位，再生成行銷文。")
            return
        mode = "租賃" if is_rent else "銷售"
        # 指示規則可由 config.json 的 gemini_instructions 覆寫（用 {mode} 代入銷售/租賃），
        # 方便不重新打包就調整文案風格。
        DEFAULT_RULES = (
            "請以「{mode}」角度撰寫繁體中文房地產行銷文，務必遵守：\n"
            "1. 此為已完工成屋，一律用現在式描述，禁止使用「預計」「將落成」等未來語氣。\n"
            "2. 正確區分坪數：建物坪數與土地坪數不可混淆；若為公寓／華廈/大樓等區分所有建物，"
            "土地屬「持份」，不可把地號基地總面積說成本戶獨有土地。\n"
            "3. 內容分點涵蓋：本案亮點、附近環境優點(生活機能、學區、公園、山海景觀)、"
            "重大建設(如捷運、國道、百貨、商場)、交通動線(到主要地點開車或步行幾分鐘、距離多遠)、"
            "格局採光通風(可看建物測量成果圖、土地可看地籍圖)、周邊同類產品行情(查實價登錄)；"
            "其中環境／建設／交通／行情若無確切資料，僅作概括描述，不要捏造具體建案名稱、路名或數字。\n"
            "4. 全文共約4~5點，以「一、二、三、…」分點條列，每一點不超過36個中文字。\n"
            "5. 只輸出分點內容，不要前言、標題或結語。"
        )
        rules = (cfg.get("gemini_instructions") or DEFAULT_RULES).replace("{mode}", mode)
        # 🔥 若「訴求重點」欄位已有內容，一併納入並請 AI 加強修飾後融入
        existing = selling_text.get("1.0", "end-1c").strip()
        _parts = [f"你是專業房地產文案。以下為物件資料（皆為事實，未提供的數字請勿自行捏造）：\n{data_text}"]
        if existing:
            _parts.append("「訴求重點」欄位已有以下內容，請將其重點納入並加強修飾後融入行銷文：\n" + existing)
        _parts.append(rules)
        prompt = "\n\n".join(_parts)

        # 🔥 先讓使用者確認／修改提示詞，按「送出」才生成、「取消」則不動作
        prompt = _edit_prompt_dialog(prompt)
        if not prompt or not prompt.strip():
            return

        ai_btn.config(state=tk.DISABLED, text="✨ 生成中…請稍候")
        # 🔥 把進度印到主控台，讓使用者知道正在跑（原本只有按鈕小字，不易察覺）
        try:
            update_message(f"[AI] 🚀 已送出，正在呼叫 Google Gemini（{model}）生成行銷文，約需 10~40 秒，請稍候…")
        except Exception:
            pass

        def worker():
            msg = None
            text = ""
            offer_key = False  # 失敗時是否提議更換金鑰（額度滿／金鑰無效時）
            try:
                import requests  # 確保此執行緒作用域可用（模組頂層雖有，仍明確 import 保險）
                url = (f"https://generativelanguage.googleapis.com/v1beta/models/"
                       f"{model}:generateContent?key={api_key}")
                payload = {"contents": [{"parts": [{"text": prompt}]}]}
                # 🔥 gemini-2.5-flash 是「思考模型」，關閉思考可加快回應，
                #    並避免『思考佔滿輸出額度→回傳沒有文字(parts)』導致沒回填也沒訊息。
                if model.startswith("gemini-2.5-flash"):
                    payload["generationConfig"] = {"thinkingConfig": {"thinkingBudget": 0}}
                resp = requests.post(url, json=payload, timeout=90)
                if resp.status_code == 200:
                    j = resp.json()
                    # 🔥 穩健解析：合併所有 parts 的文字；若無文字給清楚原因(被擋/截斷)
                    cands = j.get("candidates") or []
                    if cands:
                        parts = ((cands[0].get("content") or {}).get("parts")) or []
                        text = "".join(p.get("text", "") for p in parts).strip()
                        if not text:
                            fr = cands[0].get("finishReason", "") or "未知"
                            msg = (f"AI 沒有回傳文字（finishReason={fr}）。\n"
                                   "可能被安全機制擋下、或內容過長被截斷，\n請調整提示詞後再試一次。")
                    else:
                        pf = j.get("promptFeedback", {}) or {}
                        br = pf.get("blockReason", "") or "無候選內容"
                        msg = f"AI 未產生內容（{br}）。\n請調整提示詞後再試。"
                else:
                    try:
                        ej = resp.json().get("error", {})
                        status = ej.get("status", "")
                        emsg = ej.get("message", "")
                    except Exception:
                        status, emsg = "", (resp.text or "")[:200]
                    if resp.status_code == 429 or status == "RESOURCE_EXHAUSTED":
                        offer_key = True
                        msg = ("⚠️ 免費額度可能已用完，或呼叫太頻繁。\n\n"
                               "可過幾分鐘（或明天額度重置後）再試。")
                    elif resp.status_code in (400, 401, 403) and (
                            "api key" in emsg.lower() or "API_KEY" in status
                            or status in ("INVALID_ARGUMENT", "PERMISSION_DENIED", "UNAUTHENTICATED")):
                        offer_key = True
                        msg = "⚠️ API 金鑰無效或未授權。"
                    else:
                        msg = f"生成失敗（{resp.status_code} {status}）：\n{emsg[:150]}"
            except Exception as e:
                msg = f"連線失敗：{e}\n\n請確認網路是否正常。"

            def done():
                ai_btn.config(state=tk.NORMAL, text="✨ Google AI 生成行銷文")
                if msg or not text:
                    try:
                        update_message(f"[AI] ❌ 生成失敗：{(msg or '未取得生成內容').splitlines()[0]}")
                    except Exception:
                        pass
                    # 額度滿／金鑰無效時，才詢問是否更換金鑰（需要時才彈出）
                    if offer_key and messagebox.askyesno(
                            "生成失敗", (msg or "") + "\n\n要現在更換 API 金鑰嗎？"):
                        _change_gemini_key()
                    elif not offer_key:
                        messagebox.showerror("生成失敗", msg or "未取得生成內容，請再試一次。")
                else:
                    selling_text.delete("1.0", "end")
                    selling_text.insert("1.0", text)
                    try:
                        update_message(f"[AI] ✅ 生成完成，已回填「訴求重點」（{len(text)} 字）")
                    except Exception:
                        pass
                    try:
                        messagebox.showinfo("生成完成", "已將 AI 行銷文回填到「訴求重點」欄位，請自行核對後再使用。")
                    except Exception:
                        pass
            selling_text.after(0, done)

        threading.Thread(target=worker, daemon=True).start()

    def _change_gemini_key():
        cfg = load_config() or {}
        new = _ask_gemini_key(current=(cfg.get("gemini_api_key") or ""))
        if new is None:   # 按取消
            return
        cfg["gemini_api_key"] = new.strip()
        try:
            save_config(cfg, show_message=False)
        except Exception:
            pass
        messagebox.showinfo("完成", "API 金鑰已更新。" if new.strip() else "已清除 API 金鑰。")

    ai_btn_row = tk.Frame(selling_frame)
    ai_btn_row.pack(fill=tk.X, pady=(0, 4))
    tk.Label(ai_btn_row, text="", font=label_font, width=15).pack(side=tk.LEFT, padx=5)
    ai_btn = tk.Button(ai_btn_row, text="✨ Google AI 生成行銷文",
                       command=_ai_generate, bg="#4285F4", fg="white",
                       font=label_font, cursor="hand2", relief=tk.RAISED, padx=10)
    ai_btn.pack(side=tk.LEFT, padx=5)
    tk.Label(ai_btn_row,
             text="⚠️ AI 生成內容僅供參考，務必自行核對（行情、重大建設、交通等多為 AI 推測，可能與事實不符）",
             font=label_font, fg="#C0392B").pack(side=tk.LEFT, padx=8)

    # === 另顯示於後台內容（文字框）===
    backend_text_row = tk.Frame(selling_frame)
    backend_text_row.pack(fill=tk.X, pady=3)
    tk.Label(backend_text_row, text="另顯示於後台內容：", font=label_font, width=15, anchor="ne").pack(side=tk.LEFT, padx=5)

    backend_text = tk.Text(backend_text_row, font=entry_font, width=60, height=5, wrap=tk.WORD)
    backend_text.insert("1.0", data.get("另顯示於後台內容", ""))
    backend_text.pack(side=tk.LEFT, padx=5)
    backend_text.bind('<Button-3>', _show_text_menu)  # 🔥 右鍵文字選單
    entries["另顯示於後台內容"] = backend_text

    # === 其他條件（租件專用）===
    other_conditions_frame = create_section(scrollable_frame, "📋 其他條件")

    # 🔥 租約公證（租件專用，下拉選單）
    notarization_row = tk.Frame(other_conditions_frame)
    notarization_row.pack(fill=tk.X, pady=3)
    tk.Label(notarization_row, text="租約公證：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    notarization_var = tk.StringVar(value=data.get("租約公證", ""))
    notarization_combo = ttk.Combobox(notarization_row, textvariable=notarization_var, font=combo_font, width=15,
                                      values=["需要", "不需要"])
    notarization_combo.pack(side=tk.LEFT, padx=5)
    entries["租約公證"] = notarization_var

    # 🔥 起造人名義（租件專用，下拉選單）
    builder_row = tk.Frame(other_conditions_frame)
    builder_row.pack(fill=tk.X, pady=3)
    tk.Label(builder_row, text="起造人名義：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)
    builder_var = tk.StringVar(value=data.get("起造人名義", ""))
    builder_combo = ttk.Combobox(builder_row, textvariable=builder_var, font=combo_font, width=15,
                                 values=["地主", "承租方"])
    builder_combo.pack(side=tk.LEFT, padx=5)
    entries["起造人名義"] = builder_var

    # 🔥 適合行業別（租件專用，複選框）
    industry_label_row = tk.Frame(other_conditions_frame)
    industry_label_row.pack(fill=tk.X, pady=3)
    tk.Label(industry_label_row, text="適合行業別：", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)

    # 全選/取消選擇按鈕
    tk.Button(industry_label_row, text="全選", font=("微軟正黑體", 9),
              command=lambda: select_all_industries(True)).pack(side=tk.LEFT, padx=5)
    tk.Button(industry_label_row, text="取消選擇", font=("微軟正黑體", 9),
              command=lambda: select_all_industries(False)).pack(side=tk.LEFT, padx=5)

    # 取得已儲存的適合行業別資料
    適合行業別_data = data.get("適合行業別", {})
    if isinstance(適合行業別_data, str):  # 相容舊版本（字串格式）
        適合行業別_data = {}

    適合行業別_vars = {}
    行業選項 = ["住宅", "餐飲業", "美容業", "服飾業", "娛樂休閒", "電競電玩", "運動健身",
                "百貨超市", "生活服務", "電器通訊", "大賣場", "電商業", "倉庫", "物流業"]

    # 第一行：前7個選項
    industry_row1 = tk.Frame(other_conditions_frame)
    industry_row1.pack(fill=tk.X, pady=3)
    tk.Label(industry_row1, text="", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)  # 空白對齊

    for option in 行業選項[:7]:
        var = tk.BooleanVar(value=適合行業別_data.get(option, False))
        cb = tk.Checkbutton(industry_row1, text=option, variable=var, font=entry_font)
        cb.pack(side=tk.LEFT, padx=5)
        適合行業別_vars[option] = var

    # 第二行：後7個選項
    industry_row2 = tk.Frame(other_conditions_frame)
    industry_row2.pack(fill=tk.X, pady=3)
    tk.Label(industry_row2, text="", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)  # 空白對齊

    for option in 行業選項[7:]:
        var = tk.BooleanVar(value=適合行業別_data.get(option, False))
        cb = tk.Checkbutton(industry_row2, text=option, variable=var, font=entry_font)
        cb.pack(side=tk.LEFT, padx=5)
        適合行業別_vars[option] = var

    # 第三行：其他（自訂輸入）
    industry_row3 = tk.Frame(other_conditions_frame)
    industry_row3.pack(fill=tk.X, pady=3)
    tk.Label(industry_row3, text="", font=label_font, width=30, anchor="e").pack(side=tk.LEFT, padx=5)  # 空白對齊

    其他_var = tk.BooleanVar(value=適合行業別_data.get("其他_checked", False))
    其他_cb = tk.Checkbutton(industry_row3, text="其他：", variable=其他_var, font=entry_font)
    其他_cb.pack(side=tk.LEFT, padx=5)
    適合行業別_vars["其他_checked"] = 其他_var

    其他_entry = tk.Entry(industry_row3, font=entry_font, width=30)
    其他_entry.insert(0, 適合行業別_data.get("其他_text", ""))
    其他_entry.pack(side=tk.LEFT, padx=5)
    適合行業別_vars["其他_text"] = 其他_entry

    entries["適合行業別"] = 適合行業別_vars

    # 定義全選/取消選擇函數
    def select_all_industries(select: bool):
        """全選或取消選擇所有行業別（不包含其他）"""
        for option in 行業選項:
            if option in 適合行業別_vars:
                適合行業別_vars[option].set(select)

    # 🔥🔥🔥 在所有相關變數定義完成後，定義交易類型切換函數 🔥🔥🔥
    def toggle_transaction_type_with_parking():
        """根據交易類型顯示/隱藏對應欄位，並更新車位選單"""
        transaction_type = transaction_var.get()

        if transaction_type == "售件":
            # 🔥 設定售件的視覺樣式（淺藍色系）
            editor_window.title("🔵 物件資料編輯器 - 售件 🔵")
            scrollable_frame.config(bg="#E3F2FD")  # 淺藍色背景
            canvas.config(bg="#E3F2FD")
            main_frame.config(bg="#E3F2FD")

            # 🔥 更新所有區塊標題的顏色為淺藍色
            for title_label in section_titles:
                title_label.config(bg="#64B5F6", fg="black")  # 淺藍色標題，黑色文字

            # 顯示售件欄位
            total_price_container.pack(fill=tk.X, pady=3)
            sale_fields_frame.pack(fill=tk.X)
            # 顯示土地增值稅和貸款資訊（售件專用）- 插入在建物資訊之前
            land_value_tax_frame.master.pack(fill=tk.X, pady=5, before=building_frame.master)
            loan_frame.master.pack(fill=tk.X, pady=5, before=building_frame.master)
            # 隱藏租件欄位
            rent_fields_frame.pack_forget()
            # 🔥 隱藏其他條件區塊（租件專用）
            other_conditions_frame.master.pack_forget()
            # 🔥 隱藏稅賦配置區塊（租件專用）
            tax_frame.master.pack_forget()
            # 🔥 更新車位選單為售件選項
            park_combo['values'] = ["有，公設內", "有，獨立產權", "無"]
            current_val = park_var.get()
            if current_val == "有":
                park_var.set("有，公設內")
            # 🔥 顯示售件才需要的其他設施項目
            # 使用 after= 確保插回原始位置，避免 pack_forget 後 pack 會插到末尾
            guarantee_row.pack(fill=tk.X, pady=3, after=管理費_row)
            factory_reg_row.pack(fill=tk.X, pady=3, after=guarantee_row)
            division_row.pack(fill=tk.X, pady=3, after=factory_row)
            crane_row.pack(fill=tk.X, pady=3, after=division_row)
            樑下_row.pack(fill=tk.X, pady=3, after=crane_row)
        else:  # 租件
            # 🔥 設定租件的視覺樣式（淺紅色系）
            editor_window.title("🔴 物件資料編輯器 - 租件 🔴")
            scrollable_frame.config(bg="#FFEBEE")  # 淺紅色背景
            canvas.config(bg="#FFEBEE")
            main_frame.config(bg="#FFEBEE")

            # 🔥 更新所有區塊標題的顏色為淺紅色
            for title_label in section_titles:
                title_label.config(bg="#EF9A9A", fg="black")  # 淺紅色標題，黑色文字

            # 隱藏售件專屬欄位
            total_price_container.pack_forget()
            sale_fields_frame.pack_forget()
            # 隱藏土地增值稅和貸款資訊（租件不需要）
            land_value_tax_frame.master.pack_forget()
            loan_frame.master.pack_forget()
            # 顯示租件欄位
            rent_fields_frame.pack(fill=tk.X)
            # 🔥 顯示其他條件和稅賦配置區塊（租件專用）- 插入在訴求重點之前
            other_conditions_frame.master.pack(fill=tk.X, pady=5, before=selling_frame.master)
            tax_frame.master.pack(fill=tk.X, pady=5, before=selling_frame.master)
            # 🔥 更新車位選單為租件選項（只有「有」和「無」）
            park_combo['values'] = ["有", "無"]
            current_val = park_var.get()
            if current_val in ["有，公設內", "有，獨立產權"]:
                park_var.set("有")
            # 🔥 隱藏租件不需要的其他設施項目（履約保證、廠房保存登記、分割銷售、天車、樑下）
            # 🔥 注意：管理費已統一在租件價格資訊區，會隨 rent_fields_frame 顯示/隱藏
            # 🔥 中脊在租件中需要顯示
            guarantee_row.pack_forget()
            factory_reg_row.pack_forget()
            division_row.pack_forget()
            crane_row.pack_forget()
            樑下_row.pack_forget()
            # 中脊不隱藏，租件需要

    # 🔥 綁定交易類型切換函數
    transaction_var.trace_add('write', lambda *args: toggle_transaction_type_with_parking())
    # 初始化
    toggle_transaction_type_with_parking()

    # ========================================
    # 🔥 修復 Windows IME 輸入法焦點問題
    # ========================================
    # 這個修復會在視窗載入後強制觸發一次焦點循環，解決中文輸入法卡住的問題
    def fix_ime_focus():
        """修復 Windows IME 焦點問題"""
        try:
            # 方法1：強制觸發一次焦點切換
            editor_window.focus_force()
            editor_window.update_idletasks()

            # 方法2：找到第一個 Entry 控制項並設定焦點
            for widget in scrollable_frame.winfo_children():
                if isinstance(widget, tk.Frame):
                    for child in widget.winfo_children():
                        if isinstance(child, tk.Frame):
                            for subchild in child.winfo_children():
                                if isinstance(subchild, tk.Entry):
                                    subchild.focus_set()
                                    subchild.icursor(tk.END)
                                    return
        except:
            pass

    # 延遲執行修復（等待視窗完全載入）
    editor_window.after(100, fix_ime_focus)

    # ========================================
    # === 🔥 位置圖區塊（完整修正版）===
    # ========================================
    
    # === 位置圖區塊 ===
    map_frame = create_section(scrollable_frame, "🗺️ 位置圖")

    # 第1列：經緯度輸入（不包含按鈕）
    coord_row = tk.Frame(map_frame)
    coord_row.pack(fill=tk.X, pady=3)
    
    tk.Label(coord_row, text="緯度：", font=label_font, width=20, anchor="e").pack(side=tk.LEFT, padx=5)
    lat_entry = tk.Entry(coord_row, font=entry_font, width=10)
    lat_entry.insert(0, data.get('緯度', ''))
    lat_entry.pack(side=tk.LEFT, padx=5)
    entries["緯度"] = lat_entry

    tk.Label(coord_row, text="經度：", font=entry_font).pack(side=tk.LEFT, padx=5)
    lng_entry = tk.Entry(coord_row, font=entry_font, width=10)
    lng_entry.insert(0, data.get('經度', ''))
    lng_entry.pack(side=tk.LEFT, padx=5)
    entries["經度"] = lng_entry
    
    # 🔥 地段地號資訊標籤（預設隱藏）
    location_info_label = tk.Label(coord_row, text="", 
                                   font=entry_font,  # 🔥 改為與緯經度一樣大
                                   fg="#1976D2", 
                                   bg="#E3F2FD",
                                   padx=15, pady=2)
    # 🔥 預設不顯示（沒有 pack）
    
    # 檢查是否有從 data.json 載入的資訊
    if data.get('資料來源') == get_data_json_path() and data.get('地段'):
        info_text = f"📍 {data.get('縣市', '')}{data.get('行政區', '')} {data.get('地段', '')}{data.get('地號', '')}"
        location_info_label.config(text=info_text)
        location_info_label.pack(side=tk.LEFT, padx=10)  # 🔥 有資料才顯示
    
    def on_coord_change(*args):
        """當座標被手動修改時，清除並隱藏地段資訊"""
        if location_info_label.cget('text'):
            location_info_label.config(text="")
            location_info_label.pack_forget()  # 🔥 隱藏標籤（包括藍色底框）
            # 清除資料來源標記
            data.pop('資料來源', None)
            data.pop('縣市', None)
            data.pop('行政區', None)
            data.pop('地段', None)
            data.pop('地號', None)

    def bind_coord_change():
        lat_entry.bind('<KeyRelease>', on_coord_change)
        lng_entry.bind('<KeyRelease>', on_coord_change)
    
    editor_window.after(1000, bind_coord_change)

    # 從 data.json 讀取座標函數（但不建立按鈕）
    def load_coords_from_json():
        try:
            if not os.path.exists(get_data_json_path()):
                show_large_message("錯誤", "找不到 data.json 檔案")
                return
            
            with open(get_data_json_path(), 'r', encoding='utf-8') as f:
                data_list = json.load(f)
            
            if not data_list:
                show_large_message("錯誤", "data.json 檔案為空")
                return
            
            # 如果只有一組，直接載入
            if len(data_list) == 1:
                selected_data = data_list[0]
            else:
                # 多組資料選擇
                select_window = tk.Toplevel(editor_window)
                select_window.title("選擇地段資料")
                select_window.geometry("700x450")
                select_window.transient(editor_window)
                select_window.grab_set()
                
                x = editor_window.winfo_x() + (editor_window.winfo_width() - 700) // 2
                y = editor_window.winfo_y() + (editor_window.winfo_height() - 450) // 2
                select_window.geometry(f"+{x}+{y}")
                
                tk.Label(select_window, text="📋 data.json 中有多組資料，請選擇：",
                        font=("Microsoft JhengHei", 14, "bold"), fg="#1976D2").pack(pady=15)

                list_frame = tk.Frame(select_window)
                list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

                scrollbar = tk.Scrollbar(list_frame)
                scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

                listbox = tk.Listbox(list_frame, font=("Microsoft JhengHei", 12), height=14, yscrollcommand=scrollbar.set)
                listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
                scrollbar.config(command=listbox.yview)

                for i, item in enumerate(data_list):
                    display_text = f"{i+1}. {item.get('city','')}{item.get('area','')} {item.get('section','')}{item.get('lot_number','')} ({item.get('coordinates','')})"
                    listbox.insert(tk.END, display_text)

                listbox.selection_set(0)

                selected_index = [None]

                def on_select():
                    selection = listbox.curselection()
                    if selection:
                        selected_index[0] = selection[0]
                        select_window.destroy()

                def on_cancel():
                    select_window.destroy()

                button_frame = tk.Frame(select_window)
                button_frame.pack(pady=20)

                tk.Button(button_frame, text="[OK] 確定", command=on_select,
                        font=("Microsoft JhengHei", 14, "bold"), bg='#4CAF50', fg='white',
                        width=10, height=2).pack(side=tk.LEFT, padx=10)

                tk.Button(button_frame, text="✗ 取消", command=on_cancel,
                        font=("Microsoft JhengHei", 14, "bold"), bg='#F44336', fg='white',
                        width=10, height=2).pack(side=tk.LEFT, padx=10)

                listbox.bind('<Double-Button-1>', lambda e: on_select())

                select_window.wait_window()

                if selected_index[0] is None:
                    return

                selected_data = data_list[selected_index[0]]
            
            # 處理座標
            coordinates = selected_data.get('coordinates', '')
            if not coordinates or ',' not in coordinates:
                show_large_message("錯誤", "座標格式錯誤")
                return
            
            parts = coordinates.split(',')
            if len(parts) != 2:
                show_large_message("錯誤", "座標格式錯誤")
                return
            
            lat = parts[0].strip()
            lng = parts[1].strip()
            
            lat_entry.delete(0, tk.END)
            lng_entry.delete(0, tk.END)
            lat_entry.insert(0, lat)
            lng_entry.insert(0, lng)
            
            # 儲存地段資訊
            data['緯度'] = lat
            data['經度'] = lng
            data['資料來源'] = get_data_json_path()
            data['縣市'] = selected_data.get('city', '')
            data['行政區'] = selected_data.get('area', '')
            data['地段'] = selected_data.get('section', '')
            data['地號'] = selected_data.get('lot_number', '')
            
            # 顯示地段資訊
            info_text = f"📍 {data['縣市']}{data['行政區']} {data['地段']}{data['地號']}"
            location_info_label.config(text=info_text)
            location_info_label.pack(side=tk.LEFT, padx=10)  # 🔥 顯示標籤
            
            # 更新地圖
            if map_widget is not None:
                try:
                    map_widget.set_position(float(lat), float(lng))
                    map_widget.delete_all_marker()
                    map_widget.set_marker(float(lat), float(lng), text="物件位置")
                except:
                    pass
            
            # 🔥 自動嘗試載入都市計畫分區資料
            try:
                # 🔥 從 data 中解析地段地號
                target_lots = parse_land_lots_from_data(data)
                print(f"[檢查] 從土地標示解析出 {len(target_lots)} 筆地段地號")

                planning_data = load_urban_planning_data(target_lots=target_lots)
                if planning_data:
                    # 將使用分區（去重後用頓號分隔）填入「都市土地使用分區」欄位
                    # 🔥 若為「都市計畫外」則不帶入
                    zones_text = '、'.join(planning_data['zones'])
                    if zones_text and zones_text != '都市計畫外' and '都市土地使用分區' in entries:
                        entries['都市土地使用分區'].delete(0, tk.END)
                        entries['都市土地使用分區'].insert(0, zones_text)
                        print(f"[OK] 自動載入都市計畫分區：{zones_text}")
                    elif zones_text == '都市計畫外':
                        print(f"[略過] 都市計畫分區為「都市計畫外」，不帶入欄位")

                    # 🔥 將建蔽率（去重後用頓號分隔）填入「建蔽率」欄位
                    building_coverage_text = '、'.join(planning_data.get('building_coverage_ratios', []))
                    if building_coverage_text and '建蔽率' in entries:
                        entries['建蔽率'].delete(0, tk.END)
                        entries['建蔽率'].insert(0, building_coverage_text)
                        print(f"[OK] 自動載入建蔽率：{building_coverage_text}")

                    # 🔥 將容積率（去重後用頓號分隔）填入「容積率」欄位
                    floor_area_text = '、'.join(planning_data.get('floor_area_ratios', []))
                    if floor_area_text and '容積率' in entries:
                        entries['容積率'].delete(0, tk.END)
                        entries['容積率'].insert(0, floor_area_text)
                        print(f"[OK] 自動載入容積率：{floor_area_text}")

                    success_msg = f"已載入：{data['縣市']}{data['行政區']} {data['地段']}{data['地號']}\n\n緯度：{lat}\n經度：{lng}\n\n✓ 都市計畫使用分區：{zones_text}\n✓ 建蔽率：{building_coverage_text}\n✓ 容積率：{floor_area_text}"
                else:
                    success_msg = f"已載入：{data['縣市']}{data['行政區']} {data['地段']}{data['地號']}\n\n緯度：{lat}\n經度：{lng}\n\n(未找到符合條件的都市計畫分區資料)"
            except Exception as e:
                print(f"[提示] 無法載入都市計畫分區資料：{e}")
                success_msg = f"已載入：{data['縣市']}{data['行政區']} {data['地段']}{data['地號']}\n\n緯度：{lat}\n經度：{lng}"

            show_large_message("成功", success_msg)

        except Exception as e:
            show_large_message("錯誤", f"讀取失敗：{e}")
    
    # 預覽地圖函數（但不建立按鈕）
    def preview_map():
        lat = lat_entry.get()
        lng = lng_entry.get()
        
        if not lat or not lng:
            show_large_message("錯誤", "請先輸入經緯度座標")
            return
        
        import webbrowser
        map_url = f"https://www.google.com/maps?q={lat},{lng}&z=17"
        webbrowser.open(map_url)


    # 第2列：圖層選擇 + 縮放顯示
    layer_row = tk.Frame(map_frame)
    layer_row.pack(fill=tk.X, pady=3)

    tk.Label(layer_row, text="地圖圖層：", font=label_font, width=20, anchor="e").pack(side=tk.LEFT, padx=5)
    layer_var = tk.StringVar(value=data.get('地圖圖層', 'EMAP - 通用電子地圖(內政部)'))
    
    # 🔥 圖層選單字體加大
    layer_combo = ttk.Combobox(layer_row, textvariable=layer_var, 
                            font=("Microsoft JhengHei", 11),  # 🔥 從 entry_font 改為固定 11
                            width=32,
                            values=[
                                "EMAP - 通用電子地圖(內政部)",
                                "PHOTO_MIX - 衛星影像+地籍(內政部)",
                                "PHOTO2 - 衛星影像(內政部)",
                                "GOOGLE_SATELLITE - Google 衛星影像",
                                "GOOGLE_STREET - Google 街道地圖",
                                "GOOGLE_HYBRID - Google 混合地圖",
                                "OSM - OpenStreetMap 標準版",
                                "OSM_DE - OpenStreetMap 德國版",
                                "EMAP5 - 電子地圖(含等高線)",
                            ])
    layer_combo.pack(side=tk.LEFT, padx=5)
    entries["地圖圖層"] = layer_var

    tk.Label(layer_row, text="縮放：", font=entry_font).pack(side=tk.LEFT, padx=(15, 5))
    zoom_var = tk.StringVar(value=data.get('地圖縮放', '16'))
    zoom_display = tk.Label(layer_row, textvariable=zoom_var, font=("Microsoft JhengHei", 13, "bold"), 
                            fg="#1976D2", width=3, relief=tk.SUNKEN, bg="white")
    zoom_display.pack(side=tk.LEFT, padx=5)
    entries["地圖縮放"] = zoom_var

    # 🔥 第3列：所有按鈕在同一列（6個按鈕）
    button_row = tk.Frame(map_frame)
    button_row.pack(fill=tk.X, pady=10)
    
    # 建立按鈕容器並置中
    btn_frame = tk.Frame(button_row)
    btn_frame.pack()
    
    # 🔥 優化：定義所有按鈕功能（移除不必要的操作）
    def update_map_position():
        """更新地圖（優化版）"""
        try:
            lat = float(lat_entry.get())
            lng = float(lng_entry.get())
            current_layer = layer_var.get()
            
            # 簡化圖層切換
            if 'EMAP5' in current_layer:
                map_widget.set_tile_server("https://wmts.nlsc.gov.tw/wmts/EMAP5/default/GoogleMapsCompatible/{z}/{y}/{x}", max_zoom=20)
            elif 'PHOTO_MIX' in current_layer:
                map_widget.set_tile_server("https://wmts.nlsc.gov.tw/wmts/PHOTO_MIX/default/GoogleMapsCompatible/{z}/{y}/{x}", max_zoom=20)
            elif 'PHOTO2' in current_layer:
                map_widget.set_tile_server("https://wmts.nlsc.gov.tw/wmts/PHOTO2/default/GoogleMapsCompatible/{z}/{y}/{x}", max_zoom=20)
            elif 'GOOGLE_SATELLITE' in current_layer:
                map_widget.set_tile_server("https://mt0.google.com/vt/lyrs=s&x={x}&y={y}&z={z}", max_zoom=22)
            elif 'GOOGLE_HYBRID' in current_layer:
                map_widget.set_tile_server("https://mt0.google.com/vt/lyrs=y&x={x}&y={y}&z={z}", max_zoom=22)
            elif 'GOOGLE_STREET' in current_layer:
                map_widget.set_tile_server("https://mt0.google.com/vt/lyrs=m&x={x}&y={y}&z={z}", max_zoom=22)
            elif 'OSM_DE' in current_layer:
                map_widget.set_tile_server("https://a.tile.openstreetmap.de/{z}/{x}/{y}.png", max_zoom=19)
            elif 'OSM' in current_layer:
                map_widget.set_tile_server("https://tile.openstreetmap.org/{z}/{x}/{y}.png", max_zoom=19)
            elif 'EMAP' in current_layer:
                map_widget.set_tile_server("https://wmts.nlsc.gov.tw/wmts/EMAP/default/GoogleMapsCompatible/{z}/{y}/{x}", max_zoom=20)
            else:
                map_widget.set_tile_server("https://wmts.nlsc.gov.tw/wmts/EMAP/default/GoogleMapsCompatible/{z}/{y}/{x}", max_zoom=20)
            
            map_widget.set_position(lat, lng)
            map_widget.delete_all_marker()
            map_widget.set_marker(lat, lng, text="物件位置")
            
        except:
            pass  # 🔥 移除彈窗，加快速度
    
    def set_marker_at_center():
        """設定錨點並自動儲存"""
        try:
            position = map_widget.get_position()
            lat = f"{position[0]:.6f}"
            lng = f"{position[1]:.6f}"
            
            # 🔥🔥🔥 檢查座標是否改變，清除地段資訊 🔥🔥🔥
            old_lat = lat_entry.get().strip()
            old_lng = lng_entry.get().strip()
            
            if old_lat != lat or old_lng != lng:
                # 座標改變了，清除地段資訊
                if location_info_label.cget('text'):
                    location_info_label.config(text="")
                    location_info_label.pack_forget()  # 🔥 隱藏標籤
                    data.pop('資料來源', None)
                    data.pop('縣市', None)
                    data.pop('行政區', None)
                    data.pop('地段', None)
                    data.pop('地號', None)



            lat = f"{position[0]:.6f}"
            lng = f"{position[1]:.6f}"
            zoom = zoom_var.get()
            
            # 🔥 更新輸入框
            lat_entry.delete(0, tk.END)
            lng_entry.delete(0, tk.END)
            lat_entry.insert(0, lat)
            lng_entry.insert(0, lng)
            
            # 🔥 更新地圖標記
            map_widget.delete_all_marker()
            map_widget.set_marker(position[0], position[1], text="物件位置")
            
            # 🔥 更新 data 字典
            data['緯度'] = lat
            data['經度'] = lng
            data['地圖縮放'] = zoom
            
            # 🔥🔥🔥 自動儲存到 data_final.json 🔥🔥🔥
            try:
                data_final_path = get_data_final_path()
                with open(data_final_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                
                # 顯示成功訊息
                update_message(f"[完成] 錨點已設定並儲存：{lat}, {lng}")
                show_large_message(
                    "錨點已設定", 
                    f"新的錨點座標已儲存：\n\n"
                    f"緯度：{lat}\n"
                    f"經度：{lng}\n"
                    f"縮放級別：{zoom}\n\n"
                    f"[完成] 已自動儲存到 data_final.json"
                )
            except Exception as e:
                show_large_message("警告", f"座標已更新但儲存失敗：{e}")
                
        except Exception as e:
            show_large_message("錯誤", f"設定錨點失敗：{e}")

    # def save_current_position():  #暫時刪除不顯示儲存位置按鈕
    #     """儲存位置並更新輸入框"""
    #     try:
    #         position = map_widget.get_position()
    #         lat = f"{position[0]:.6f}"
    #         lng = f"{position[1]:.6f}"
    #         zoom = zoom_var.get()
            
    #         # 🔥 更新輸入框（確保同步）
    #         lat_entry.delete(0, tk.END)
    #         lng_entry.delete(0, tk.END)
    #         lat_entry.insert(0, lat)
    #         lng_entry.insert(0, lng)
            
    #         # 🔥 更新 data 字典
    #         data['緯度'] = lat
    #         data['經度'] = lng
    #         data['地圖縮放'] = zoom
            
    #         # 🔥 顯示儲存資訊
    #         update_message(f"✅ 位置已儲存：{lat}, {lng}, 縮放 {zoom}")
    #         show_large_message(
    #             "位置已儲存", 
    #             f"當前地圖位置已儲存：\n\n"
    #             f"緯度：{lat}\n"
    #             f"經度：{lng}\n"
    #             f"縮放級別：{zoom}\n\n"
    #             f"✅ 座標已更新到輸入框\n"
    #             f"[警告] 記得按下方「儲存」按鈕才會寫入 data_final.json"
    #         )
    #     except Exception as e:
    #         show_large_message("錯誤", f"儲存位置失敗：{e}")


    def capture_map_screenshot():
        """截取地圖（含座標檢查）"""
        try:
            # 🔥 檢查當前地圖中心和輸入框座標是否一致
            current_position = map_widget.get_position()
            current_lat = f"{current_position[0]:.6f}"
            current_lng = f"{current_position[1]:.6f}"
            
            input_lat = lat_entry.get().strip()
            input_lng = lng_entry.get().strip()
            
            # 🔥 如果座標不同，提示用戶
            if current_lat != input_lat or current_lng != input_lng:
                # 🔥🔥🔥 自訂大字體對話框 🔥🔥🔥
                dialog = tk.Toplevel(editor_window)
                dialog.title("地圖位置已變更")
                dialog.geometry("650x550")
                dialog.transient(editor_window)
                dialog.grab_set()
                
                # 置中顯示
                dialog.update_idletasks()
                x = editor_window.winfo_x() + (editor_window.winfo_width() - 650) // 2
                y = editor_window.winfo_y() + (editor_window.winfo_height() - 400) // 2
                dialog.geometry(f"+{x}+{y}")
                
                # 內容框架
                content_frame = tk.Frame(dialog, padx=30, pady=20)
                content_frame.pack(fill=tk.BOTH, expand=True)
                
                # 標題
                tk.Label(content_frame, 
                        text="[警告] 地圖位置已變更",
                        font=("Microsoft JhengHei", 16, "bold"),
                        fg="#FF6B00").pack(pady=10)
                
                # 訊息內容
                message_text = (
                    f"目前地圖中心與錨點不同：\n\n"
                    f"錨點座標：\n"
                    f"緯度 {input_lat}, 經度 {input_lng}\n\n"
                    f"地圖中心：\n"
                    f"緯度 {current_lat}, 經度 {current_lng}\n\n"
                    f"是否要更新地圖中心為錨點？"
                )
                
                tk.Label(content_frame,
                        text=message_text,
                        font=("Microsoft JhengHei", 12),
                        justify=tk.LEFT).pack(pady=15)
                
                # 選項說明
                options_text = (
                    "【是】- 更新錨點並截圖\n"
                    "【否】- 恢復原錨點後截圖\n"
                    "【取消】- 取消截圖"
                )
                
                tk.Label(content_frame,
                        text=options_text,
                        font=("Microsoft JhengHei", 11),
                        fg="#666666",
                        justify=tk.LEFT).pack(pady=10)
                
                # 按鈕框架
                button_frame = tk.Frame(content_frame)
                button_frame.pack(pady=20)
                
                # 結果變數
                result = [None]
                
                def on_yes():
                    result[0] = True
                    dialog.destroy()
                
                def on_no():
                    result[0] = False
                    dialog.destroy()
                
                def on_cancel():
                    result[0] = None
                    dialog.destroy()
                
                # 三個大按鈕
                tk.Button(button_frame, text="是 (Y)", command=on_yes,
                        font=("Microsoft JhengHei", 12, "bold"), 
                        bg='#4CAF50', fg='white',
                        width=10, height=2).pack(side=tk.LEFT, padx=10)
                
                tk.Button(button_frame, text="否 (N)", command=on_no,
                        font=("Microsoft JhengHei", 12, "bold"), 
                        bg='#FF9800', fg='white',
                        width=10, height=2).pack(side=tk.LEFT, padx=10)
                
                tk.Button(button_frame, text="取消", command=on_cancel,
                        font=("Microsoft JhengHei", 12, "bold"), 
                        bg='#F44336', fg='white',
                        width=10, height=2).pack(side=tk.LEFT, padx=10)
                
                # 綁定快捷鍵
                dialog.bind('y', lambda e: on_yes())
                dialog.bind('Y', lambda e: on_yes())
                dialog.bind('n', lambda e: on_no())
                dialog.bind('N', lambda e: on_no())
                dialog.bind('<Escape>', lambda e: on_cancel())
                
                # 等待對話框關閉
                editor_window.wait_window(dialog)
                
                result = result[0]
                
                if result is None:  # 取消
                    update_message("[警告] 已取消截圖")
                    return
                elif result:  # 是 - 更新錨點
                    # 更新為地圖中心
                    lat_entry.delete(0, tk.END)
                    lng_entry.delete(0, tk.END)
                    lat_entry.insert(0, current_lat)
                    lng_entry.insert(0, current_lng)
                    
                    data['緯度'] = current_lat
                    data['經度'] = current_lng
                    
                    # 自動儲存
                    try:
                        data_final_path = get_data_final_path()
                        with open(data_final_path, 'w', encoding='utf-8') as f:
                            json.dump(data, f, ensure_ascii=False, indent=2)
                        print(f"[OK] 錨點已更新並儲存：{current_lat}, {current_lng}")
                    except Exception as e:
                        print(f"[警告] 儲存失敗：{e}")
                    
                    # 更新標記
                    map_widget.delete_all_marker()
                    map_widget.set_marker(float(current_lat), float(current_lng), text="物件位置")
                    
                else:  # 否 - 恢復原座標
                    # 恢復到輸入框的座標
                    try:
                        map_widget.set_position(float(input_lat), float(input_lng))
                        map_widget.delete_all_marker()
                        map_widget.set_marker(float(input_lat), float(input_lng), text="物件位置")
                        editor_window.update()  # 強制更新畫面
                        import time
                        time.sleep(0.1)  # 等待地圖載入
                        print(f"[OK] 已恢復到原錨點：{input_lat}, {input_lng}")
                    except Exception as e:
                        print(f"[警告] 恢復座標失敗：{e}")
            
            # 🔥 執行截圖（使用輸入框的座標）
            zoom = int(zoom_var.get())
            lat = float(lat_entry.get())
            lng = float(lng_entry.get())
            current_layer = layer_var.get()
            
            # 提取圖層代碼
            if 'EMAP5' in current_layer:
                layer_code = 'EMAP5'
            elif 'PHOTO_MIX' in current_layer:
                layer_code = 'PHOTO_MIX'
            elif 'PHOTO2' in current_layer:
                layer_code = 'PHOTO2'
            elif 'GOOGLE_SATELLITE' in current_layer:
                layer_code = 'GOOGLE_SATELLITE'
            elif 'GOOGLE_HYBRID' in current_layer:
                layer_code = 'GOOGLE_HYBRID'
            elif 'GOOGLE_STREET' in current_layer:
                layer_code = 'GOOGLE_STREET'
            elif 'OSM_DE' in current_layer:
                layer_code = 'OSM_DE'
            elif 'OSM' in current_layer:
                layer_code = 'OSM'
            elif 'EMAP' in current_layer:
                layer_code = 'EMAP'
            else:
                layer_code = 'EMAP'
            
            print(f"開始截圖：圖層={layer_code}, 縮放={zoom}, 座標={lat},{lng}")
            update_message(f"正在生成地圖...")
            
            # 生成地圖
            map_image_path = generate_map_image(lat, lng, zoom, layer_code)
            
            if not map_image_path:
                show_large_message("錯誤", "地圖生成失敗")
                return
            
            from PIL import Image, ImageTk
            screenshot = Image.open(map_image_path)
            
            # 顯示預覽視窗
            preview_window = tk.Toplevel(editor_window)
            preview_window.title("地圖預覽")
            preview_window.geometry("680x750")
            preview_window.transient(editor_window)
            preview_window.grab_set()
            
            preview_window.update_idletasks()
            x = editor_window.winfo_x() + (editor_window.winfo_width() - 680) // 2
            y = editor_window.winfo_y() + (editor_window.winfo_height() - 750) // 2
            preview_window.geometry(f"+{x}+{y}")
            
            canvas = tk.Canvas(preview_window)
            scrollbar = tk.Scrollbar(preview_window, orient="vertical", command=canvas.yview)
            scrollable_frame = tk.Frame(canvas)
            scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            
            tk.Label(scrollable_frame, text="📸 地圖預覽", font=("Microsoft JhengHei", 14, "bold"), fg="#1976D2").pack(pady=10)
            tk.Label(scrollable_frame, text=f"座標：{lat:.6f}, {lng:.6f}\n圖層：{layer_code}\n縮放：{zoom}", 
                    font=("Microsoft JhengHei", 10)).pack(pady=5)
            
            img_frame = tk.Frame(scrollable_frame, relief=tk.SUNKEN, borderwidth=2)
            img_frame.pack(pady=10, padx=20)
            photo = ImageTk.PhotoImage(screenshot)
            img_label = tk.Label(img_frame, image=photo)
            img_label.image = photo
            img_label.pack()
            
            button_frame = tk.Frame(scrollable_frame)
            button_frame.pack(pady=20)
            
            def save_screenshot():
                # 🔥 取得正確的截圖路徑（4.其他相關 資料夾）
                target_path = get_screenshot_path()
                if target_path:
                    screenshot.save(target_path, 'PNG')
                    update_message(f"[完成] 地圖已儲存至：{target_path}")
                    show_large_message("成功", f"地圖已儲存至：\n{target_path}")
                else:
                    # 如果無法取得路徑，回退到當前目錄
                    screenshot.save("gui_map_screenshot.png", 'PNG')
                    update_message("[完成] 地圖已儲存（當前目錄）")
                    show_large_message("警告", "無法取得 data.json 資訊\n\n地圖已儲存至當前目錄：gui_map_screenshot.png")

                preview_window.destroy()
                try:
                    os.remove(map_image_path)
                except:
                    pass
            
            def retake():
                preview_window.destroy()
                try:
                    os.remove(map_image_path)
                except:
                    pass
                editor_window.after(200, capture_map_screenshot)
            
            def cancel():
                try:
                    os.remove(map_image_path)
                except:
                    pass
                preview_window.destroy()
            
            tk.Button(button_frame, text="💾 儲存", command=save_screenshot,
                    font=("Microsoft JhengHei", 12, "bold"), bg='#4CAF50', fg='white',
                    width=10, height=2).pack(side=tk.LEFT, padx=10)
            tk.Button(button_frame, text="🔄 重新生成", command=retake,
                    font=("Microsoft JhengHei", 12), bg='#FF9800', fg='white',
                    width=10, height=2).pack(side=tk.LEFT, padx=10)
            tk.Button(button_frame, text="[錯誤] 取消", command=cancel,
                    font=("Microsoft JhengHei", 12), bg='#F44336', fg='white',
                    width=10, height=2).pack(side=tk.LEFT, padx=10)
            
        except Exception as e:
            show_large_message("錯誤", f"截圖失敗：{e}")
            import traceback
            traceback.print_exc()

    # 🔥 建立6個按鈕（全部在同一列）
    tk.Button(btn_frame, text="讀取 data 坐標", command=load_coords_from_json,
            font=("Microsoft JhengHei", 11, "bold"), bg='#4CAF50', fg='white',
            padx=10, pady=6).pack(side=tk.LEFT, padx=3)
    
    tk.Button(btn_frame, text="瀏覽器預覽", command=preview_map,
            font=("Microsoft JhengHei", 11, "bold"), bg='#2196F3', fg='white',
            padx=10, pady=6).pack(side=tk.LEFT, padx=3)
    
    tk.Button(btn_frame, text="🔄 更新地圖", command=update_map_position,
            font=("Microsoft JhengHei", 11, "bold"), bg='#FF9800', fg='white',
            padx=10, pady=6).pack(side=tk.LEFT, padx=3)
    
    tk.Button(btn_frame, text="📍 設定錨點", command=set_marker_at_center,
            font=("Microsoft JhengHei", 11, "bold"), bg="#9C27B0", fg='white',
            padx=10, pady=6).pack(side=tk.LEFT, padx=3)
    
    # tk.Button(btn_frame, text="💾 儲存位置", command=save_current_position,
    #     font=("Microsoft JhengHei", 11, "bold"), bg='#9C27B0', fg='white',
    #     padx=10, pady=6).pack(side=tk.LEFT, padx=3)
    
    tk.Button(btn_frame, text="📸 截取地圖", command=capture_map_screenshot,
            font=("Microsoft JhengHei", 11, "bold"), bg='#E91E63', fg='white',
            padx=10, pady=6).pack(side=tk.LEFT, padx=3)
    
    # 🔥 優化：圖層變更時不彈訊息，直接更新
    def on_layer_change(*args):
        try:
            if map_widget:
                update_map_position()
        except:
            pass
    
    layer_var.trace('w', on_layer_change)

    # 第4列：互動式地圖
    map_widget_frame = tk.Frame(map_frame)
    map_widget_frame.pack(fill=tk.BOTH, expand=True, pady=10)

    try:
        import tkintermapview
        from PIL import Image, ImageGrab, ImageTk
        
        map_widget = tkintermapview.TkinterMapView(map_widget_frame, width=600, height=400, corner_radius=0)
        map_widget.pack(pady=10)
        
        initial_lat = float(data.get('緯度', '22.65496')) if data.get('緯度') else 22.65496
        initial_lng = float(data.get('經度', '120.33575')) if data.get('經度') else 120.33575
        initial_zoom = int(data.get('地圖縮放', '16'))
        
        map_widget.set_position(initial_lat, initial_lng)
        map_widget.set_zoom(initial_zoom)
        
        current_layer = data.get('地圖圖層', 'EMAP - 通用電子地圖(內政部))')
        
        if 'EMAP5' in current_layer:
            map_widget.set_tile_server("https://wmts.nlsc.gov.tw/wmts/EMAP5/default/GoogleMapsCompatible/{z}/{y}/{x}", max_zoom=20)
        elif 'PHOTO_MIX' in current_layer:
            map_widget.set_tile_server("https://wmts.nlsc.gov.tw/wmts/PHOTO_MIX/default/GoogleMapsCompatible/{z}/{y}/{x}", max_zoom=20)
        elif 'PHOTO2' in current_layer:
            map_widget.set_tile_server("https://wmts.nlsc.gov.tw/wmts/PHOTO2/default/GoogleMapsCompatible/{z}/{y}/{x}", max_zoom=20)
        elif 'GOOGLE_SATELLITE' in current_layer:
            map_widget.set_tile_server("https://mt0.google.com/vt/lyrs=s&x={x}&y={y}&z={z}", max_zoom=22)
        elif 'GOOGLE_HYBRID' in current_layer:
            map_widget.set_tile_server("https://mt0.google.com/vt/lyrs=y&x={x}&y={y}&z={z}", max_zoom=22)
        elif 'GOOGLE_STREET' in current_layer:
            map_widget.set_tile_server("https://mt0.google.com/vt/lyrs=m&x={x}&y={y}&z={z}", max_zoom=22)
        elif 'OSM_DE' in current_layer:
            map_widget.set_tile_server("https://a.tile.openstreetmap.de/{z}/{x}/{y}.png", max_zoom=19)
        elif 'OSM' in current_layer:
            map_widget.set_tile_server("https://tile.openstreetmap.org/{z}/{x}/{y}.png", max_zoom=19)
        elif 'EMAP' in current_layer:
            map_widget.set_tile_server("https://wmts.nlsc.gov.tw/wmts/EMAP/default/GoogleMapsCompatible/{z}/{y}/{x}", max_zoom=20)
        else:
            map_widget.set_tile_server("https://wmts.nlsc.gov.tw/wmts/EMAP/default/GoogleMapsCompatible/{z}/{y}/{x}", max_zoom=20)
        
        marker = map_widget.set_marker(initial_lat, initial_lng, text="物件位置")

        # 🔥 完全禁用默認的滾輪縮放，改為 Ctrl+滾輪縮放
        def on_map_mouse_wheel(event):
            """自定義滾輪事件：只有按住 Ctrl 才能縮放地圖"""
            if event.state & 0x0004:  # 檢查 Ctrl 鍵是否按下（0x0004 = Control modifier）
                # Ctrl 被按下，允許縮放
                try:
                    current_zoom = int(zoom_var.get())
                    if event.delta > 0:  # 向上滾
                        new_zoom = min(current_zoom + 1, 22)  # 最大 zoom 22
                    else:  # 向下滾
                        new_zoom = max(current_zoom - 1, 1)  # 最小 zoom 1

                    if new_zoom != current_zoom:
                        map_widget.set_zoom(new_zoom)
                        zoom_var.set(str(new_zoom))
                        last_known_zoom[0] = new_zoom
                except Exception as e:
                    print(f"縮放錯誤: {e}")

                return "break"  # 阻止事件繼續傳播
            else:
                # Ctrl 沒按下，阻止地圖縮放，但允許頁面滾動
                return "break"  # 🔥 關鍵：阻止地圖處理滾輪事件

        # 🔥 解除 tkintermapview 預設綁定的所有滾輪事件
        # 先解除所有可能的滾輪綁定（包含元件和畫布層級）
        try:
            # 解除元件層級的綁定
            map_widget.unbind("<MouseWheel>")
            map_widget.unbind("<Button-4>")  # Linux 向上滾
            map_widget.unbind("<Button-5>")  # Linux 向下滾
            # 解除畫布層級的綁定
            map_widget.canvas.unbind("<MouseWheel>")
            map_widget.canvas.unbind("<Button-4>")
            map_widget.canvas.unbind("<Button-5>")
        except:
            pass

        # 現在綁定我們的自訂滾輪處理（同時在元件和畫布層級）
        map_widget.bind("<MouseWheel>", on_map_mouse_wheel)
        map_widget.bind("<Button-4>", on_map_mouse_wheel)  # Linux
        map_widget.bind("<Button-5>", on_map_mouse_wheel)  # Linux
        map_widget.canvas.bind("<MouseWheel>", on_map_mouse_wheel)
        map_widget.canvas.bind("<Button-4>", on_map_mouse_wheel)
        map_widget.canvas.bind("<Button-5>", on_map_mouse_wheel)

        # 🔥 禁用 tkintermapview 內部的滾輪縮放功能
        # 覆蓋所有可能的縮放方法
        if hasattr(map_widget, '_mouse_zoom'):
            def disabled_mouse_zoom(event):
                pass  # 不執行任何操作
            map_widget._mouse_zoom = disabled_mouse_zoom

        if hasattr(map_widget, 'mouse_zoom'):
            def disabled_mouse_zoom2(event):
                pass
            map_widget.mouse_zoom = disabled_mouse_zoom2

        # 注意：不使用 unbind_all，避免影響整個應用的滾輪事件
        # 透過 return "break" 已經可以阻止事件傳播到地圖內部處理

        last_known_zoom = [initial_zoom]

        def check_zoom_change():
            try:
                if hasattr(map_widget, '_zoom'):
                    current = int(map_widget._zoom)
                elif hasattr(map_widget, 'zoom'):
                    current = int(map_widget.zoom)
                else:
                    editor_window.after(200, check_zoom_change)
                    return
                
                if current != last_known_zoom[0]:
                    last_known_zoom[0] = current
                    zoom_var.set(str(current))
            except:
                pass
            
            editor_window.after(200, check_zoom_change)
        
        editor_window.after(1000, check_zoom_change)
        
        print("[OK] 互動式地圖已載入")

    except ImportError:
        no_map_label = tk.Label(map_widget_frame,
                            text="[警告] 未安裝 tkintermapview 套件\n\n請執行：pip install tkintermapview",
                            font=("Microsoft JhengHei", 12), fg='red')
        no_map_label.pack(pady=50)
    except Exception as e:
        error_label = tk.Label(map_widget_frame,
                            text=f"[錯誤] 地圖載入失敗：{e}",
                            font=("Microsoft JhengHei", 10), fg='red')
        error_label.pack(pady=50)

    # === 承辦人資訊 ===
    agent_frame = create_section(scrollable_frame, "👤 承辦人資訊")

    # 載入通訊錄（按店別分組）
    # 🔥 臨時強制使用備用函數以獲取詳細日誌
    load_func = _fallback_load_contacts_from_excel  # 強制使用備用函數
    # load_func = load_contacts_from_excel if load_contacts_from_excel else _fallback_load_contacts_from_excel

    # 記錄載入狀態到檔案（已停用）
    def write_log(msg):
        pass  # 不再寫入 log 檔案

    write_log("=" * 60)
    write_log("開始載入通訊錄")
    write_log(f"使用函數：{'main.load_contacts_from_excel' if load_contacts_from_excel else '_fallback_load_contacts_from_excel'}")

    contacts_data = load_func('通訊錄.xlsx')
    stores_list = contacts_data.get('stores', [])
    contacts_by_store = contacts_data.get('contacts_by_store', {})

    # 記錄載入結果
    write_log(f"通訊錄載入完成")
    write_log(f"  店別數量：{len(stores_list)}")
    write_log(f"  店別列表：{stores_list}")
    total_contacts = sum(len(contacts) for contacts in contacts_by_store.values())
    write_log(f"  總聯絡人數：{total_contacts}")

    # 如果沒有載入任何資料，顯示警告
    if not stores_list or total_contacts == 0:
        write_log("警告：通訊錄沒有載入任何資料！")
        import tkinter.messagebox as messagebox
        messagebox.showwarning("通訊錄警告",
            f"通訊錄載入異常\n\n"
            f"店別數量：{len(stores_list)}\n"
            f"聯絡人數量：{total_contacts}\n\n"
            f"請檢查 通訊錄.xlsx 檔案是否存在且格式正確")

    # 取得已儲存的承辦人列表
    saved_agents = data.get("承辦人列表", [])
    if isinstance(saved_agents, str):  # 相容舊版本（單一承辦人字串格式）
        saved_agents = []

    # 儲存承辦人選擇的變數、店別變數和電話標籤
    agent_store_vars = []  # 店別選擇
    agent_name_vars = []  # 姓名選擇
    agent_name_combos = []  # 姓名下拉選單（用於更新選項）
    agent_phone_labels = []  # 電話顯示
    set_default_var = tk.BooleanVar(value=False)  # 🔥 預設勾選框變數

    # 建立6個承辦人選擇下拉選單
    for i in range(6):
        agent_row = tk.Frame(agent_frame)
        agent_row.pack(fill=tk.X, pady=3)

        tk.Label(agent_row, text=f"承辦人 {i+1}：", font=label_font, width=15, anchor="e").pack(side=tk.LEFT, padx=5)

        # 店別下拉選單
        store_var = tk.StringVar(value="")
        store_combo = ttk.Combobox(agent_row, textvariable=store_var, font=combo_font, width=12,
                                    values=[""] + stores_list, state="readonly")
        store_combo.pack(side=tk.LEFT, padx=5)
        agent_store_vars.append(store_var)

        # 姓名下拉選單（初始為空）
        name_var = tk.StringVar(value="")
        name_combo = ttk.Combobox(agent_row, textvariable=name_var, font=combo_font, width=12,
                                  values=[""], state="readonly")
        name_combo.pack(side=tk.LEFT, padx=5)
        agent_name_vars.append(name_var)
        agent_name_combos.append(name_combo)

        # 電話顯示標籤（縮小寬度以容納勾選框）
        phone_label = tk.Label(agent_row, text="", font=entry_font, fg="blue", width=12, anchor="w")
        phone_label.pack(side=tk.LEFT, padx=3)
        agent_phone_labels.append(phone_label)

        # 🔥 只在承辦人1添加「預設」勾選框（直接接在電話後面）
        if i == 0:
            default_check = tk.Checkbutton(agent_row, text="預設", variable=set_default_var,
                                          font=entry_font, command=lambda: save_default_agent())
            default_check.pack(side=tk.LEFT, padx=0)

        # 當選擇店別時，更新姓名選項
        def on_store_change(idx=i):
            """當店別改變時，更新姓名下拉選單"""
            selected_store = agent_store_vars[idx].get()
            if selected_store and selected_store in contacts_by_store:
                names = [c['name'] for c in contacts_by_store[selected_store]]
                agent_name_combos[idx]['values'] = [""] + names
            else:
                agent_name_combos[idx]['values'] = [""]
            # 清空姓名和電話
            agent_name_vars[idx].set("")
            agent_phone_labels[idx].config(text="")

        store_var.trace('w', lambda *args, idx=i: on_store_change(idx))

        # 當選擇姓名時，更新電話顯示並檢查重複
        def on_name_change(idx=i):
            """當姓名改變時，更新電話顯示並檢查重複"""
            selected_store = agent_store_vars[idx].get()
            selected_name = agent_name_vars[idx].get()

            if selected_name and selected_store:
                # 檢查是否重複選擇
                for j in range(6):
                    if j != idx and agent_name_vars[j].get() == selected_name:
                        # 發現重複，清空當前選擇
                        agent_name_vars[idx].set("")
                        agent_phone_labels[idx].config(text="")
                        show_large_message("重複選擇", f"承辦人「{selected_name}」已被選擇！\n請選擇其他人員。")
                        return

                # 顯示電話
                store_contacts = contacts_by_store.get(selected_store, [])
                for contact in store_contacts:
                    if contact['name'] == selected_name:
                        agent_phone_labels[idx].config(text=contact['phone'])
                        break
            else:
                agent_phone_labels[idx].config(text="")

        name_var.trace('w', lambda *args, idx=i: on_name_change(idx))

        # 初始化：載入已儲存的值
        if i < len(saved_agents):
            saved_name = saved_agents[i]
            # 🔥 修正：尋找該姓名所屬的店別，如果找不到也要顯示姓名
            found_store = None
            found_phone = ""
            for store, contacts in contacts_by_store.items():
                for contact in contacts:
                    if contact['name'] == saved_name:
                        found_store = store
                        found_phone = contact.get('phone', '')
                        break
                if found_store:
                    break

            if found_store:
                # 找到店別，正常載入
                store_var.set(found_store)
                # 觸發店別改變會更新姓名選項
                # 然後設定姓名
                editor_window.after(100, lambda idx=i, n=saved_name, p=found_phone: (
                    agent_name_vars[idx].set(n),
                    agent_phone_labels[idx].config(text=p)
                ))
            else:
                # 🔥 找不到店別，但仍然顯示姓名（手動填入模式）
                print(f"[警告] 承辦人 {i+1}「{saved_name}」在通訊錄中找不到，但仍載入姓名")
                # 將姓名選項設為包含此姓名
                agent_name_combos[i]['values'] = ["", saved_name]
                agent_name_vars[i].set(saved_name)
                agent_phone_labels[i].config(text="(通訊錄中無此人)", fg="orange")

    entries["承辦人列表"] = agent_name_vars

    # 🔥 儲存預設承辦人的函數
    def save_default_agent():
        """儲存承辦人1為預設值"""
        if set_default_var.get():
            # 取得承辦人1的店別和姓名
            store = agent_store_vars[0].get()
            name = agent_name_vars[0].get()

            if not store or not name:
                update_message("[提示] 請先選擇承辦人1的店別和姓名")
                set_default_var.set(False)
                return

            # 儲存到 config.json
            try:
                from config_manager import load_config, save_config
                config = load_config()
                if 'default_agent' not in config:
                    config['default_agent'] = {}

                config['default_agent'] = {
                    'store': store,
                    'name': name
                }

                save_config(config)
                update_message(f"[成功] 已設定預設承辦人：{store} - {name}")
            except Exception as e:
                update_message(f"[錯誤] 儲存預設承辦人失敗：{e}")
                set_default_var.set(False)
        else:
            # 取消預設值
            try:
                from config_manager import load_config, save_config
                config = load_config()
                if 'default_agent' in config:
                    del config['default_agent']
                    save_config(config)
                    update_message("[提示] 已取消預設承辦人")
            except Exception as e:
                update_message(f"[錯誤] 取消預設承辦人失敗：{e}")

    # 🔥 載入預設承辦人（如果存在）
    try:
        from config_manager import load_config
        config = load_config()
        default_agent = config.get('default_agent', {})

        if default_agent and 'store' in default_agent and 'name' in default_agent:
            # 如果承辦人1沒有已儲存的值，則使用預設值
            if not saved_agents or not saved_agents[0]:
                default_store = default_agent['store']
                default_name = default_agent['name']

                # 設定承辦人1的店別和姓名
                agent_store_vars[0].set(default_store)
                # 延遲設定姓名，確保店別選項已更新
                editor_window.after(150, lambda: agent_name_vars[0].set(default_name))

                update_message(f"[提示] 已自動載入預設承辦人：{default_store} - {default_name}")

            # 勾選「預設」框
            if (agent_store_vars[0].get() == default_agent['store'] and
                agent_name_vars[0].get() == default_agent['name']):
                set_default_var.set(True)
    except Exception as e:
        update_message(f"[提示] 載入預設承辦人失敗：{e}")

    # ========================================
    # 🔥 儲存和取消按鈕
    # ========================================
    button_frame = tk.Frame(scrollable_frame)
    button_frame.pack(pady=20)
    
    def save_data():
        """儲存資料"""
        try:
            # 收集所有欄位資料
            for key, widget in entries.items():
                if key.startswith("格局_"):
                    field = key.replace("格局_", "")
                    if "格局" not in data:
                        data["格局"] = {}
                    data["格局"][field] = widget.get()
                elif key.startswith("主建物_樓層"):
                    # 跳過，稍後統一處理
                    continue
                elif key.startswith("車位_型式_"):
                    continue
                elif key.startswith("車位_"):
                    field = key.replace("車位_", "")
                    if "車位" not in data:
                        data["車位"] = {}
                    if isinstance(widget, tk.Entry):
                        data["車位"][field] = widget.get()
                    else:
                        data["車位"][field] = widget.get()
                elif key.startswith("所有權型態_"):
                    continue
                elif key == "隔局":
                    # 處理隔局複選框
                    continue
                elif key == "適合行業別":
                    # 處理適合行業別複選框
                    continue
                elif key == "承辦人列表":
                    # 處理承辦人列表
                    continue
                elif key == "稅賦":
                    # 處理稅賦單選框
                    continue
                elif key == "訴求重點":
                    data[key] = widget.get("1.0", "end-1c")
                elif key == "另顯示於後台內容":
                    data[key] = widget.get("1.0", "end-1c")
                elif isinstance(widget, tk.BooleanVar):
                    data[key] = widget.get()
                elif isinstance(widget, tk.StringVar):
                    data[key] = widget.get()
                elif isinstance(widget, tk.Entry):
                    data[key] = widget.get()

            # 處理隔局複選框
            if "隔局" in entries:
                隔局列表 = []
                for option, var in entries["隔局"].items():
                    if var.get():
                        隔局列表.append(option)
                data["格局選項"] = 隔局列表

            # 處理適合行業別複選框
            if "適合行業別" in entries:
                適合行業別_dict = {}
                for key, widget in entries["適合行業別"].items():
                    if key == "其他_text":
                        # 處理其他文字輸入框
                        適合行業別_dict[key] = widget.get()
                    else:
                        # 處理複選框
                        適合行業別_dict[key] = widget.get()
                data["適合行業別"] = 適合行業別_dict

            # 處理承辦人列表
            if "承辦人列表" in entries:
                承辦人列表 = []
                for agent_var in entries["承辦人列表"]:
                    agent_name = agent_var.get().strip()
                    if agent_name:  # 只保存非空的承辦人
                        承辦人列表.append(agent_name)
                data["承辦人列表"] = 承辦人列表

            # 處理稅賦單選框
            if "稅賦" in entries:
                稅賦資料 = {}
                for tax_item, var in entries["稅賦"].items():
                    value = var.get()
                    # ⭐ 將「清除選擇」轉回空字串
                    if value == "清除選擇":
                        value = ""
                    稅賦資料[tax_item] = value
                data["稅賦"] = 稅賦資料

            # 🔥 處理主建物各層資料
            主建物明細 = {}
            for i in range(9):
                floor_name_key = f"主建物_樓層{i+1}_名稱"
                floor_area_key = f"主建物_樓層{i+1}_坪數"

                if floor_name_key in entries and floor_area_key in entries:
                    floor_name = entries[floor_name_key].get().strip()
                    floor_area = entries[floor_area_key].get().strip()

                    if floor_name and floor_area:
                        主建物明細[floor_name] = floor_area

            data["主建物"] = 主建物明細

            # 🔥 處理附屬建物各項資料
            附屬建物明細 = {}
            for i in range(9):
                attached_type_key = f"附屬建物_項目{i+1}_類型"
                attached_area_key = f"附屬建物_項目{i+1}_坪數"

                if attached_type_key in entries and attached_area_key in entries:
                    attached_type = entries[attached_type_key].get().strip()
                    attached_area = entries[attached_area_key].get().strip()

                    if attached_type and attached_area:
                        附屬建物明細[attached_type] = attached_area

            data["附屬建物"] = 附屬建物明細
            
            # 處理車位型式的複選
            型式列表 = []
            if entries.get("車位_型式_平面") and entries["車位_型式_平面"].get():
                型式列表.append("平面")
            if entries.get("車位_型式_機械") and entries["車位_型式_機械"].get():
                型式列表.append("機械")
            
            if "車位" not in data:
                data["車位"] = {}
            data["車位"]["型式"] = "、".join(型式列表) if 型式列表 else ""
            
            # 處理所有權型態的複選
            所有權型態列表 = []
            if entries.get("所有權型態_單獨所有") and entries["所有權型態_單獨所有"].get():
                所有權型態列表.append("單獨所有")
            if entries.get("所有權型態_分別共有") and entries["所有權型態_分別共有"].get():
                所有權型態列表.append("分別共有")
            if entries.get("所有權型態_公同共有") and entries["所有權型態_公同共有"].get():
                所有權型態列表.append("公同共有")
            if entries.get("所有權型態_地上權") and entries["所有權型態_地上權"].get():
                所有權型態列表.append("地上權")
            
            data["所有權型態"] = "、".join(所有權型態列表) if 所有權型態列表 else ""
            
            # 處理現況說明
            if entries.get("現況說明"):
                data["現況說明"] = entries["現況說明"].get()

            # 🔥 智能截圖：只有當截圖不存在時才自動截圖
            map_captured = False
            screenshot_needed = not os.path.exists("gui_map_screenshot.png")

            if screenshot_needed:
                update_message("[進行中] 正在自動截取地圖...")
                try:
                    if map_widget is not None:
                        # 取得當前的地圖參數
                        lat = float(entries["緯度"].get())
                        lng = float(entries["經度"].get())
                        zoom = int(zoom_var.get())
                        current_layer = layer_var.get()

                        # 從圖層名稱提取圖層代碼
                        if 'EMAP5' in current_layer:
                            layer_code = 'EMAP5'
                        elif 'PHOTO_MIX' in current_layer:
                            layer_code = 'PHOTO_MIX'
                        elif 'PHOTO2' in current_layer:
                            layer_code = 'PHOTO2'
                        elif 'GOOGLE_SATELLITE' in current_layer:
                            layer_code = 'GOOGLE_SATELLITE'
                        elif 'GOOGLE_HYBRID' in current_layer:
                            layer_code = 'GOOGLE_HYBRID'
                        elif 'GOOGLE_STREET' in current_layer:
                            layer_code = 'GOOGLE_STREET'
                        elif 'OSM_DE' in current_layer:
                            layer_code = 'OSM_DE'
                        elif 'OSM' in current_layer:
                            layer_code = 'OSM'
                        elif 'EMAP' in current_layer:
                            layer_code = 'EMAP'
                        else:
                            layer_code = 'EMAP'

                        print(f"自動生成地圖：圖層={layer_code}, 縮放={zoom}")

                        # 生成地圖
                        temp_map_path = generate_map_image(lat, lng, zoom, layer_code)

                        if temp_map_path and os.path.exists(temp_map_path):
                            # 將臨時地圖複製為最終截圖
                            import shutil
                            shutil.copy2(temp_map_path, "gui_map_screenshot.png")

                            # 清理臨時檔案
                            try:
                                os.remove(temp_map_path)
                            except:
                                pass

                            if os.path.exists("gui_map_screenshot.png"):
                                file_size = os.path.getsize("gui_map_screenshot.png")
                                print(f"[完成] 地圖自動生成成功：{file_size} bytes")
                                map_captured = True
                        else:
                            print("[警告] 地圖生成失敗")
                except Exception as e:
                    print(f"[錯誤] 自動截圖失敗：{e}")
                    import traceback
                    traceback.print_exc()

            # 儲存到檔案
            data_final_path = get_data_final_path()
            with open(data_final_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            # 🔥 顯示儲存資訊（精簡版 + 案件 + 欄位數，避免大片留白）
            # 取出案件名（從 data_final_path 倒推案件資料夾）
            case_name = ""
            try:
                # data_final_path 結構：<案件>/4.其他相關/data_final.json
                case_name = os.path.basename(os.path.dirname(os.path.dirname(data_final_path)))
            except Exception:
                pass

            # 計算「有值」的欄位數（過濾掉空字串和「【沒得選】」之類預設值）
            filled_count = sum(
                1 for v in data.values()
                if v and v not in ("", "【沒得選】", [], {}, None)
            )

            saved_info = f"案件：{case_name}\n"
            saved_info += f"已儲存 {filled_count}/{len(data)} 個欄位\n"
            saved_info += f"📍 座標：{data.get('緯度', '')}, {data.get('經度', '')}\n"
            saved_info += f"🔍 縮放 {data.get('地圖縮放', '')}　🗺️ 圖層 {data.get('地圖圖層', '').split(' - ')[0] if data.get('地圖圖層') else ''}\n"

            # 截圖狀態（精簡成一行）
            if os.path.exists("gui_map_screenshot.png"):
                kb = os.path.getsize("gui_map_screenshot.png") // 1024
                if screenshot_needed and map_captured:
                    saved_info += f"📸 地圖截圖：自動截取成功 ({kb} KB)"
                else:
                    saved_info += f"📸 地圖截圖：已存在 ({kb} KB)"
            else:
                saved_info += "⚠️ 地圖截圖：截取失敗，請手動截圖"

            update_message("[完成] 資料已儲存")
            show_large_message("資料已儲存", saved_info)
            
        except Exception as e:
            update_message(f"[錯誤] 儲存失敗: {e}")
            messagebox.showwarning("警告", f"儲存失敗：{e}")
            import traceback
            traceback.print_exc()

    
    def cancel_edit():
        """取消編輯"""
        result = messagebox.askyesno("確認", "確定要取消編輯？\n未儲存的變更將會遺失。")
        if result:
            editor_window.destroy()
    
    tk.Button(button_frame, text="💾 儲存", command=save_data,
             font=("Microsoft JhengHei", 12, "bold"), bg='#4CAF50', fg='white',
             width=10, height=2).pack(side=tk.LEFT, padx=20)

    tk.Button(button_frame, text="生成物件調查表", command=generate_report_from_button,
             font=("Microsoft JhengHei", 12, "bold"), bg='#2196F3', fg='white',
             width=15, height=2).pack(side=tk.LEFT, padx=20)

    # 🔥 匯出封存按鈕（支持 Shift+點擊強制重設）
    def export_archive_wrapper(event=None):
        """包裝函數，檢測 Shift 鍵"""
        force_reset = False
        if event and (event.state & 0x0001):  # 檢測 Shift 鍵
            force_reset = True
            update_message("[強制] Shift+點擊，將重新設定匯出選項")

        # 使用 after 確保按鈕視覺狀態正確更新
        def run_export():
            try:
                export_archive(force_reset=force_reset)
            finally:
                # 確保按鈕恢復正常狀態
                try:
                    export_btn.config(relief=tk.RAISED)
                    export_btn.update_idletasks()
                except:
                    pass

        # 延遲執行，讓按鈕視覺效果正確顯示
        editor_window.after(10, run_export)

    export_btn = tk.Button(button_frame, text="📦 匯出封存",
                          font=("Microsoft JhengHei", 12, "bold"), bg='#FF9800', fg='white',
                          width=12, height=2)
    export_btn.pack(side=tk.LEFT, padx=20)
    export_btn.bind('<Button-1>', export_archive_wrapper)

    tk.Button(button_frame, text="[錯誤] 取消", command=cancel_edit,
             font=("Microsoft JhengHei", 12, "bold"), bg='#F44336', fg='white',
             width=10, height=2).pack(side=tk.LEFT, padx=20)
    
    # 組裝滾動區域
    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # 🔥 綁定滑鼠滾輪事件（修正：只綁定到 canvas，不影響其他元件）
    def on_mousewheel(event):
        canvas.yview_scroll(int(-1*(event.delta/120)), "units")

    # 直接綁定到 canvas，不使用 bind_all
    canvas.bind("<MouseWheel>", on_mousewheel)

    # 當視窗關閉時的處理
    def on_closing():
        editor_window.destroy()

    editor_window.protocol("WM_DELETE_WINDOW", on_closing)

    # 🔥 確保視窗可以接收輸入
    editor_window.focus_force()
    editor_window.update_idletasks()

    # 🔥 所有 UI 已建立 → 一次掃描所有可輸入文字的元件，補上滑鼠右鍵選單
    #    （此時全部 Entry/Text/Combobox 都已建好，同步綁定最可靠）
    _bind_text_menu_recursive(editor_window)

    # 🔥🔥🔥 所有 UI 建立完成 → 移除載入中遮罩 🔥🔥🔥
    _remove_loading_overlay()


